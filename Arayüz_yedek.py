import os
import sys

# Windows konsolu (cp1254) emoji/Unicode karakterleri basamadığı için
# çıktıyı UTF-8'e zorla; aksi halde print(...) ifadeleri çökertir.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

import time
import threading
import traceback
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.style import Style
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics
from openpyxl import Workbook
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    import tid_generator_logic
    import sgd_generator_logic
    import stt_generator_logic
    import dgöygö_generator_logic
    import kmtd_generator_logic
    import sitet_generator_logic
    import alt_sistem_test_logic
    import dtet_ytet_generator_logic
    import text_cleanup
    import html_generation
    import pdf_extraction
except ImportError as e:
    messagebox.showerror(
        "Modül Hatası",
        f"Gerekli bir modül yüklenemedi: {e}\nLütfen programı yeniden kurun veya bağımlılıkları kontrol edin."
    )
    sys.exit(1)
    
start1_time = time.time()

def pre_process_files(file_paths, status_callback=None):
    all_chunks = []
    
    for file_path in file_paths:
        if status_callback:
            status_callback(f"Dosya işleniyor: {os.path.basename(file_path)}")
        
        chunks = tid_generator_logic.extract_book_chunks(file_path)
        
        if status_callback:
            status_callback(f"{len(chunks)} adet chunk bulundu.")
        
        all_chunks.extend(chunks)

    if not all_chunks:
        if status_callback:
            status_callback("Chunk bulunamadı.", is_error=True)
        return None, None

    if status_callback:
        status_callback(f"Toplam {len(all_chunks)} adet chunk bulundu.")
        status_callback("Chunk'lar embedding ile analiz ediliyor...\n")

    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        yerel_model_yolu = os.path.join(base_dir, "HuggingFaceEmbeddings", "all-MiniLM-L6-v2")
        embedder = HuggingFaceEmbeddings(model_name=yerel_model_yolu)
        
        embeddings = np.array(embedder.embed_documents(all_chunks))
        
        center = np.mean(embeddings, axis=0)
        similarities = [
            np.dot(e, center) / (np.linalg.norm(e) * np.linalg.norm(center))
            for e in embeddings
        ]
        sorted_indices = np.argsort(similarities)[::-1]
        
        return all_chunks, sorted_indices

    except Exception as e:
        if status_callback:
            status_callback(f"Embedding Hatası: {str(e)}", is_error=True)
        return None, None

class TIDGeneratorApp:
    def __init__(self, master):
        self.master = master
        master.title("EHSİM Elektronik Harp Sistemleri Müh. Tic. A.Ş.")
        master.geometry("1130x950")
                                                            
        self.style = Style(theme="litera")
        target_bg = self.style.colors.light
        
        self.master.configure(bg=target_bg)
        self.style.configure("TLabel", background=target_bg)
        self.style.configure("TCheckbutton", background=target_bg)
        self.style.configure("secondary.TLabel", background=target_bg)
        self.style.configure("TFrame", background=target_bg) 

        self.dark_blue = "#0052cc"
        self.style.configure("primary.TLabel", foreground=self.dark_blue, background=target_bg)
        self.style.configure("primary.TButton", background=self.dark_blue, foreground="white")
        self.style.configure("primary.Outline.TButton", foreground=self.dark_blue, bordercolor=self.dark_blue)
        self.style.configure("primary.TCombobox", fieldbackground=target_bg, background=target_bg, foreground=self.dark_blue)
        self.style.configure("success.TButton", background=self.style.colors.success, foreground="white")
        self.style.configure("Thin.Vertical.TScrollbar", width=8, arrowsize=10)
        self.style.configure("Black.TCheckbutton", foreground="black", background=target_bg)
        
        try:
            self.style.configure("info.TButton", background="#17a2b8", foreground="white")
        except:
            pass

        self.file_paths = []
        self.template_file_path = None
        self.entry_widgets = {}
        
        self.last_generated_output = ""   
        self.raw_output_cache = ""        
        
        self.tree_data = {}
        self.flat_data = {}
        self.checkbox_vars = {}
        
        self.last_tid_list = []
        self.last_sgd_list = []
        self.last_stt_list = []
        self.last_dgoygo_list = []
        self.last_sitet_list = []
        self.last_alt_sistem_test_list = []
        self.last_dtet_ytet_list = []

        # --- Ana yatay yerleşim: SOLDA mevcut form (kaydırılabilir), SAĞDA Copilot sohbet ---
        left_container = ttk.Frame(master, style="light")
        left_container.pack(side="left", fill="both", expand=True)

        self.canvas = tk.Canvas(left_container, bg=target_bg, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(left_container, orient="vertical", command=self.canvas.yview,
                                       style="Thin.Vertical.TScrollbar")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.inner_frame = ttk.Frame(self.canvas, padding=20, style="light")
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.inner_frame, anchor="nw")
        self.inner_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.master.bind("<MouseWheel>", self._on_mousewheel)

        self._create_header()
        self._create_input_widgets(self.inner_frame)
        self._create_output_widgets(self.inner_frame)
        self._create_buttons(self.inner_frame)

        # SAĞ TARAF: Doküman Copilot sohbet paneli
        self._create_chat_panel(master, target_bg)

    def _create_header(self):
        header_frame = ttk.Frame(self.inner_frame, style="light")
        header_frame.pack(fill="x", pady=(5, 10))

        ttk.Label(
            header_frame,
            text="AI ile V Modeldeki Teknik Dokümanların Üretimi",
            font=("Segoe UI", 18, "bold"),
            foreground=self.dark_blue,
            style="primary.TLabel"
        ).pack(side="left", pady=(0, 10), anchor='w', expand=True, fill='x')

        logo_path = os.path.join(os.path.dirname(__file__), "ehsim logo.png")
        self.ehsim_logo = self._load_logo(logo_path)
        if self.ehsim_logo:
            logo_label = ttk.Label(header_frame, image=self.ehsim_logo, style="light")
            logo_label.image = self.ehsim_logo
            logo_label.pack(side="right", pady=(0, 10), padx=(10, 0), anchor='e')

        icon_path = os.path.join(os.path.dirname(__file__), "icon.png")
        if os.path.exists(icon_path):
            try:
                icon = tk.PhotoImage(file=icon_path)
                self.master.iconphoto(True, icon)
            except Exception as e:
                print(f"Icon yüklenemedi: {e}")

    def _load_logo(self, path):
        try:
            if not os.path.exists(path):
                raise FileNotFoundError
            img = Image.open(path).resize((100, 88), Image.Resampling.LANCZOS)
            return ImageTk.PhotoImage(img)
        except Exception as e:
            print(f"Logo yüklenemedi: {e}")
            return None

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_textarea_mousewheel(self, event):
        self.output_text_area.yview_scroll(int(-1 * (event.delta / 120)), "units")
        return "break"

    def _create_input_widgets(self, parent_frame):
        def create_input_row(parent, label_text, entry_key,
                             is_file_selector=False, is_count_entry=False,
                             is_cross=False, is_template_selector=False):
            frame = ttk.Frame(parent, style="light")
            frame.pack(fill="x", pady=8)

            ttk.Label(
                frame, text=f"{label_text}:", font=("Segoe UI", 10),
                foreground="#333333",
                width=39, anchor="w"
            ).pack(side="left")

            if is_file_selector:
                lbl = ttk.Label(frame, text="Seçilmedi", relief="flat",
                                style="secondary.TLabel", width=20, anchor="w", padding=(5, 5))
                lbl.pack(side="left", expand=True, fill="x", padx=5)
                self.entry_widgets[entry_key] = lbl
                ttk.Button(frame, text="...", command=self.select_files,
                           style="primary.Outline.TButton", width=3).pack(side="left", padx=(5, 0))

            elif is_template_selector:
                lbl = ttk.Label(frame, text="Seçilmedi", relief="flat",
                                style="secondary.TLabel", width=20, anchor="w", padding=(5, 5))
                lbl.pack(side="left", expand=True, fill="x", padx=5)
                self.entry_widgets[entry_key] = lbl
                ttk.Button(frame, text="...", command=self.select_template_file,
                           style="primary.Outline.TButton", width=3).pack(side="left", padx=(5, 0))

            elif is_count_entry:
                entry = ttk.Entry(frame, width=5, validate="key")
                entry['validatecommand'] = (entry.register(self.validate_number), '%P')
                entry.pack(side="left", padx=(0, 5))
                self.entry_widgets[entry_key] = entry
                entry.insert(0, "0")

            elif is_cross:
                lbl = ttk.Label(frame, text="X", font=("Segoe UI", 10, "bold"),
                                foreground="#FF0000", width=5)
                lbl.pack(side="left", padx=(0, 5))
                self.entry_widgets[entry_key] = lbl

            else:
                entry = ttk.Entry(frame)
                entry.pack(side="left", fill="x", padx=15, expand=True)
                self.entry_widgets[entry_key] = entry

        create_input_row(parent_frame, "Proje İsmi", "proje_ismi")
        create_input_row(parent_frame, "Girdi Dosyaları (PDF/TXT)", "proje_bilesenleri", is_file_selector=True)
        create_input_row(parent_frame, "Şablon Dosyası (.docx)", "template_file", is_template_selector=True)

        ttk.Label(
            parent_frame,
            text="TEKNİK DOKÜMANLAR:",
            font=("Segoe UI", 10, "bold"),
            foreground=self.dark_blue,
            style="primary.TLabel"
        ).pack(pady=(15, 3), anchor="w")

        labels_frame = ttk.Frame(parent_frame, style="light")
        labels_frame.pack(pady=(5, 10), anchor="w", fill="x")
        ttk.Label(
            labels_frame,
            text="Gereksinim Dokümanları ve Madde Sayısı",
            font=("Segoe UI", 10, "bold", "underline"),
            foreground=self.dark_blue
        ).pack(side="left", padx=(0, 40))
        ttk.Label(
            labels_frame,
            text="Test Dokümanları",
            font=("Segoe UI", 10, "bold", "underline"),
            foreground=self.dark_blue
        ).pack(side="left", padx=(45, 0))

        docs_frame = ttk.Frame(parent_frame, style="light", padding=10)
        docs_frame.pack(fill="x")
        left_col = ttk.Frame(docs_frame, style="light")
        right_col = ttk.Frame(docs_frame, style="light")
        left_col.pack(side="left", expand=True, fill="x")
        right_col.pack(side="right", expand=True, fill="x")

        create_input_row(left_col, "Kullanıcı Gereksinimi (User Requirement)", "teknik_ister", is_count_entry=True)
        create_input_row(left_col, "Sistem Gereksinimi (System Requirements)", "sistem_gereksinimi", is_count_entry=True)
        create_input_row(left_col, "Alt Sistem Gereksinimleri (Subsystem Requirements)", "sistem_tanimlama_testi", is_count_entry=True)
        create_input_row(left_col, "Donanım-Yazılım Geliştirme (Hardware/Software Design)", "donanim_yazilim_gelistirme_ozeti", is_count_entry=True)

        label_width = 39
        label_font = ("Segoe UI", 10)
        label_color = "#333333"

        kmtd_frame = ttk.Frame(right_col, style="light")
        kmtd_frame.pack(fill="x", pady=15)
        ttk.Label(kmtd_frame, text="Kabul Testi (Acceptance Test):", font=label_font, foreground=label_color,
                  width=label_width, anchor="w").pack(side="left")
        self.checkbox_vars["generate_kmtd"] = tk.BooleanVar(value=True)
        ttk.Checkbutton(kmtd_frame, text="Üret", variable=self.checkbox_vars["generate_kmtd"],
                        style="Black.TCheckbutton").pack(side="left", padx=(0, 5))

        sitet_frame = ttk.Frame(right_col, style="light")
        sitet_frame.pack(fill="x", pady=15)
        ttk.Label(sitet_frame, text="Sistem İşletme Test Tanımı (SITET):", font=label_font, foreground=label_color,
                  width=label_width, anchor="w").pack(side="left")
        self.checkbox_vars["generate_sitet"] = tk.BooleanVar(value=True)
        ttk.Checkbutton(sitet_frame, text="Üret", variable=self.checkbox_vars["generate_sitet"],
                        style="Black.TCheckbutton").pack(side="left", padx=(0, 5))

        alt_test_frame = ttk.Frame(right_col, style="light")
        alt_test_frame.pack(fill="x", pady=15)
        ttk.Label(alt_test_frame, text="Alt Sistem Testi (Subsystem Testing):", font=label_font, foreground=label_color,
                  width=label_width, anchor="w").pack(side="left")
        self.checkbox_vars["generate_alt_sistem_testi"] = tk.BooleanVar(value=True)
        ttk.Checkbutton(alt_test_frame, text="Üret", variable=self.checkbox_vars["generate_alt_sistem_testi"],
                        style="Black.TCheckbutton").pack(side="left", padx=(0, 5))

        dtet_frame = ttk.Frame(right_col, style="light")
        dtet_frame.pack(fill="x", pady=15)
        ttk.Label(dtet_frame, text="Donanım-Yazılım Testi (Hardware-Software Testing):", font=label_font, foreground=label_color,
                  width=label_width, anchor="w").pack(side="left")
        self.checkbox_vars["generate_dtet_ytet"] = tk.BooleanVar(value=True)
        ttk.Checkbutton(dtet_frame, text="Üret", variable=self.checkbox_vars["generate_dtet_ytet"],
                        style="Black.TCheckbutton").pack(side="left", padx=(0, 5))

    def validate_number(self, P):
        return P.isdigit() or P == ""

    def _create_output_widgets(self, parent_frame):
        ttk.Label(
            parent_frame,
            text="İSTENEN ÇIKTI TÜRÜ:",
            font=("Segoe UI", 10, "bold"),
            foreground=self.dark_blue,
            style="primary.TLabel"
        ).pack(pady=(5, 8), anchor="w")

        output_frame = ttk.Frame(parent_frame, style="light", padding=5)
        output_frame.pack(fill="x")
        ttk.Label(output_frame, text="Çıktı Formatı:").pack(side="left")
        self.format_combo = ttk.Combobox(
            output_frame,
            values=["txt", "pdf", "excel", "html", "docx", "şablon"],
            state="readonly",
            width=8,
            style="primary.TCombobox"
        )
        self.format_combo.set("pdf")
        self.format_combo.pack(side="left", padx=(5, 20))
        ttk.Label(output_frame, text="Durum/Çıktı Konsolu:").pack(side="left")

        console_frame = ttk.Frame(parent_frame, style="light")
        console_frame.pack(pady=(10, 15), fill="both", expand=True)

        self.text_scrollbar = ttk.Scrollbar(console_frame, orient="vertical")
        self.text_scrollbar.pack(side="right", fill="y")

        self.output_text_area = tk.Text(
            console_frame,
            height=12,
            relief="solid",
            borderwidth=1,
            state=tk.DISABLED,
            yscrollcommand=self.text_scrollbar.set,
            font=("Segoe UI", 10)
        )
        self.output_text_area.pack(side="left", fill="both", expand=True)
        self.text_scrollbar.config(command=self.output_text_area.yview)
        self.output_text_area.bind("<MouseWheel>", self._on_textarea_mousewheel)

    def _create_buttons(self, parent_frame):
        button_frame = ttk.Frame(parent_frame, style="light")
        button_frame.pack(side="bottom", fill="x", pady=20)

        self.reset_button = ttk.Button(
            button_frame,
            text="İptal / Sıfırla",
            command=self.reset_app,
            bootstyle="danger", 
            width=15
        )
        self.reset_button.pack(side="right", padx=(10, 0), pady=5)

        self.download_docs_button = ttk.Button(
            button_frame,
            text="Dokümanları İndir",
            command=self.download_docs,
            style="primary.TButton"
        )
        self.download_docs_button.pack(side="right", pady=5)

        self.classify_button = ttk.Button(
            button_frame,
            text="Dokümanları Sınıflandır",
            command=self.start_classification,
            style="primary.TButton",  
            width=21
        )
        self.classify_button.pack(side="right", padx=10, pady=5)

        self.create_docs_button = ttk.Button(
            button_frame,
            text="Dokümanları Üret",
            command=self.start_generation,
            style="primary.TButton"
        )
        self.create_docs_button.pack(side="right", padx=10, pady=5)

    def reset_app(self):
        cevap = messagebox.askyesno("Sıfırla", "Tüm seçimler, yüklenen dosyalar ve üretilen veriler silinecek.\nEmin misiniz?")
        if not cevap:
            return

        self.file_paths = []
        self.template_file_path = None
        self.entry_widgets["proje_bilesenleri"].config(text="Seçilmedi")
        self.entry_widgets["template_file"].config(text="Seçilmedi")
        self.entry_widgets["proje_ismi"].delete(0, tk.END)
        
        sayac_alanlari = ["teknik_ister", "sistem_gereksinimi", "sistem_tanimlama_testi", "donanim_yazilim_gelistirme_ozeti"]
        for key in sayac_alanlari:
            self.entry_widgets[key].delete(0, tk.END)
            self.entry_widgets[key].insert(0, "0")

        for var in self.checkbox_vars.values():
            var.set(True)

        self.last_generated_output = ""
        self.raw_output_cache = ""
        self.tree_data.clear()
        self.flat_data.clear()
        self.last_tid_list = []      
        self.last_sgd_list = []
        self.last_stt_list = []
        self.last_dgoygo_list = []
        self.last_sitet_list = []
        self.last_alt_sistem_test_list = []
        self.last_dtet_ytet_list = []

        self.output_text_area.config(state=tk.NORMAL)
        self.output_text_area.delete("1.0", tk.END)
        self.output_text_area.config(state=tk.DISABLED)

        self.create_docs_button.config(state=tk.NORMAL, text="Dokümanları Üret", style="primary.TButton")
        self.download_docs_button.config(state=tk.NORMAL)
        self.classify_button.config(state=tk.NORMAL, text="Dokümanları Sınıflandır")

        messagebox.showinfo("Bilgi", "Uygulama başarıyla sıfırlandı.")

    def select_files(self):
        paths = filedialog.askopenfilenames(
            title="Girdi Dosyalarını Seç",
            filetypes=[("PDF ve TXT dosyaları", "*.pdf *.txt")]
        )
        if paths:
            self.file_paths = list(paths)
            names = [os.path.basename(p) for p in self.file_paths]
            display = ", ".join(names) if len(names) <= 3 else f"{len(names)} dosya seçildi"
            self.entry_widgets["proje_bilesenleri"].config(text=display)
            self.update_status_text(f"Seçilen dosya(lar): {display}", clear=True)

    def select_template_file(self):
        path = filedialog.askopenfilename(
            title="Şablon Dosyasını Seç (docx)",
            filetypes=[("DOCX dosyaları", "*.docx"), ("Tüm Dosyalar", "*.*")]
        )
        if path:
            self.template_file_path = path
            self.entry_widgets["template_file"].config(text=os.path.basename(path))
            self.update_status_text(f"Seçilen şablon: {os.path.basename(path)}")
        else:
            self.template_file_path = None
            self.entry_widgets["template_file"].config(text="Seçilmedi")

    def update_status_text(self, message, is_error=False, is_complete=False, clear=False):
        def _inner():
            self.output_text_area.config(state=tk.NORMAL)
            if clear:
                self.output_text_area.delete("1.0", tk.END)

            tag = "msg"
            if is_error:
                tag = "error"
                self.output_text_area.tag_config(tag, foreground=self.style.colors.danger,
                                                 font=("Segoe UI", 10, "bold"))
                self.output_text_area.insert(tk.END, f"\n[HATA] {message}", tag)
            elif is_complete:
                tag = "complete"
                self.output_text_area.tag_config(tag, foreground=self.style.colors.success,
                                                 font=("Segoe UI", 10, "bold"))
                self.output_text_area.insert(tk.END, f"\n{message}", tag)
            else:
                self.output_text_area.insert(tk.END, f"\n{message}", tag)

            self.output_text_area.see(tk.END)
            self.output_text_area.config(state=tk.DISABLED)
        self.master.after(0, _inner)

    # V-Modelindeki doküman türleri: (flat_data type, başlık, bacak)
    # 'req' = sol bacak (gereksinim), 'test' = sağ bacak (doğrulama)
    VMODEL_SECTIONS = [
        ("TID",       "Kullanıcı Gereksinimi (User Requirement)",                "req"),
        ("KMTD",      "Kabul Testi (Acceptance Test)",               "test"),
        ("SGD",       "Sistem Gereksinimi (System Requirements)",            "req"),
        ("SITET",     "Sistem İşletme Test Tanımı (SITET)",                  "test"),
        ("STT",       "Alt Sistem Gereksinimleri (Subsystem Requirements)",  "req"),
        ("AST",       "Alt Sistem Testi (Subsystem Test)",                   "test"),
        ("DGÖ-YGÖ",   "Donanım-Yazılım Geliştirme (Hardware/Software Design)", "req"),
        ("DTET-YTET", "Donanım-Yazılım Testi (Hardware-Software Testing)",   "test"),
    ]

    def _register_pdf_font(self):
        """Türkçe karakterleri destekleyen bir font ailesi kaydeder; (normal, kalın) döner."""
        candidates = [
            ("VModelFont", "VModelFont-Bold",
             r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        ]
        for reg_name, bold_name, reg_path, bold_path in candidates:
            try:
                if os.path.exists(reg_path):
                    pdfmetrics.registerFont(TTFont(reg_name, reg_path))
                    if os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont(bold_name, bold_path))
                    else:
                        bold_name = reg_name
                    return reg_name, bold_name
            except Exception:
                continue
        # Yedek: reportlab ile gelen Bitstream Vera (Türkçe destekli)
        try:
            import reportlab
            fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
            pdfmetrics.registerFont(TTFont("VModelFont", os.path.join(fonts_dir, "Vera.ttf")))
            pdfmetrics.registerFont(TTFont("VModelFont-Bold", os.path.join(fonts_dir, "VeraBd.ttf")))
            return "VModelFont", "VModelFont-Bold"
        except Exception:
            return "Helvetica", "Helvetica-Bold"

    def _generate_vmodel_pdf(self, path, proje_ismi):
        """flat_data'yı kullanarak V-Model yapısında biçimli bir PDF üretir."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        reg, bold = self._register_pdf_font()

        title_style = ParagraphStyle("vm_title", fontName=bold, fontSize=20,
                                     textColor=colors.HexColor("#1F3864"), leading=24, spaceAfter=4)
        sub_style = ParagraphStyle("vm_sub", fontName=reg, fontSize=11,
                                   textColor=colors.HexColor("#555555"), leading=15)
        legend_style = ParagraphStyle("vm_legend", fontName=reg, fontSize=9,
                                      textColor=colors.HexColor("#555555"), leading=13)
        sec_style = ParagraphStyle("vm_sec", fontName=bold, fontSize=13,
                                   textColor=colors.white, leading=16)
        cell_style = ParagraphStyle("vm_cell", fontName=reg, fontSize=9, leading=12)
        cellid_style = ParagraphStyle("vm_cellid", fontName=bold, fontSize=9, leading=12,
                                      textColor=colors.HexColor("#1F3864"))
        head_style = ParagraphStyle("vm_head", fontName=bold, fontSize=9, leading=12,
                                    textColor=colors.white)

        # Renkler: sol bacak (gereksinim) mavi, sağ bacak (test) yeşil tonları
        REQ_COLOR = colors.HexColor("#1F3864")
        REQ_LIGHT = colors.HexColor("#D6E0F0")
        TEST_COLOR = colors.HexColor("#375623")
        TEST_LIGHT = colors.HexColor("#E2EFDA")

        story = []
        story.append(Paragraph("V-MODEL TEKNİK DOKÜMAN PAKETİ", title_style))
        story.append(Paragraph(f"Proje: <b>{proje_ismi}</b>", sub_style))
        story.append(Paragraph(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}", sub_style))
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1F3864")))
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            "<font color='#1F3864'>■</font> Sol Bacak: Gereksinim Dokümanları &nbsp;&nbsp;&nbsp; "
            "<font color='#375623'>■</font> Sağ Bacak: Test / Doğrulama Dokümanları", legend_style))
        story.append(Spacer(1, 12))

        # Sayfa genişliği ~ 170mm kullanılabilir
        col_widths = [26 * mm, 118 * mm, 26 * mm]
        any_section = False

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_section = True

            is_req = (leg == "req")
            head_color = REQ_COLOR if is_req else TEST_COLOR
            light_color = REQ_LIGHT if is_req else TEST_LIGHT
            bound_header = "Kaynak / Bağlı" if is_req else "Doğruladığı Madde"

            # Bölüm başlığı bandı
            sec_band = Table([[Paragraph(f"{title}  ({len(items)} madde)", sec_style)]],
                             colWidths=[sum(col_widths)])
            sec_band.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), head_color),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(sec_band)

            # Tablo verisi
            data = [[Paragraph("ID", head_style),
                     Paragraph("Açıklama", head_style),
                     Paragraph(bound_header, head_style)]]
            for d in items:
                data.append([
                    Paragraph(str(d.get("ID", "")), cellid_style),
                    Paragraph(str(d.get("content", "")), cell_style),
                    Paragraph(str(d.get("bound_to", "-")), cell_style),
                ])

            tbl = Table(data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), head_color),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_color]),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 16))

        if not any_section:
            # flat_data boşsa ham metni sığdır
            story.append(Paragraph("Üretilen doküman verisi bulunamadı. Ham çıktı:", sub_style))
            story.append(Spacer(1, 6))
            for line in (self.last_generated_output or "").split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), cell_style))

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=18 * mm, bottomMargin=18 * mm,
            title=f"{proje_ismi} - V-Model Doküman Paketi"
        )
        doc.build(story)

    # ================================================================== #
    #  DOKÜMAN COPILOT — sağ panel, tek madde revizyonu (grounding'li)    #
    # ================================================================== #

    # --- Kıdemli Sistem Mühendisi system prompt (paydaş kuralları 1-4) ---
    COPILOT_SYSTEM_PROMPT = (
        "ROL: Sen kıdemli bir Sistem Mühendisisin (INCOSE / sistem mühendisliği "
        "standartlarına hâkim). Görevin, sana verilen TEK bir gereksinim veya test "
        "maddesini kullanıcının isteğine göre yeniden yazmaktır. Aşağıdaki kurallara "
        "İSTİSNASIZ uyarsın:\n"
        "1) SADE ve SAF: Yanıtın YALNIZCA revize edilmiş madde metnidir. 'Tabii ki', "
        "'İşte güncellenmiş hâli' gibi laf kalabalığı, giriş/kapanış cümlesi, açıklama, "
        "yorum, başlık, ID, numara, madde işareti, yıldız veya tırnak KESİNLİKLE YOK. "
        "Tek paragraf, en fazla 2-3 cümle.\n"
        "2) BİRİMLER: Tüm sayısal değerleri metrik/teknik birimlerle ver "
        "(ms, s, Hz, kHz, MHz, dB, dBm, m, km, %, °C). Birimsiz çıplak sayı bırakma.\n"
        "3) TEKNİK DİL: Günlük dil kullanma; sistem mühendisliği terminolojisi kullan. "
        "Cümleler net, ölçülebilir ve doğrulanabilir olmalı ('sistem ... -malıdır/-melidir', "
        "'doğrulanmalıdır', 'karşılanmalıdır').\n"
        "4) DSB KURALI (SAYI UYDURMAK KESİNLİKLE YASAK): Bir sayısal değeri YALNIZCA şu iki "
        "durumda yazabilirsin: (a) kullanıcı o sayıyı isteğinde AÇIKÇA vermişse, ya da "
        "(b) sayı zaten MEVCUT METİN'de varsa. Kullanıcı 'bir değer/süre/eşik/mesafe ekle' "
        "diyor ama sayının kendisini SÖYLEMİYORSA, o sayıyı TAHMİN ETME — makul/tipik bir "
        "değer bile olsa UYDURMA. Bunun yerine değerin geçtiği yere harfi harfine 'DSB' yaz "
        "(DSB = Daha Sonra Belirlenecek), birimi koru. "
        "Örnek: kullanıcı 'tespit süresi ekle' dedi ama süreyi vermediyse → "
        "'Sistem, hedefleri DSB ms içinde tespit etmelidir.'\n"
        "ÖRNEKLER (uy):\n"
        "- MEVCUT: 'Sistem, hedefleri tespit etmelidir.' | İSTEK: 'tespit süresi ekle' "
        "(kullanıcı sayı VERMEDİ) → DOĞRU: 'Sistem, hedefleri DSB ms içinde tespit etmelidir.' | "
        "YANLIŞ: 'Sistem, hedefleri 100 ms içinde tespit etmelidir.' (100 UYDURMA, yasak!).\n"
        "- MEVCUT: 'Sistem çalışmalıdır.' | İSTEK: 'çalışma sıcaklığı -20 ile +55 °C olsun' "
        "(değer VERİLDİ) → DOĞRU: 'Sistem, -20 ile +55 °C sıcaklık aralığında çalışmalıdır.'\n"
        "ÇIKTI: Yalnızca revize edilmiş madde metni. Başka HİÇBİR ŞEY yazma."
    )

    # DSB ripple'ının yayılacağı test tipi maddeler
    TEST_TYPES = ("KMTD", "SITET", "AST", "DTET-YTET")
    # Revize edilince bağlı testleri yeniden üretilecek gereksinim tipleri
    REQ_TYPES = ("TID", "SGD", "STT", "DGÖ-YGÖ")

    def _create_chat_panel(self, master, target_bg):
        """Sağ tarafta üretilen maddeleri revize etmek için sohbet paneli."""
        panel = ttk.Frame(master, style="light", width=370)
        panel.pack(side="right", fill="y")
        panel.pack_propagate(False)

        head = ttk.Frame(panel, style="light")
        head.pack(fill="x", padx=10, pady=(12, 4))
        ttk.Label(head, text="🤖 Doküman Copilot", font=("Segoe UI", 13, "bold"),
                  foreground=self.dark_blue, style="primary.TLabel").pack(side="left")

        ttk.Label(panel,
                  text="Üretilen bir maddeyi revize ettir. Örn:\n«UR-004'ü daha teknik yaz, güvenlik standardı ekle»",
                  font=("Segoe UI", 8), foreground="#666", justify="left",
                  wraplength=340, style="secondary.TLabel").pack(fill="x", padx=10, pady=(0, 6))

        hist_frame = ttk.Frame(panel, style="light")
        hist_frame.pack(fill="both", expand=True, padx=10)
        sb = ttk.Scrollbar(hist_frame, orient="vertical")
        sb.pack(side="right", fill="y")
        self.chat_history = tk.Text(hist_frame, wrap="word", relief="solid", borderwidth=1,
                                    state=tk.DISABLED, font=("Segoe UI", 9), yscrollcommand=sb.set)
        self.chat_history.pack(side="left", fill="both", expand=True)
        sb.config(command=self.chat_history.yview)
        self.chat_history.tag_config("user", foreground="#0052cc", font=("Segoe UI", 9, "bold"))
        self.chat_history.tag_config("bot", foreground="#1a7f37")
        self.chat_history.tag_config("err", foreground="#c1121f")
        self.chat_history.tag_config("info", foreground="#666", font=("Segoe UI", 8, "italic"))

        entry_frame = ttk.Frame(panel, style="light")
        entry_frame.pack(fill="x", padx=10, pady=10)
        self.chat_entry = ttk.Entry(entry_frame, font=("Segoe UI", 10))
        self.chat_entry.pack(side="left", fill="x", expand=True)
        self.chat_entry.bind("<Return>", lambda e: self._chat_send())
        self.chat_send_btn = ttk.Button(entry_frame, text="Gönder", style="primary.TButton",
                                        command=self._chat_send, width=8)
        self.chat_send_btn.pack(side="left", padx=(6, 0))

        self._chat_append("Merhaba! Önce dokümanları üret, sonra bir maddeyi bana revize ettir. "
                          "Örn: «SR-002'yi daha ölçülebilir yap».", "info")

    def _chat_append(self, text, tag="bot"):
        def _inner():
            self.chat_history.config(state=tk.NORMAL)
            prefix = {"user": "\n👤 Sen:\n", "bot": "\n🤖 Copilot:\n", "err": "\n⚠️ "}.get(tag, "\n")
            self.chat_history.insert(tk.END, prefix + text + "\n", tag)
            self.chat_history.see(tk.END)
            self.chat_history.config(state=tk.DISABLED)
        self.master.after(0, _inner)

    def _find_target_id(self, msg):
        """Mesajdan flat_data'daki bir madde ID'sini yakalar (Türkçe/nokta duyarsız)."""
        def norm(s):
            s = s.upper()
            for a, b in (("İ", "I"), ("Ö", "O"), ("Ğ", "G"), ("Ü", "U"),
                         ("Ş", "S"), ("Ç", "C"), ("I", "I")):
                s = s.replace(a, b)
            return s
        nmsg = norm(msg)
        # Uzun ID'leri (DTET-YTET-001 gibi) önce dene ki kısa parçalar yanlış eşleşmesin
        for key in sorted(self.flat_data.keys(), key=len, reverse=True):
            if norm(key) in nmsg:
                return key
        return None

    def _chat_send(self):
        msg = self.chat_entry.get().strip()
        if not msg:
            return
        if not self.flat_data:
            self._chat_append("Henüz üretilmiş doküman yok. Önce 'Dokümanları Üret' ile üretim yap.", "err")
            return
        self.chat_entry.delete(0, tk.END)
        self._chat_append(msg, "user")
        self.chat_send_btn.config(state=tk.DISABLED)
        threading.Thread(target=self._chat_worker, args=(msg,), daemon=True).start()

    def _chat_worker(self, msg):
        try:
            from llm_handler import call_gemma3_api
            target = self._find_target_id(msg)
            if not target:
                self._chat_append("Hangi maddeyi revize edeceğimi bulamadım. Lütfen ID yaz "
                                  "(ör: UR-004, SR-002).", "err")
                return
            old = self.flat_data[target].get("content", "")
            self._chat_append(f"{target} revize ediliyor...", "info")

            # DSB kararını PYTHON'da deterministik ver (4B model 'değer verildi mi' ayrımını
            # güvenilir yapamıyor). 3 durum: (1) değer isteniyor+sayı var → kullan,
            # (2) değer isteniyor+sayı yok → DSB, (3) sadece ton/ifade değişikliği → sayıları koru.
            import re as _re
            _VALUE_KW = ("değer", "sayı", "sayısal", "süre", "eşik", "sınır", "menzil", "mesafe",
                         "sıcaklık", "hız", "frekans", "oran", "gecikme", "tolerans", "doğruluk",
                         "hassasiyet", "kapasite", "aralık", "metrik", "ölçülebilir", "birim",
                         "limit", "seviye", "band", "bant", "voltaj", "akım", "güç", "boyut", "ağırlık")
            # ÖNEMLİ: "SR-001" gibi ID'lerin içindeki rakamları (001) "kullanıcı sayı verdi"
            # sanmamak için, sayı kontrolünden ÖNCE mesajdan ID'leri temizle.
            _msg_no_id = _re.sub(_re.escape(target), "", msg, flags=_re.I) if target else msg
            _msg_no_id = _re.sub(r"[A-Za-zÇĞİÖŞÜçğıöşü]{2,}(?:-[A-Za-zÇĞİÖŞÜçğıöşü]+)*-\d+", "", _msg_no_id)
            _has_number = bool(_re.search(r"\d", _msg_no_id))
            _asks_value = any(k in msg.lower() for k in _VALUE_KW)
            if _has_number:
                dsb_directive = ("NOT: Kullanıcı somut sayısal değer(ler) verdi. Bu değerleri AYNEN "
                                 "kullan. Bu istekte 'DSB' YAZMA.")
            elif _asks_value:
                dsb_directive = ("NOT: Kullanıcı bir sayısal değer/eşik istiyor ama sayının kendisini "
                                 "VERMEDİ. O değerin yerine MUTLAKA 'DSB' yaz (birimi koru) — sayı UYDURMA.")
            else:
                dsb_directive = ("NOT: Bu istek yalnızca ifade/ton değişikliğidir. Mevcut metindeki TÜM "
                                 "sayıları ve değerleri AYNEN KORU; yeni sayı ekleme, 'DSB' YAZMA.")
            prompt = (
                f"MADDE ID: {target}\n"
                f"MEVCUT METİN: \"{old}\"\n\n"
                f"KULLANICI İSTEĞİ: {msg}\n\n"
                f"{dsb_directive}\n\n"
                "Bu maddeyi yukarıdaki nota ve sistem mühendisliği kurallarına göre revize et. "
                "Yanıt olarak SADECE revize edilmiş madde metnini ver."
            )
            resp = call_gemma3_api(prompt, max_tokens=280, temperature=0.1,
                                   system_message=self.COPILOT_SYSTEM_PROMPT)
            if not resp or not resp.strip():
                self._chat_append(f"{target} için model cevap vermedi. LM Studio açık ve model yüklü mü?", "err")
                return
            new = self._clean_revision(resp)
            if "DSB" in new.upper():
                new = text_cleanup.dsb_temizle(new)   # DSB varsa çelişen uydurma sayıları temizle
            self._apply_revision(target, old, new)

            # Ana konsolda da GÖRÜNÜR yap (kullanıcı genelde sol tarafa bakıyor)
            self.update_status_text(f"\n━━━ COPILOT · {target} GÜNCELLENDİ ━━━", is_complete=True)
            self.update_status_text(f"ESKİ: {old}")
            self.update_status_text(f"YENİ: {new}", is_complete=True)

            # --- DEĞİŞİKLİK ETKİ ANALİZİ: bir GEREKSİNİM değişince, ona bağlı test(ler)i
            #     yeni gereksinim metnine göre YENİDEN ÜRET (kullanıcı tercihi).
            #     Gereksinimde DSB varsa üretilen teste de DSB notu eklenir. ---
            extra = ""
            if self.flat_data[target].get("type") in self.REQ_TYPES:
                self._chat_append(f"{target} değişti → bağlı test(ler) yeniden üretiliyor...", "info")
                affected = self._ripple_regenerate(target)
                if affected:
                    extra = ("\n🔗 Bağlı test(ler) yeni gereksinime göre güncellendi: "
                             + ", ".join(affected))
            self._chat_append(f"✅ {target} güncellendi:\n{new}{extra}\n\n"
                              "(İndirdiğinde pdf/excel/html/word çıktısına yansır.)", "bot")
        except Exception as e:
            self._chat_append(f"Hata: {e}", "err")
        finally:
            self.master.after(0, lambda: self.chat_send_btn.config(state=tk.NORMAL))

    def _clean_revision(self, text):
        """4B modelin eklediği başlık/etiket/markdown/liste artıklarını temizler, tek paragraf yapar."""
        import re
        text = (text or "").strip().strip('"').strip()
        # markdown vurgu ve işaretlerini kaldır
        text = re.sub(r'[*_`#]+', '', text)
        # uydurma standart kodlarını (örn. "(GS-005)") her yerden temizle
        text = re.sub(r'\s*\((?:GS|STD|MIL-?STD|DS|TS)[-\s]?\d+\)', '', text, flags=re.I)

        label = re.compile(r'^(madde\s*id|madde|revizyon|revize|id|çıktı|cikti|cevap|output|sonuç|sonuc|not)\b[\s:–-]*', re.I)
        verb = re.compile(r'(malı|meli|olmalı|etmeli|dır|dir|dur|dür)\b', re.I)

        kept = []
        for ln in text.split("\n"):
            s = ln.strip().lstrip("-•·*").strip()
            if not s:
                continue
            low = s.lower().replace("̇", "")  # Türkçe İ combining-dot'u temizle
            if label.match(low):
                # etiketi kırp; geriye anlamlı bir gereksinim cümlesi kalıyorsa tut, yoksa (başlık) at
                stripped = label.sub('', s).strip()
                if verb.search(stripped.lower()) and len(stripped.split()) >= 4:
                    kept.append(stripped)
                continue
            kept.append(s)

        out = " ".join(kept)
        out = re.sub(r'\s+', ' ', out)
        out = re.sub(r'\s+([.,;:])', r'\1', out)   # " ." -> "."
        return out.strip(' -–•:')

    def _sync_item_text(self, item_id, old_text, new_text):
        """Bir maddenin yeni metnini ham çıktı metnine ve ilgili last_*_list'e yansıtır."""
        # 1) ham metin (txt çıktısı / konsol tutarlılığı)
        if old_text and old_text in (self.last_generated_output or ""):
            self.last_generated_output = self.last_generated_output.replace(old_text, new_text)
        # 2) ilgili last_*_list (sınıflandırma tutarlılığı) — best effort
        t = self.flat_data.get(item_id, {}).get("type", "")
        list_map = {
            "TID":       (self.last_tid_list, "TID_ID", "TID_Aciklama"),
            "SGD":       (self.last_sgd_list, "SGD_ID", "SGD_Aciklama"),
            "STT":       (self.last_stt_list, "STT_ID", "STT_Aciklama"),
            "SITET":     (self.last_sitet_list, "SITET_ID", "SITET_Aciklama"),
            "AST":       (self.last_alt_sistem_test_list, "AST_ID", "AST_Aciklama"),
            "DGÖ-YGÖ":   (self.last_dgoygo_list, "ID", "Aciklama"),
            "DTET-YTET": (self.last_dtet_ytet_list, "ID", "Aciklama"),
        }
        if t in list_map:
            lst, idk, txtk = list_map[t]
            for it in lst:
                if it.get(idk) == item_id:
                    it[txtk] = new_text
                    break

    def _apply_revision(self, target, old, new):
        """Revizyonu flat_data + ham metin + ilgili last_*_list üzerinde uygular."""
        self.flat_data[target]["content"] = new     # tüm çıktıların (pdf/excel/html/docx) kaynağı
        self._sync_item_text(target, old, new)
        self.update_status_text(f"[Copilot] {target} güncellendi.", is_complete=True)

    def _ripple_regenerate(self, requirement_id):
        """
        DEĞİŞİKLİK ETKİ ANALİZİ (TAM KASKAD): Bir gereksinim revize edilince, ona bağlı
        HEM alt gereksinimleri HEM testleri yeni metne göre yeniden üretir. Alt gereksinimler
        için işlem özyinelemeli olarak aşağı iner (SR→SSR→HSD ve her birinin testi).
        Üst maddede 'DSB' varsa, türeyen tüm maddelerde de ilgili değer DSB olur (uydurma yok).
        """
        child_req_gen = {
            "TID": sgd_generator_logic.generate_sgd_from_ur,
            "SGD": stt_generator_logic.generate_subsystem_req_from_sgd,
            "STT": dgöygö_generator_logic.generate_dgöygö_from_ssr,
        }
        test_gen = {
            "TID": kmtd_generator_logic.generate_kmtd_from_tid,
            "SGD": sitet_generator_logic.generate_sitet_from_sgd,
            "STT": alt_sistem_test_logic.generate_subsystem_test,
            "DGÖ-YGÖ": dtet_ytet_generator_logic.generate_dtet_ytet_from_dgoygo,
        }
        dsb_kural = (" [KURAL: Yukarıdaki maddede 'DSB' geçen değer belirsizdir; türeteceğin "
                     "maddede de o değer için sayı UYDURMA, yerine 'DSB' yaz.]")
        note = (" (Not: Üst gereksinimde DSB bulunduğu için bu maddede de ilgili değer DSB'dir.)")
        affected, visited = [], set()

        def kaskad(parent_id):
            if parent_id in visited:
                return
            visited.add(parent_id)
            parent = self.flat_data.get(parent_id, {})
            ptype = parent.get("type")
            ptext = parent.get("content", "")
            dsb = "DSB" in (ptext or "").upper()
            girdi = (ptext + dsb_kural) if dsb else ptext
            cocuklar = [(k, v) for k, v in list(self.flat_data.items())
                        if v.get("bound_to") == parent_id]
            for iid, it in cocuklar:
                itype = it.get("type")
                if itype in self.TEST_TYPES:
                    gen_fn = test_gen.get(ptype)
                elif itype in self.REQ_TYPES:
                    gen_fn = child_req_gen.get(ptype)
                else:
                    gen_fn = None
                if not gen_fn:
                    continue
                try:
                    raw = gen_fn(girdi, "Proje")
                except Exception as e:
                    self.update_status_text(f"[Copilot] {iid} yeniden üretilemedi: {e}", is_error=True)
                    continue
                if not raw:
                    continue
                yeni = text_cleanup.temizle(raw, test=(itype in self.TEST_TYPES))
                if not yeni:
                    continue
                if dsb:
                    # DSB ile çelişen uydurma sayı/örnekleri temizle ('(örneğin 100g)', '100g DSB'...)
                    yeni = text_cleanup.dsb_temizle(yeni)
                    if "DSB" not in yeni.upper():   # hiç DSB kalmadıysa notu ekle
                        yeni = (yeni.rstrip() + note).strip()
                eski = it.get("content", "")
                it["content"] = yeni
                self._sync_item_text(iid, eski, yeni)
                affected.append(iid)
                if itype in self.REQ_TYPES:   # alt gereksinimin de altını güncelle
                    kaskad(iid)

        kaskad(requirement_id)
        if affected:
            self.update_status_text(
                f"[Copilot] {requirement_id} değişti → bağlı madde(ler) güncellendi: "
                f"{', '.join(affected)}", is_complete=True)
        return affected

    def _ripple_dsb(self, requirement_id):
        """
        Paydaş kuralı 5 — İZLENEBİLİRLİK YAYILIMI (ripple effect):
        Bir gereksinim maddesi 'DSB' içerecek şekilde revize edildiğinde, ona BAĞLI olan
        test maddelerini (KMTD/SITET/AST/DTET-YTET) bulur ve onların test kriterine de
        'DSB' notu ekler. Bağ ilişkisi: test maddesinin bound_to == gereksinim ID'si.
        Etkilenen test ID'lerinin listesini döndürür.
        """
        note = " (Not: İlgili gereksinimde DSB bulunduğu için test kriteri de DSB'dir.)"
        affected = []
        for item_id, d in list(self.flat_data.items()):
            if d.get("type") in self.TEST_TYPES and d.get("bound_to") == requirement_id:
                old_c = d.get("content", "")
                if "DSB" in old_c.upper():
                    continue  # zaten DSB'li → tekrar ekleme
                new_c = (old_c.rstrip() + note).strip()
                d["content"] = new_c
                self._sync_item_text(item_id, old_c, new_c)
                affected.append(item_id)
        if affected:
            self.update_status_text(
                f"[Copilot] DSB yayılımı → {requirement_id} bağlı test(ler)i güncellendi: "
                f"{', '.join(affected)}", is_complete=True)
        return affected

    # ------------------------------------------------------------------ #
    #  EXCEL / HTML / WORD çıktıları (hepsi flat_data + VMODEL_SECTIONS)  #
    # ------------------------------------------------------------------ #
    def _export_excel(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile çok sayfalı, biçimli Excel üretir."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def style_header(ws, ncols, color="1F3864"):
            fill = PatternFill("solid", fgColor=color)
            for c in range(1, ncols + 1):
                cell = ws.cell(row=1, column=c)
                cell.fill = fill
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

        wb = Workbook()
        # --- Kapak ---
        info = wb.active
        info.title = "Proje Bilgisi"
        info.append(["V-MODEL İZLENEBİLİRLİK RAPORU"])
        info.append(["Proje", proje_ismi])
        info.append(["Tarih", datetime.now().strftime("%d.%m.%Y %H:%M")])
        info.append(["Toplam Madde", len(self.flat_data)])
        info["A1"].font = Font(bold=True, size=14, color="1F3864")
        info.column_dimensions["A"].width = 20
        info.column_dimensions["B"].width = 50

        leg_tr = {"req": "Gereksinim", "test": "Test"}

        # --- İzlenebilirlik Matrisi ---
        ws = wb.create_sheet("Izlenebilirlik Matrisi")
        ws.append(["ID", "Tür", "Bacak", "Açıklama", "Bağlı Olduğu"])
        for type_key, title, leg in self.VMODEL_SECTIONS:
            for d in self.flat_data.values():
                if d.get("type") == type_key:
                    ws.append([d.get("ID", ""), type_key, leg_tr.get(leg, leg),
                               d.get("content", ""), d.get("bound_to", "Yok")])
        style_header(ws, 5)
        for col, w in zip("ABCDE", [16, 12, 12, 80, 22]):
            ws.column_dimensions[col].width = w
        for r in range(2, ws.max_row + 1):
            ws.cell(row=r, column=4).alignment = Alignment(wrap_text=True, vertical="top")

        # --- Her tür için ayrı sayfa ---
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            safe = type_key.replace("/", "-")[:31]
            sh = wb.create_sheet(safe)
            sh.append([f"{type_key} ID", "Açıklama", "Bağlı Olduğu"])
            for d in items:
                sh.append([d.get("ID", ""), d.get("content", ""), d.get("bound_to", "Yok")])
            style_header(sh, 3, "1F3864" if leg == "req" else "375623")
            for col, w in zip("ABC", [16, 90, 22]):
                sh.column_dimensions[col].width = w
            for r in range(2, sh.max_row + 1):
                sh.cell(row=r, column=2).alignment = Alignment(wrap_text=True, vertical="top")

        wb.save(path)

    def _export_html(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile bağımsız, biçimli HTML raporu üretir."""
        import html as _html
        esc = lambda s: _html.escape(str(s))
        leg_color = {"req": "#1F3864", "test": "#375623"}
        leg_light = {"req": "#D6E0F0", "test": "#E2EFDA"}

        p = []
        p.append("<!DOCTYPE html><html lang='tr'><head><meta charset='utf-8'>")
        p.append(f"<title>{esc(proje_ismi)} - V-Model Doküman Paketi</title>")
        p.append("<style>"
                 "body{font-family:'Segoe UI',Arial,sans-serif;margin:24px;color:#222;background:#f7f9fc}"
                 "h1{color:#1F3864;margin-bottom:0}.sub{color:#666;margin:2px 0 14px}"
                 "table{border-collapse:collapse;width:100%;margin-bottom:22px;background:#fff;"
                 "box-shadow:0 1px 3px rgba(0,0,0,.08)}"
                 "th,td{border:1px solid #ccc;padding:7px 9px;text-align:left;vertical-align:top;font-size:14px}"
                 ".sec{color:#fff;padding:8px 12px;font-weight:bold;font-size:15px;margin-top:8px;border-radius:3px 3px 0 0}"
                 "td.id{font-weight:bold;white-space:nowrap}"
                 ".legend span{display:inline-block;margin-right:18px}"
                 ".chip{display:inline-block;width:12px;height:12px;border-radius:2px;margin-right:5px;vertical-align:middle}"
                 "</style></head><body>")
        p.append("<h1>V-Model Teknik Doküman Paketi</h1>")
        p.append(f"<div class='sub'>Proje: <b>{esc(proje_ismi)}</b> &nbsp;|&nbsp; "
                 f"{datetime.now().strftime('%d.%m.%Y %H:%M')} &nbsp;|&nbsp; Toplam {len(self.flat_data)} madde</div>")
        p.append("<div class='legend'>"
                 "<span><span class='chip' style='background:#1F3864'></span>Sol Bacak: Gereksinim</span>"
                 "<span><span class='chip' style='background:#375623'></span>Sağ Bacak: Test / Doğrulama</span></div><br>")

        any_sec = False
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_sec = True
            col = leg_color.get(leg, "#333"); light = leg_light.get(leg, "#eee")
            bound_h = "Kaynak / Bağlı" if leg == "req" else "Doğruladığı Madde"
            p.append(f"<div class='sec' style='background:{col}'>{esc(title)} ({len(items)} madde)</div>")
            p.append(f"<table><tr style='background:{col};color:#fff'>"
                     f"<th>ID</th><th>Açıklama</th><th>{bound_h}</th></tr>")
            for i, d in enumerate(items):
                bg = "#ffffff" if i % 2 == 0 else light
                p.append(f"<tr style='background:{bg}'><td class='id'>{esc(d.get('ID',''))}</td>"
                         f"<td>{esc(d.get('content',''))}</td><td>{esc(d.get('bound_to','-'))}</td></tr>")
            p.append("</table>")
        if not any_sec:
            p.append("<p><i>Üretilen doküman verisi bulunamadı.</i></p>")
        p.append("</body></html>")

        with open(path, "w", encoding="utf-8") as f:
            f.write("".join(p))

    def _export_docx(self, path, proje_ismi, template_path=None):
        """Kurumsal .docx şablonunu doldurur (veya sıfırdan Word üretir). Başarı=True."""
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Eksik Kütüphane",
                "Word (.docx) çıktısı için 'python-docx' gerekli.\n\n"
                "Terminalde şunu çalıştırın:\n    pip install python-docx")
            return False

        doc = Document(template_path) if template_path else Document()

        if template_path:
            mapping = {"{{PROJE_ADI}}": proje_ismi, "{{TARIH}}": datetime.now().strftime("%d.%m.%Y")}
            for para in doc.paragraphs:
                for k, v in mapping.items():
                    if k in para.text:
                        para.text = para.text.replace(k, v)
            doc.add_page_break()
        else:
            doc.add_heading("V-MODEL TEKNİK DOKÜMAN PAKETİ", level=0)
            doc.add_paragraph(f"Proje: {proje_ismi}")
            doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            doc.add_heading(title, level=2)
            t = doc.add_table(rows=1, cols=3)
            try:
                t.style = "Light Grid Accent 1"
            except Exception:
                t.style = "Table Grid"
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = "ID", "Açıklama", "Bağlı Olduğu"
            for d in items:
                c = t.add_row().cells
                c[0].text = str(d.get("ID", ""))
                c[1].text = str(d.get("content", ""))
                c[2].text = str(d.get("bound_to", "Yok"))
        doc.save(path)
        return True

    def download_docs(self):
        fmt = self.format_combo.get().lower()
        if fmt != "şablon" and not self.last_generated_output and not self.flat_data:
            messagebox.showwarning("Uyarı", "İndirilecek içerik yok. Önce doküman üretin.")
            return

        proje_ismi = self.entry_widgets["proje_ismi"].get().strip().replace(" ", "_") or "yapay_zeka_dokumanlar"

        if fmt == "excel":
            ext, ftypes = ".xlsx", [("Excel Workbook", "*.xlsx")]
        elif fmt in ("şablon", "docx"):
            ext, ftypes = ".docx", [("Word Document", "*.docx")]
        else:
            ext, ftypes = f".{fmt}", [(f"{fmt.upper()} files", f"*.{fmt}")]

        init_name = f"{proje_ismi}_Dokumanlar{ext}"
        path = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=ftypes + [("All files", "*.*")],
            initialfile=init_name
        )
        if not path:
            return

        try:
            if fmt == "txt":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.last_generated_output or "")

            elif fmt == "pdf":
                self._generate_vmodel_pdf(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "excel":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Excel'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_excel(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "html":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "HTML'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_html(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt in ("şablon", "docx"):
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Word'e yerleştirilecek veri yok. Önce doküman üretin.")
                    return
                tmpl = self.template_file_path if fmt == "şablon" else None
                if fmt == "şablon" and not tmpl:
                    messagebox.showerror("Hata", "Şablon dosyası seçilmedi (.docx).")
                    return
                ok = self._export_docx(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje", tmpl)
                if not ok:
                    return

            messagebox.showinfo("Başarılı", f"Dosya kaydedildi:\n{path}")
        except Exception as e:
            messagebox.showerror("Hata", f"Kaydetme hatası: {e}")
            self.update_status_text(f"Kaydetme hatası: {e}", is_error=True)

    def start_generation(self):
        if not self.file_paths:
            messagebox.showerror("Hata", "Girdi dosyalarını seçin.")
            return

        try:
            proje_ismi = self.entry_widgets["proje_ismi"].get().strip()
            if not proje_ismi:
                raise ValueError("Proje İsmi boş bırakılamaz.")

            doc_counts = {
                "max_tids": int(self.entry_widgets["teknik_ister"].get() or 0),
                "max_sgds": int(self.entry_widgets["sistem_gereksinimi"].get() or 0),
                "max_stts": int(self.entry_widgets["sistem_tanimlama_testi"].get() or 0),
                "max_dgöygös": int(self.entry_widgets["donanim_yazilim_gelistirme_ozeti"].get() or 0),
            }

            doc_flags = {
                "generate_kmtd": self.checkbox_vars["generate_kmtd"].get(),
                "generate_sitet": self.checkbox_vars["generate_sitet"].get(),
                "generate_alt_sistem_testi": self.checkbox_vars["generate_alt_sistem_testi"].get(),
                "generate_dtet_ytet": self.checkbox_vars["generate_dtet_ytet"].get(),
            }

            if sum(v > 0 for v in doc_counts.values()) == 0 and not any(doc_flags.values()):
                raise ValueError("En az bir doküman sayısı veya test seçilmeli.")

        except ValueError as e:
            messagebox.showerror("Hata", f"Geçersiz giriş: {e}")
            return

        self.last_generated_output = ""
        self.tree_data.clear()
        self.flat_data.clear()
        self.update_status_text("--- YAPAY ZEKA ÜRETİMİ BAŞLADI ---\n", clear=True)

        self.create_docs_button.config(state=tk.DISABLED, text="İŞLENİYOR...", style="success.TButton")
        self.download_docs_button.config(state=tk.DISABLED)
        self.classify_button.config(state=tk.DISABLED)

        thread = threading.Thread(
            target=self.run_ai_process,
            args=(self.file_paths, doc_counts, doc_flags,
                  self.format_combo.get().lower(), proje_ismi)
        )
        thread.start()

    def run_ai_process(self, file_paths, doc_counts, doc_flags, output_format, proje_ismi):
        total_start_time = time.time()

        # NOT: Alt Sistem Gereksinimleri (max_stts) artık kaynak dosyadan (chunk)
        # değil, üretilen SGD listesinden türetildiği için bu kontrole dahil değildir.
        kaynak_dokuman_gerekli = (
            doc_counts["max_tids"] > 0 or
            doc_counts["max_sgds"] > 0 or
            doc_counts["max_dgöygös"] > 0
        )

        all_chunks = None
        sorted_indices = None

        if kaynak_dokuman_gerekli:
            all_chunks, sorted_indices = pre_process_files(file_paths, self.update_status_text)
            
            if not all_chunks:
                self.update_status_text("Hata: Kaynak dosyalardan veri alınamadı, işlem durduruluyor.", is_error=True)
                self.master.after(0, lambda: self._reset_buttons_state())
                return

        try:
            if doc_counts.get("max_tids", 0) > 0:
                self.update_status_text("Kullanıcı Gereksinimi üretimi başlıyor...")
                result = tid_generator_logic.run_generation_logic(
                    file_paths=None,
                    max_tids=doc_counts["max_tids"],
                    output_format=output_format,
                    project_name=proje_ismi,
                    status_callback=self.update_status_text,
                    precomputed_chunks=all_chunks,
                    precomputed_indices=sorted_indices
                )
                if result.get("result"):
                    self.last_tid_list = result.get("tid_list", [])
                    output = f"\n--- KULLANICI GEREKSİNİMİ (User Requirement) --- {proje_ismi} ---\n\n"
                    for item in self.last_tid_list:
                        output += f"{item['TID_ID']} | {item['TID_Aciklama']}\n"
                        self.flat_data[item['TID_ID']] = {
                            'type': 'TID', 'bound_to': 'Yok',
                            'ID': item['TID_ID'], 'content': item['TID_Aciklama']
                        }
                    self.last_generated_output += output
                    self.update_status_text(output, is_complete=True)

            if doc_flags["generate_kmtd"]:
                if self.last_tid_list:
                    self.update_status_text("KMTD üretimi başlıyor...")
                    try:
                        result = kmtd_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_tid_list,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            kmtd_list = result.get("kmtd_list", [])
                            output = f"\n\n--- KABUL MUAYENE TESTİ (Acceptance Test) --- {proje_ismi} ---\n\n"
                            for item in kmtd_list:
                                output += f"{item['KMTD_ID']} | {item['KMTD_Aciklama']}\n"
                                self.flat_data[item['KMTD_ID']] = {
                                    'type': 'KMTD', 'bound_to': item['Bound_TID'],
                                    'ID': item['KMTD_ID'], 'content': item['KMTD_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"KMTD Hatası: {e}", is_error=True)

            if doc_counts.get("max_sgds", 0) > 0:
                self.update_status_text("SGD üretimi başlıyor...")
                if self.last_tid_list:
                    # İZLENEBİLİRLİK: her Kullanıcı Gereksiniminden (UR) türeyen SGD (Bound_TID)
                    result = sgd_generator_logic.run_generation_from_requirements(
                        requirement_list=self.last_tid_list,
                        max_sgds=doc_counts["max_sgds"],
                        project_name=proje_ismi,
                        status_callback=self.update_status_text
                    )
                else:
                    # UR yoksa eski yöntem: doğrudan kaynak dosyadan (genel bağ)
                    result = sgd_generator_logic.run_generation_logic(
                        file_paths=None,
                        max_sgds=doc_counts["max_sgds"],
                        output_format=output_format,
                        project_name=proje_ismi,
                        status_callback=self.update_status_text,
                        precomputed_chunks=all_chunks,
                        precomputed_indices=sorted_indices
                    )
                if result.get("result"):
                    self.last_sgd_list = result.get("sgd_list", [])
                    output = f"\n\n--- SİSTEM GEREKSİNİMİ (System Requirements) --- {proje_ismi} ---\n\n"
                    for item in self.last_sgd_list:
                        output += f"{item['SGD_ID']} | {item['SGD_Aciklama']}\n"
                        self.flat_data[item['SGD_ID']] = {
                            'type': 'SGD', 'bound_to': item.get('Bound_TID', 'TID-Genel'),
                            'ID': item['SGD_ID'], 'content': item['SGD_Aciklama']
                        }
                    self.last_generated_output += output
                    self.update_status_text(output, is_complete=True)

            if doc_flags["generate_sitet"]:
                if self.last_sgd_list:
                    self.update_status_text("SITET üretimi başlıyor...")
                    try:
                        if 'sitet_generator_logic' in sys.modules:
                            result = sitet_generator_logic.run_generation_from_requirements(
                                requirement_list=self.last_sgd_list,
                                project_name=proje_ismi,
                                status_callback=self.update_status_text
                            )
                            if result.get("result"):
                                self.last_sitet_list = result.get("sitet_list", [])
                                output = f"\n\n--- Sistem İşletme Test Tanımı (SITET) LİSTESİ --- {proje_ismi} ---\n\n"
                                for item in self.last_sitet_list:
                                    output += f"{item['SITET_ID']} | {item['SITET_Aciklama']}\n"
                                    self.flat_data[item['SITET_ID']] = {
                                        'type': 'SITET', 'bound_to': item.get('Bound_SGD', 'SGD'),
                                        'ID': item['SITET_ID'], 'content': item['SITET_Aciklama']
                                    }
                                self.last_generated_output += output
                                self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"SITET Hatası: {e}", is_error=True)

            if doc_counts.get("max_stts", 0) > 0:
                self.update_status_text("Alt Sistem Gereksinimleri üretimi başlıyor...")
                if not self.last_sgd_list:
                    self.update_status_text(
                        "Alt Sistem Gereksinimleri için önce Sistem Gereksinimi (SGD) üretmelisiniz.",
                        is_error=True
                    )
                else:
                    try:
                        result = stt_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_sgd_list,
                            max_stts=doc_counts["max_stts"],
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            self.last_stt_list = result.get("stt_list", [])
                            output = f"\n\n--- ALT SİSTEM GEREKSİNİMLERİ (Subsystem Requirements) --- {proje_ismi} ---\n\n"
                            for item in self.last_stt_list:
                                output += f"{item['STT_ID']} | {item['STT_Aciklama']}\n"
                                self.flat_data[item['STT_ID']] = {
                                    'type': 'STT', 'bound_to': item.get('Bound_SGD', 'SGD'),
                                    'ID': item['STT_ID'], 'content': item['STT_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"Alt Sistem Gereksinimleri Hatası: {e}", is_error=True)

            if doc_flags["generate_alt_sistem_testi"]:
                if self.last_stt_list:
                    self.update_status_text("Alt Sistem Testi üretimi başlıyor...")
                    try:
                        result = alt_sistem_test_logic.run_generation_from_requirements(
                            requirement_list=self.last_stt_list,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            self.last_alt_sistem_test_list = result.get("ast_list", [])
                            output = f"\n\n--- ALT SİSTEM TESTİ (Subsystem Test) LİSTESİ --- {proje_ismi} ---\n\n"
                            for item in self.last_alt_sistem_test_list:
                                output += f"{item['AST_ID']} | {item['AST_Aciklama']}\n"
                                self.flat_data[item['AST_ID']] = {
                                    'type': 'AST', 'bound_to': item.get('Bound_STT', 'ASG'),
                                    'ID': item['AST_ID'], 'content': item['AST_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"Alt Sistem Testi Hatası: {e}", is_error=True)

            if doc_counts.get("max_dgöygös", 0) > 0:
                self.update_status_text("DGÖ-YGÖ üretimi başlıyor...")
                try:
                    if self.last_stt_list:
                        # İZLENEBİLİRLİK: her Alt Sistem Gereksiniminden (SSR) türeyen DGÖ-YGÖ (Bound_STT)
                        result = dgöygö_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_stt_list,
                            max_items=doc_counts["max_dgöygös"],
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                    else:
                        # SSR yoksa eski yöntem: doğrudan kaynak dosyadan (genel bağ)
                        result = dgöygö_generator_logic.run_generation_logic(
                            file_paths=None,
                            max_items=doc_counts["max_dgöygös"],
                            output_format=output_format,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text,
                            precomputed_chunks=all_chunks,
                            precomputed_indices=sorted_indices
                        )
                    if result.get("result"):
                        self.last_dgoygo_list = result.get("dgoygo_list", [])
                        output = f"\n\n--- DONANIM-YAZILIM GELİŞTİRME (Hardware/Software Design) --- {proje_ismi} ---\n\n"
                        for item in self.last_dgoygo_list:
                            output += f"{item['ID']} | {item['Aciklama']}\n"
                            self.flat_data[item['ID']] = {
                                'type': 'DGÖ-YGÖ', 'bound_to': item.get('Bound_STT', 'Genel'),
                                'ID': item['ID'], 'content': item['Aciklama']
                            }
                        self.last_generated_output += output
                        self.update_status_text(output, is_complete=True)
                except Exception as e:
                    self.update_status_text(f"DGÖ-YGÖ Modül Hatası: {e}", is_error=True)

            if doc_flags["generate_dtet_ytet"] and self.last_dgoygo_list:
                try:
                    if 'dtet_ytet_generator_logic' in sys.modules:
                        result = dtet_ytet_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_dgoygo_list,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            self.last_dtet_ytet_list = result.get("test_list", [])
                            output = f"\n\n--- DONANIM-YAZILIM TESTİ (Hardware-Software Testing) --- {proje_ismi} ---\n\n"
                            for item in self.last_dtet_ytet_list:
                                output += f"{item['ID']} | {item['Aciklama']}\n"
                                self.flat_data[item['ID']] = {
                                    'type': 'DTET-YTET', 'bound_to': item.get('Bound_DGÖYGÖ', 'Genel'),
                                    'ID': item['ID'], 'content': item['Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                except Exception as e:
                    self.update_status_text(f"DTET-YTET Hatası: {e}", is_error=True)

            total_end_time = time.time()
            total_duration = total_end_time - total_start_time
            minutes = int(total_duration // 60)
            seconds = int(total_duration % 60)
            
            if not self.last_generated_output.strip():
                self.update_status_text("Hiçbir doküman üretilemedi.", is_error=True)
            else:
                final_msg = f"Toplam Geçen Süre: {minutes} dakika {seconds} saniye\n--- YAPAY ZEKA ÜRETİMİ TAMAMLANDI ---\n"
                self.update_status_text(final_msg, is_complete=True)
                self.raw_output_cache = self.last_generated_output  

        except Exception as e:
            self.update_status_text(f"KRİTİK HATA: {e}", is_error=True)
            traceback.print_exc()
        finally:
            self.master.after(0, lambda: self._reset_buttons_state())

    def _reset_buttons_state(self):
        self.create_docs_button.config(state=tk.NORMAL, text="Dokümanları Üret", style="primary.TButton")
        self.download_docs_button.config(state=tk.NORMAL)
        self.classify_button.config(state=tk.NORMAL, text="Dokümanları Sınıflandır")
        self.reset_button.config(state=tk.NORMAL)

    def start_classification(self):
        if not self.last_tid_list and not self.last_sgd_list and not self.last_stt_list and not self.last_dgoygo_list:
            messagebox.showwarning("Uyarı", "Sınıflandırılacak veri (Kullanıcı Gereksinimi, SGD, STT veya DGÖ) yok. Önce doküman üretin.")
            return
        threading.Thread(target=self.run_classification_process, daemon=True).start()

    def run_classification_process(self):
        """
        GÜNCELLENMİŞ VERSİYON:
        1. Sınıflandırılan metni anında ekrana (Yeşil) basar.
        2. Arka planda indirme dosyası için birleştirir.
        3. DGÖ-YGÖ Desteği eklendi.
        """
        self.classify_button.config(text="SINIFLANDIRILIYOR...", state=tk.DISABLED)
        self.create_docs_button.config(state=tk.DISABLED)
        
        self.update_status_text("\n--- SINIFLANDIRMA İŞLEMİ BAŞLATILDI ---")
        final_full_text = self.raw_output_cache + "\n\n" + "="*50 + "\n   SINIFLANDIRILMIŞ DOKÜMANLAR RAPORU   \n" + "="*50 + "\n"

        if self.last_tid_list:
            self.update_status_text("⏳ Kullanıcı Gereksinimi Listesi Sınıflandırılıyor...")
            try:
                result_tid = tid_generator_logic.classify_existing_tid_list(
                    self.last_tid_list,
                    status_callback=None 
                )
                if result_tid["result"]:
                    header = "\n\n--- SINIFLANDIRILMIŞ KULLANICI GEREKSİNİMİ (User Requirement) ---\n"
                    content = result_tid["classified_text"]
                    full_block = header + content + "\n"
                    
                    final_full_text += full_block
                    
                    self.update_status_text(full_block, is_complete=True)
                else:
                    self.update_status_text(f"⚠️ Kullanıcı Gereksinimi Sınıflandırma başarısız: {result_tid.get('message')}", is_error=True)
            except Exception as e:
                self.update_status_text(f"⚠️ Kullanıcı Gereksinimi Hatası: {e}", is_error=True)


        if self.last_sgd_list:
            self.update_status_text("⏳ SGD Listesi Sınıflandırılıyor...")
            try:
                result_sgd = sgd_generator_logic.classify_sgd_requirements(
                    self.last_sgd_list,
                    status_callback=None
                )
                if result_sgd["result"]:
                    header = "\n\n--- SINIFLANDIRILMIŞ SİSTEM GEREKSİNİMİ (System Requirements) ---\n"
                    content = result_sgd["classified_text"]
                    full_block = header + content + "\n"

                    final_full_text += full_block

                    self.update_status_text(full_block, is_complete=True)
                else:
                    self.update_status_text(f"⚠️ SGD Sınıflandırma başarısız: {result_sgd.get('message')}", is_error=True)
            except Exception as e:
                self.update_status_text(f"⚠️ SGD Hatası: {e}", is_error=True)

        if self.last_stt_list:
            self.update_status_text("⏳ STT Listesi Sınıflandırılıyor...")
            try:
                result_stt = stt_generator_logic.classify_stt_requirements(
                    self.last_stt_list,
                    status_callback=None
                )
                if result_stt["result"]:

                    header = "\n\n--- SINIFLANDIRILMIŞ ALT SİSTEM GEREKSİNİMLERİ (Subsystem Requirements) ---\n"
                    content = result_stt["classified_text"]
                    full_block = header + content + "\n"

                    final_full_text += full_block

                    self.update_status_text(full_block, is_complete=True)
                else:
                    self.update_status_text(f"⚠️ STT Sınıflandırma başarısız: {result_stt.get('message')}", is_error=True)
            except Exception as e:
                self.update_status_text(f"⚠️ STT Hatası: {e}", is_error=True)


        if self.last_dgoygo_list:
            self.update_status_text("⏳ DGÖ-YGÖ Listesi Sınıflandırılıyor...")
            try:

                if 'dgöygö_generator_logic' in sys.modules:
                    result_dgo = dgöygö_generator_logic.classify_existing_dgo_list(
                        self.last_dgoygo_list,
                        status_callback=None
                    )
                    
                    if result_dgo["result"]:
                        header = "\n\n--- SINIFLANDIRILMIŞ DONANIM-YAZILIM GELİŞTİRME (Hardware/Software Design) ---\n"
                        content = result_dgo["classified_text"]
                        full_block = header + content + "\n"


                        final_full_text += full_block
                        self.update_status_text(full_block, is_complete=True)
                    else:
                        self.update_status_text(f"⚠️ DGÖ Sınıflandırma başarısız: {result_dgo.get('message')}", is_error=True)
                else:
                     self.update_status_text("⚠️ DGÖ Modülü bulunamadı.", is_error=True)

            except Exception as e:
                self.update_status_text(f"⚠️ DGÖ Sınıflandırma Hatası: {e}", is_error=True)
                traceback.print_exc()

        self.last_generated_output = final_full_text
        
        self.update_status_text("\n--- TÜM SINIFLANDIRMA İŞLEMLERİ BİTTİ ---", is_complete=True)

        self._reset_buttons_state()

if __name__ == "__main__":
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    root = ttk.Window()
    app = TIDGeneratorApp(root)
    root.mainloop()