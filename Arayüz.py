import csv
import os
import sys

# Windows konsolu (cp1254) emoji/Unicode karakterleri basamadığı için
# çıktıyı UTF-8'e zorla; aksi halde print(...) ifadeleri çökertir.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

import time
import threading
import traceback
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.style import Style
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics
from openpyxl import Workbook
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings
from app_identity import (
    APP_NAME, ICON_RELATIVE_PATH, apply_app_identity,
    prepare_process_identity, resource_path,
)

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    import tid_generator_logic
    import sgd_generator_logic
    import stt_generator_logic
    import dgöygö_generator_logic
    import kmtd_generator_logic
    import sitet_generator_logic
    import alt_sistem_test_logic
    import dtet_ytet_generator_logic
    import hardware_list_logic
    import hardware_generator_logic
    import hardware_list_ui
    import donanim_kartlari_gorsel
    import donanim_kartlari_ui
    import donanim_kartlari_yonetim
    import etki_analizi_ui
    import etki_analizi_izlenebilirlik
    import etki_analizi_entegrasyon
    import etki_analizi_simulasyon
    import etki_analizi_degisim_paketi
    import etki_analizi_degisim_raporlama
    import donanim_kartlari_algilama
    import mimari_cerceve_ui
    import text_cleanup
    import html_generation
    import pdf_extraction
except ImportError as e:
    messagebox.showerror(
        "Modül Hatası",
        f"Gerekli bir modül yüklenemedi: {e}\nLütfen programı yeniden kurun veya bağımlılıkları kontrol edin."
    )
    sys.exit(1)
    
start1_time = time.time()

def pre_process_files(file_paths, status_callback=None):
    all_chunks = []
    
    for file_path in file_paths:
        if status_callback:
            status_callback(f"Dosya işleniyor: {os.path.basename(file_path)}")
        
        chunks = tid_generator_logic.extract_book_chunks(file_path)
        
        if status_callback:
            status_callback(f"{len(chunks)} adet chunk bulundu.")
        
        all_chunks.extend(chunks)

    if not all_chunks:
        if status_callback:
            status_callback("Chunk bulunamadı.", is_error=True)
        return None, None

    if status_callback:
        status_callback(f"Toplam {len(all_chunks)} adet chunk bulundu.")
        status_callback("Chunk'lar embedding ile analiz ediliyor...\n")

    try:
        base_dir = os.path.dirname(os.path.abspath(__file__))
        yerel_model_yolu = os.path.join(base_dir, "HuggingFaceEmbeddings", "all-MiniLM-L6-v2")
        embedder = HuggingFaceEmbeddings(model_name=yerel_model_yolu)
        
        embeddings = np.array(embedder.embed_documents(all_chunks))
        
        center = np.mean(embeddings, axis=0)
        similarities = [
            np.dot(e, center) / (np.linalg.norm(e) * np.linalg.norm(center))
            for e in embeddings
        ]
        sorted_indices = np.argsort(similarities)[::-1]
        
        return all_chunks, sorted_indices

    except Exception as e:
        if status_callback:
            status_callback(f"Embedding Hatası: {str(e)}", is_error=True)
        return None, None

class TIDGeneratorApp:
    FILE_HINT = "Sürükle-bırak ya da '...' ile seç"

    def __init__(self, master):
        self.master = master
        self._app_icon = apply_app_identity(master)
        master.title(APP_NAME)
        master.geometry("1130x950")
                                                            
        self.style = Style(theme="litera")
        target_bg = self.style.colors.light
        
        self.master.configure(bg=target_bg)
        self.style.configure("TLabel", background=target_bg)
        self.style.configure("TCheckbutton", background=target_bg)
        self.style.configure("secondary.TLabel", background=target_bg)
        self.style.configure("TFrame", background=target_bg) 

        self.dark_blue = "#0052cc"
        self.style.configure("primary.TLabel", foreground=self.dark_blue, background=target_bg)
        self.style.configure("primary.TButton", background=self.dark_blue, foreground="white")
        self.style.configure("primary.Outline.TButton", foreground=self.dark_blue, bordercolor=self.dark_blue)
        self.style.configure("primary.TCombobox", fieldbackground=target_bg, background=target_bg, foreground=self.dark_blue)
        self.style.configure("success.TButton", background=self.style.colors.success, foreground="white")
        self.style.configure("Thin.Vertical.TScrollbar", width=8, arrowsize=10)
        self.style.configure("Black.TCheckbutton", foreground="black", background=target_bg)
        
        try:
            self.style.configure("info.TButton", background="#17a2b8", foreground="white")
        except:
            pass

        self.file_paths = []
        self.generated_document_paths = []
        self.template_file_path = None
        self.entry_widgets = {}
        
        self.last_generated_output = ""   
        self.raw_output_cache = ""        
        
        self.tree_data = {}
        self.flat_data = {}
        self.hardware_data = {}
        self.hardware_workspace = None
        self.hardware_cards_workspace = None
        self.impact_analysis_workspace = None
        self.mimari_cerceve_workspace = None
        self._hardware_generation_token = 0
        self._traceability_generation_token = 0
        self._hardware_catalog_generation_token = 0
        self._traceability_cancel_event = threading.Event()
        self._architecture_generation_state = "ready"
        self._architecture_generation_detail = ""
        self.last_traceability_report = None
        self.last_traceability_health = None
        self.last_hardware_catalog = None
        self.last_hardware_catalog_status = None
        self.last_hardware_impact_result = None
        self.checkbox_vars = {}

        # --- Dil (TR/EN) + Tema (aydınlık/karanlık) altyapısı ---
        self.lang = "tr"                 # "tr" | "en"
        self.dark = False                # aydınlık başla
        self._i18n = []                  # dil değişiminde yeniden etiketlenecek (widget, tr, en)
        self._theme_labels = []          # tema değişiminde yeniden renklenecek düz ttk.Label'lar
        self._theme_texts = []           # tema değişiminde yeniden renklenecek tk.Text/Canvas
        self._light_theme = "litera"
        self._dark_theme = "darkly"
        
        self.last_tid_list = []
        self.last_sgd_list = []
        self.last_stt_list = []
        self.last_dgoygo_list = []
        self.last_sitet_list = []
        self.last_alt_sistem_test_list = []
        self.last_dtet_ytet_list = []

        # --- Sürükle-Bırak (Drag & Drop) desteği: tkinterdnd2 varsa etkinleştir ---
        # Yoksa uygulama yine çalışır, sadece "..." ile tıklayarak dosya seçilir.
        self._dnd_enabled = False
        try:
            from tkinterdnd2 import TkinterDnD
            TkinterDnD._require(master)
            self._dnd_enabled = True
        except Exception as _dnd_err:
            print(f"Sürükle-bırak devre dışı (tkinterdnd2 yok): {_dnd_err}")

        # --- Ana yatay yerleşim: SOLDA mevcut form (kaydırılabilir), SAĞDA Copilot sohbet ---
        left_container = ttk.Frame(master, style="light")
        left_container.pack(side="left", fill="both", expand=True)

        self.canvas = tk.Canvas(left_container, bg=target_bg, highlightthickness=0)
        self.scrollbar = ttk.Scrollbar(left_container, orient="vertical", command=self.canvas.yview,
                                       style="Thin.Vertical.TScrollbar")
        self.canvas.configure(yscrollcommand=self.scrollbar.set)
        self.canvas.pack(side="left", fill="both", expand=True)
        self.scrollbar.pack(side="right", fill="y")

        self.inner_frame = ttk.Frame(self.canvas, padding=20, style="light")
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.inner_frame, anchor="nw")
        self.inner_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))
        self.master.bind("<MouseWheel>", self._on_mousewheel)

        self._create_header()
        self._create_input_widgets(self.inner_frame)
        self._create_output_widgets(self.inner_frame)
        self._create_buttons(self.inner_frame)

        # SAĞ TARAF: Doküman Copilot sohbet paneli
        self._create_chat_panel(master, target_bg)

    # ══════════════════════════════════════════════════════════════════
    #  DİL (TR/EN) + TEMA (Aydınlık/Karanlık) yardımcıları
    # ══════════════════════════════════════════════════════════════════
    def _t(self, tr, en):
        """Aktif dile göre metin döndürür (Türkçe karakterler kaynak string'de UTF-8)."""
        return tr if self.lang == "tr" else en

    def _L(self, parent, tr, en, translate=True, theme=True, **kw):
        """Çevrilebilir + tema-duyarlı ttk.Label oluşturur ve kaydeder."""
        w = ttk.Label(parent, text=(tr if self.lang == "tr" else en), **kw)
        if translate:
            self._i18n.append((w, tr, en))
        if theme and "foreground" in kw:
            # vurgu (başlık/mavi) mı yoksa gövde metni mi — oluşturma anında sabitle
            is_accent = (kw.get("foreground") == self.dark_blue)
            self._theme_labels.append((w, is_accent))
        return w

    def _reg_btn(self, widget, tr, en):
        """Bir butonu dil değişiminde yeniden etiketlenmek üzere kaydeder."""
        self._i18n.append((widget, tr, en))
        return widget

    def _toggle_lang(self):
        self.lang = "en" if self.lang == "tr" else "tr"
        for w, tr, en in self._i18n:
            try:
                w.config(text=(tr if self.lang == "tr" else en))
            except Exception:
                pass
        if hasattr(self, "lang_btn"):
            self.lang_btn.config(text=("EN" if self.lang == "tr" else "TR"))
        # Chat karşılama metni (Text kutusu içinde) — henüz sohbet başlamadıysa yeniden yaz
        if getattr(self, "_chat_has_convo", True) is False and hasattr(self, "chat_history"):
            self.chat_history.config(state=tk.NORMAL)
            self.chat_history.delete("1.0", tk.END)
            self.chat_history.config(state=tk.DISABLED)
            self._chat_append(self._t(*self._chat_greeting), "info")
        workspace = getattr(self, "hardware_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh_language()
        cards_workspace = getattr(self, "hardware_cards_workspace", None)
        if cards_workspace and cards_workspace.exists:
            cards_workspace.refresh_language()
        impact_workspace = getattr(self, "impact_analysis_workspace", None)
        if impact_workspace and impact_workspace.exists:
            impact_workspace.refresh_language()
        architecture_workspace = getattr(self, "mimari_cerceve_workspace", None)
        if architecture_workspace and architecture_workspace.exists:
            architecture_workspace.refresh_language()

    def _toggle_theme(self):
        self.dark = not self.dark
        self._apply_theme()
        if hasattr(self, "theme_btn"):
            self.theme_btn.config(text=("☀" if self.dark else "🌙"))

    # Elle palet (ttkbootstrap theme_use scrollbar element'ini çakıştırdığı için kullanılmaz)
    _PALETTE = {
        "light": dict(bg="#F5F6F7", surface="#FFFFFF", fg="#222222", muted="#5C666D",
                      entry_bg="#FFFFFF", entry_fg="#222222", accent="#0052cc"),
        "dark":  dict(bg="#1F2329", surface="#2B303A", fg="#E4E6EA", muted="#95A0A8",
                      entry_bg="#2B303A", entry_fg="#E8EAED", accent="#5AA0F2"),
    }

    def _apply_theme(self):
        """Aydınlık/karanlık tema uygular — kendi paletiyle (ttk theme_use kullanmaz)."""
        p = self._PALETTE["dark" if self.dark else "light"]
        bg, fg, entry_bg, entry_fg, muted = p["bg"], p["fg"], p["entry_bg"], p["entry_fg"], p["muted"]
        self.dark_blue = p["accent"]

        # Arka planı olan tüm özel stiller (kod her yerde style="light" kullanıyor)
        for st in ("TLabel", "light.TLabel", "light.TFrame", "TFrame",
                   "TCheckbutton", "Black.TCheckbutton"):
            try:
                self.style.configure(st, background=bg)
            except Exception:
                pass
        self.style.configure("light.TLabel", foreground=fg)
        self.style.configure("TCheckbutton", foreground=fg)
        self.style.configure("Black.TCheckbutton", foreground=fg, background=bg)
        self.style.configure("primary.TLabel", foreground=self.dark_blue, background=bg)
        self.style.configure("primary.TButton", background=self.dark_blue, foreground="white")
        self.style.configure("secondary.TLabel", background=entry_bg, foreground=muted)
        self.style.configure("TEntry", fieldbackground=entry_bg, foreground=entry_fg,
                             insertcolor=fg)
        self.style.configure("primary.TCombobox", fieldbackground=entry_bg,
                             background=bg, foreground=entry_fg)

        # Kök + kaydırılabilir alan
        self.master.configure(bg=bg)
        if hasattr(self, "canvas"):
            self.canvas.configure(bg=bg)

        # Kayıtlı düz etiketler (inline foreground'lu olanlar)
        for w, is_accent in self._theme_labels:
            try:
                w.configure(background=bg, foreground=(self.dark_blue if is_accent else fg))
            except Exception:
                pass

        # tk.Text ve benzeri (konsol, chat)
        for w in self._theme_texts:
            try:
                w.configure(background=p["surface"], foreground=fg, insertbackground=fg)
            except Exception:
                pass

        workspace = getattr(self, "hardware_workspace", None)
        if workspace and workspace.exists:
            workspace.apply_theme()
        cards_workspace = getattr(self, "hardware_cards_workspace", None)
        if cards_workspace and cards_workspace.exists:
            cards_workspace.apply_theme()
        impact_workspace = getattr(self, "impact_analysis_workspace", None)
        if impact_workspace and impact_workspace.exists:
            impact_workspace.apply_theme()
        architecture_workspace = getattr(self, "mimari_cerceve_workspace", None)
        if architecture_workspace and architecture_workspace.exists:
            architecture_workspace.apply_theme()

    def _show_sozluk(self):
        """Teknik sözlüğü ayrı bir pencerede açar (sozluk.py'den okur)."""
        try:
            import importlib, sozluk
            importlib.reload(sozluk)
            veri = sozluk.SOZLUK
        except Exception as e:
            messagebox.showerror(self._t("Hata", "Error"),
                                 self._t(f"Sözlük yüklenemedi: {e}", f"Glossary could not load: {e}"))
            return
        c = self.style.colors
        win = tk.Toplevel(self.master)
        win.title(self._t("Teknik Sözlük", "Technical Glossary"))
        win.geometry("620x640")
        win.configure(bg=c.bg)

        ttk.Label(win, text=self._t("📖 Teknik Sözlük", "📖 Technical Glossary"),
                  font=("Segoe UI", 14, "bold"), foreground=c.primary,
                  background=c.bg).pack(anchor="w", padx=14, pady=(12, 2))
        ttk.Label(win, text=self._t("Çıktıdaki kısaltmalar ve sistem mühendisliği terimleri.",
                                    "Abbreviations in the output and systems-engineering terms."),
                  font=("Segoe UI", 9), foreground=c.secondary, background=c.bg).pack(anchor="w", padx=14)

        frame = ttk.Frame(win)
        frame.pack(fill="both", expand=True, padx=10, pady=8)
        sb = ttk.Scrollbar(frame, orient="vertical")
        sb.pack(side="right", fill="y")
        txt = tk.Text(frame, wrap="word", font=("Segoe UI", 10), yscrollcommand=sb.set,
                      relief="solid", borderwidth=1, background=c.inputbg,
                      foreground=c.inputfg, padx=10, pady=8, spacing1=2, spacing3=4)
        txt.pack(side="left", fill="both", expand=True)
        sb.config(command=txt.yview)
        txt.tag_config("kat", foreground=c.primary, font=("Segoe UI", 11, "bold"), spacing1=10, spacing3=4)
        txt.tag_config("term", foreground=c.info, font=("Segoe UI", 10, "bold"))
        txt.tag_config("acik", foreground=c.inputfg, font=("Segoe UI", 10))
        txt.tag_config("vurgu_kat", foreground=c.success, font=("Segoe UI", 11, "bold"), spacing1=10, spacing3=4)
        txt.tag_config("ipucu", foreground=c.secondary, font=("Segoe UI", 9, "italic"), spacing3=6)

        idx = 0 if self.lang == "tr" else 1

        # ── B ÖZELLİĞİ: üretilen çıktıda geçen terimleri öne çıkar ──
        try:
            birlesik = " ".join(list(self.flat_data.keys()) +
                                [str(d.get("content", "")) for d in self.flat_data.values()])
        except Exception:
            birlesik = ""
        import re as _sr
        def _gecti(madde):
            # Kelime SINIRIYLA eşleştir (önek kabul) → 'DOA' artık 'TDOA' içinde eşleşmez.
            aramalar = madde[3] if len(madde) > 3 else [madde[0]]
            for a in aramalar:
                if a and _sr.search(r"\b" + _sr.escape(a), birlesik, _sr.I):
                    return True
            return False

        if birlesik.strip():
            gecenler = [m for maddeler in veri.values() for m in maddeler if _gecti(m)]
            txt.insert(tk.END, self._t("🔎 Bu çıktıda geçen terimler",
                                       "🔎 Terms found in this output") + "\n", "vurgu_kat")
            if gecenler:
                for madde in gecenler:
                    txt.insert(tk.END, f"  • {madde[0]}\n", "term")
                    txt.insert(tk.END, f"      {madde[1 + idx]}\n", "acik")
            else:
                txt.insert(tk.END, self._t("  (eşleşen terim bulunamadı)",
                                           "  (no matching terms found)") + "\n", "ipucu")
            txt.insert(tk.END, "\n" + "─" * 40 + "\n", "ipucu")
        else:
            txt.insert(tk.END, self._t(
                "İpucu: Doküman ürettikten sonra bu pencere, o çıktıda geçen terimleri en üstte öne çıkarır.\n",
                "Tip: After generating documents, this window highlights the terms found in that output at the top.\n"),
                "ipucu")

        # ── Tüm sözlük ──
        for kategori, maddeler in veri.items():
            txt.insert(tk.END, f"\n{kategori}\n", "kat")
            for madde in maddeler:
                txt.insert(tk.END, f"  • {madde[0]}\n", "term")
                txt.insert(tk.END, f"      {madde[1 + idx]}\n", "acik")
        txt.config(state=tk.DISABLED)

        ttk.Button(win, text=self._t("Kapat", "Close"), command=win.destroy,
                   style="primary.TButton", width=12).pack(pady=(0, 10))

    def _create_header(self):
        # ── Üst araç çubuğu: Tema · Dil · Sözlük ──
        toolbar = ttk.Frame(self.inner_frame, style="light")
        toolbar.pack(fill="x", pady=(0, 2))
        self.theme_btn = ttk.Button(toolbar, text="🌙", width=3, command=self._toggle_theme,
                                    style="primary.Outline.TButton")
        self.theme_btn.pack(side="right")
        self.lang_btn = ttk.Button(toolbar, text="EN", width=4, command=self._toggle_lang,
                                   style="primary.Outline.TButton")
        self.lang_btn.pack(side="right", padx=(0, 6))
        self.sozluk_btn = ttk.Button(toolbar, command=self._show_sozluk,
                                     style="primary.Outline.TButton")
        self._reg_btn(self.sozluk_btn, "📖 Sözlük", "📖 Glossary")
        self.sozluk_btn.pack(side="right", padx=(0, 6))

        header_frame = ttk.Frame(self.inner_frame, style="light")
        header_frame.pack(fill="x", pady=(5, 10))

        title = self._L(
            header_frame,
            "AI ile V Modeldeki Teknik Dokümanların Üretimi",
            "AI-based V-Model Technical Document Generation",
            font=("Segoe UI", 18, "bold"),
            foreground=self.dark_blue,
            style="primary.TLabel"
        )
        title.pack(side="left", pady=(0, 10), anchor='w', expand=True, fill='x')

        logo_path = resource_path(ICON_RELATIVE_PATH)
        self.ehsim_logo = self._load_logo(str(logo_path))
        if self.ehsim_logo:
            logo_label = ttk.Label(header_frame, image=self.ehsim_logo, style="light")
            logo_label.image = self.ehsim_logo
            logo_label.pack(side="right", pady=(0, 10), padx=(10, 0), anchor='e')

    def _load_logo(self, path):
        try:
            if not os.path.exists(path):
                raise FileNotFoundError
            img = Image.open(path).resize((100, 88), Image.Resampling.LANCZOS)
            return ImageTk.PhotoImage(img)
        except Exception as e:
            print(f"Logo yüklenemedi: {e}")
            return None

    def _on_mousewheel(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_textarea_mousewheel(self, event):
        self.output_text_area.yview_scroll(int(-1 * (event.delta / 120)), "units")
        return "break"

    def _create_input_widgets(self, parent_frame):
        def create_input_row(parent, label_tr, label_en, entry_key,
                             is_file_selector=False, is_count_entry=False,
                             is_cross=False, is_template_selector=False):
            frame = ttk.Frame(parent, style="light")
            frame.pack(fill="x", pady=8)

            self._L(
                frame, f"{label_tr}:", f"{label_en}:", font=("Segoe UI", 10),
                foreground="#333333", width=39, anchor="w"
            ).pack(side="left")

            if is_file_selector:
                hint = self.FILE_HINT if self._dnd_enabled else "Seçilmedi"
                lbl = ttk.Label(frame, text=hint, relief="solid", borderwidth=1,
                                style="secondary.TLabel", width=20, anchor="w", padding=(5, 5))
                lbl.pack(side="left", expand=True, fill="x", padx=5)
                self.entry_widgets[entry_key] = lbl
                ttk.Button(frame, text="...", command=self.select_files,
                           style="primary.Outline.TButton", width=3).pack(side="left", padx=(5, 0))
                # Hem etikete hem satır çerçevesine sürükle-bırak hedefi kaydet
                self._register_drop_target(lbl)
                self._register_drop_target(frame)

            elif is_template_selector:
                lbl = ttk.Label(frame, text=self._t("Seçilmedi", "Not selected"), relief="flat",
                                style="secondary.TLabel", width=20, anchor="w", padding=(5, 5))
                lbl.pack(side="left", expand=True, fill="x", padx=5)
                self.entry_widgets[entry_key] = lbl
                ttk.Button(frame, text="...", command=self.select_template_file,
                           style="primary.Outline.TButton", width=3).pack(side="left", padx=(5, 0))

            elif is_count_entry:
                entry = ttk.Entry(frame, width=5, validate="key")
                entry['validatecommand'] = (entry.register(self.validate_number), '%P')
                entry.pack(side="left", padx=(0, 5))
                self.entry_widgets[entry_key] = entry
                entry.insert(0, "0")

            elif is_cross:
                lbl = ttk.Label(frame, text="X", font=("Segoe UI", 10, "bold"),
                                foreground="#FF0000", width=5)
                lbl.pack(side="left", padx=(0, 5))
                self.entry_widgets[entry_key] = lbl

            else:
                entry = ttk.Entry(frame)
                entry.pack(side="left", fill="x", padx=15, expand=True)
                self.entry_widgets[entry_key] = entry

        create_input_row(parent_frame, "Proje İsmi", "Project Name", "proje_ismi")
        create_input_row(parent_frame, "Girdi Dosyaları (PDF/TXT)", "Input Files (PDF/TXT)", "proje_bilesenleri", is_file_selector=True)
        create_input_row(parent_frame, "Şablon Dosyası (.docx)", "Template File (.docx)", "template_file", is_template_selector=True)

        self._L(
            parent_frame, "TEKNİK DOKÜMANLAR:", "TECHNICAL DOCUMENTS:",
            font=("Segoe UI", 10, "bold"), foreground=self.dark_blue, style="primary.TLabel"
        ).pack(pady=(15, 3), anchor="w")

        labels_frame = ttk.Frame(parent_frame, style="light")
        labels_frame.pack(pady=(5, 10), anchor="w", fill="x")
        self._L(
            labels_frame, "Gereksinim Dokümanları ve Madde Sayısı",
            "Requirement Documents and Item Count",
            font=("Segoe UI", 10, "bold", "underline"), foreground=self.dark_blue
        ).pack(side="left", padx=(0, 40))
        self._L(
            labels_frame, "Test Dokümanları", "Test Documents",
            font=("Segoe UI", 10, "bold", "underline"), foreground=self.dark_blue
        ).pack(side="left", padx=(45, 0))

        docs_frame = ttk.Frame(parent_frame, style="light", padding=10)
        docs_frame.pack(fill="x")
        left_col = ttk.Frame(docs_frame, style="light")
        right_col = ttk.Frame(docs_frame, style="light")
        left_col.pack(side="left", expand=True, fill="x")
        right_col.pack(side="right", expand=True, fill="x")

        create_input_row(left_col, "Kullanıcı Gereksinimi (User Requirement)", "User Requirement", "teknik_ister", is_count_entry=True)
        create_input_row(left_col, "Sistem Gereksinimi (System Requirements)", "System Requirement", "sistem_gereksinimi", is_count_entry=True)
        create_input_row(left_col, "Alt Sistem Gereksinimleri (Subsystem Requirements)", "Subsystem Requirement", "sistem_tanimlama_testi", is_count_entry=True)

        label_width = 39
        label_font = ("Segoe UI", 10)
        label_color = "#333333"

        def test_row(tr, en, key):
            fr = ttk.Frame(right_col, style="light")
            fr.pack(fill="x", pady=15)
            self._L(fr, f"{tr}:", f"{en}:", font=label_font, foreground=label_color,
                    width=label_width, anchor="w").pack(side="left")
            self.checkbox_vars[key] = tk.BooleanVar(value=True)
            cb = ttk.Checkbutton(fr, variable=self.checkbox_vars[key], style="Black.TCheckbutton")
            self._reg_btn(cb, "Üret", "Generate")
            cb.config(text="Üret" if self.lang == "tr" else "Generate")
            cb.pack(side="left", padx=(0, 5))

        test_row("Kabul Testi (Acceptance Test)", "Acceptance Test", "generate_kmtd")
        test_row("Sistem Testi (System Test)", "System Test", "generate_sitet")
        test_row("Alt Sistem Testi (Subsystem Testing)", "Subsystem Test", "generate_alt_sistem_testi")

    def validate_number(self, P):
        return P.isdigit() or P == ""

    def _create_output_widgets(self, parent_frame):
        self._L(
            parent_frame, "İSTENEN ÇIKTI TÜRÜ:", "OUTPUT FORMAT:",
            font=("Segoe UI", 10, "bold"), foreground=self.dark_blue, style="primary.TLabel"
        ).pack(pady=(5, 8), anchor="w")

        output_frame = ttk.Frame(parent_frame, style="light", padding=5)
        output_frame.pack(fill="x")
        self._L(output_frame, "Çıktı Formatı:", "Format:", style="light").pack(side="left")
        self.format_combo = ttk.Combobox(
            output_frame,
            values=["txt", "pdf", "excel", "html", "docx", "şablon", "DOORS"],
            state="readonly",
            width=8,
            style="primary.TCombobox"
        )
        self.format_combo.set("pdf")
        self.format_combo.pack(side="left", padx=(5, 20))
        self._L(output_frame, "Durum/Çıktı Konsolu:", "Status / Output Console:", style="light").pack(side="left")

        console_frame = ttk.Frame(parent_frame, style="light")
        console_frame.pack(pady=(10, 15), fill="both", expand=True)

        self.text_scrollbar = ttk.Scrollbar(console_frame, orient="vertical")
        self.text_scrollbar.pack(side="right", fill="y")

        self.output_text_area = tk.Text(
            console_frame,
            height=12,
            relief="solid",
            borderwidth=1,
            state=tk.DISABLED,
            yscrollcommand=self.text_scrollbar.set,
            font=("Segoe UI", 10)
        )
        self.output_text_area.pack(side="left", fill="both", expand=True)
        self.text_scrollbar.config(command=self.output_text_area.yview)
        self.output_text_area.bind("<MouseWheel>", self._on_textarea_mousewheel)
        self._theme_texts.append(self.output_text_area)   # tema değişiminde renklen

    def _create_buttons(self, parent_frame):
        architecture_frame = ttk.Frame(parent_frame, style="light")
        architecture_frame.pack(side="bottom", fill="x", pady=(0, 4))
        self.architecture_button = ttk.Button(
            architecture_frame,
            command=self.open_architecture_framework_workspace,
            style="primary.Outline.TButton",
            width=24,
        )
        self._reg_btn(
            self.architecture_button,
            "Mimari Çerçeve",
            "Architecture Framework",
        )
        self.architecture_button.config(
            text=self._t("Mimari Çerçeve", "Architecture Framework")
        )
        self.architecture_button.pack(side="left", fill="x", expand=True, pady=3)

        button_frame = ttk.Frame(parent_frame, style="light")
        button_frame.pack(side="bottom", fill="x", pady=20)

        self.hardware_button = ttk.Button(
            button_frame, command=self.open_hardware_workspace,
            style="primary.Outline.TButton", width=18)
        self._reg_btn(self.hardware_button, "Donanım Listesi", "Hardware List")
        self.hardware_button.config(text=self._t("Donanım Listesi", "Hardware List"))
        self.hardware_button.pack(side="left", pady=5)

        self.hardware_cards_button = ttk.Button(
            button_frame, command=self.open_hardware_cards_workspace,
            style="primary.Outline.TButton", width=18)
        self._reg_btn(self.hardware_cards_button, "Donanım Kartları", "Hardware Cards")
        self.hardware_cards_button.config(text=self._t("Donanım Kartları", "Hardware Cards"))
        self.hardware_cards_button.pack(side="left", padx=(8, 0), pady=5)

        self.impact_analysis_button = ttk.Button(
            button_frame, command=self.open_impact_analysis_workspace,
            style="primary.Outline.TButton", width=16)
        self._reg_btn(self.impact_analysis_button, "Etki Analizi", "Impact Analysis")
        self.impact_analysis_button.config(text=self._t("Etki Analizi", "Impact Analysis"))
        self.impact_analysis_button.pack(side="left", padx=(8, 0), pady=5)

        self.reset_button = ttk.Button(
            button_frame, command=self.reset_app, bootstyle="danger", width=15)
        self._reg_btn(self.reset_button, "İptal / Sıfırla", "Cancel / Reset")
        self.reset_button.config(text=self._t("İptal / Sıfırla", "Cancel / Reset"))
        self.reset_button.pack(side="right", padx=(10, 0), pady=5)

        self.download_docs_button = ttk.Button(
            button_frame, command=self.download_docs, style="primary.TButton")
        self._reg_btn(self.download_docs_button, "Dokümanları İndir", "Download Documents")
        self.download_docs_button.config(text=self._t("Dokümanları İndir", "Download Documents"))
        self.download_docs_button.pack(side="right", pady=5)

        # --- Gereksinim Kalite Denetçisi (ayrı özellik; kaldırmak için bu buton + kalite_denetci.py sil) ---
        self.kalite_button = ttk.Button(
            button_frame, command=self.run_kalite_denetimi, style="primary.TButton", width=16)
        self._reg_btn(self.kalite_button, "Kaliteyi Denetle", "Check Quality")
        self.kalite_button.config(text=self._t("Kaliteyi Denetle", "Check Quality"))
        self.kalite_button.pack(side="right", padx=10, pady=5)

        self.create_docs_button = ttk.Button(
            button_frame, command=self.start_generation, style="primary.TButton")
        self._reg_btn(self.create_docs_button, "Dokümanları Üret", "Generate Documents")
        self.create_docs_button.config(text=self._t("Dokümanları Üret", "Generate Documents"))
        self.create_docs_button.pack(side="right", padx=10, pady=5)

    def open_hardware_workspace(self):
        """Akıllı Donanım Listesi mühendislik çalışma alanını açar."""
        workspace = getattr(self, "hardware_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh()
            workspace.focus()
            return

        try:
            self.hardware_workspace = hardware_list_ui.HardwareWorkspace(
                master=self.master,
                style=self.style,
                hardware_data=self.hardware_data,
                flat_data=self.flat_data,
                language_getter=lambda: self.lang,
                palette_getter=self._hardware_palette,
                project_name_getter=lambda: self.entry_widgets["proje_ismi"].get().strip() or "Proje",
                generate_callback=self.start_hardware_generation,
                on_close=self._on_hardware_workspace_closed,
            )
        except Exception as e:
            self.hardware_workspace = None
            messagebox.showerror(
                self._t("Donanım Listesi Hatası", "Hardware List Error"),
                self._t(
                    f"Donanım çalışma alanı açılamadı: {e}",
                    f"Hardware workspace could not be opened: {e}",
                ),
            )

    def open_hardware_cards_workspace(self):
        """Kanıtlı donanım kartı, BOM ve datasheet çalışma alanını açar."""
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh()
            workspace.focus()
            return
        try:
            self.hardware_cards_workspace = donanim_kartlari_ui.HardwareCardsWorkspace(
                master=self.master, style=self.style,
                language_getter=lambda: self.lang,
                palette_getter=self._hardware_palette,
                project_name_getter=lambda: self.entry_widgets["proje_ismi"].get().strip() or "Proje",
                catalog_getter=self._get_current_hardware_catalog,
                traceability_getter=self._get_current_traceability_report,
                rescan_callback=self._rescan_traceability_from_workspace,
                datasheet_callback=self._start_hardware_datasheet_ingestion,
                impact_callback=self._open_hardware_comparison,
                requirement_callback=self._open_hardware_requirement,
                on_close=self._on_hardware_cards_workspace_closed,
            )
            if self.last_hardware_impact_result:
                self.hardware_cards_workspace.set_simulation_result(
                    self.last_hardware_impact_result
                )
        except Exception as error:
            self.hardware_cards_workspace = None
            messagebox.showerror(
                self._t("Donanım Kartları Hatası", "Hardware Cards Error"),
                self._t(
                    f"Donanım Kartları çalışma alanı açılamadı: {error}",
                    f"Hardware Cards workspace could not be opened: {error}",
                ),
            )

    def open_impact_analysis_workspace(self):
        """Çok alternatifli Etki Analizi çalışma alanını açar."""
        workspace = getattr(self, "impact_analysis_workspace", None)
        if workspace and workspace.exists:
            workspace.focus()
            return

        try:
            self.impact_analysis_workspace = (
                etki_analizi_ui.ImpactAnalysisWorkspace(
                    master=self.master,
                    style=self.style,
                    language_getter=lambda: self.lang,
                    palette_getter=self._hardware_palette,
                    traceability_getter=self._get_current_traceability_report,
                    on_close=self._on_impact_analysis_workspace_closed,
                    traceability_update_callback=self._set_current_traceability_report,
                    traceability_rescan_callback=self._rescan_traceability_from_workspace,
                    traceability_cancel_callback=self._cancel_traceability_from_workspace,
                    project_info_getter=self._get_impact_project_info,
                    change_apply_callback=self._apply_approved_change_package,
                    simulation_result_callback=self._on_hardware_impact_result,
                    hardware_detail_callback=self._open_hardware_detail,
                )
            )
        except Exception as e:
            self.impact_analysis_workspace = None
            messagebox.showerror(
                self._t("Etki Analizi Hatası", "Impact Analysis Error"),
                self._t(
                    f"Etki Analizi çalışma alanı açılamadı: {e}",
                    f"Impact Analysis workspace could not be opened: {e}",
                ),
            )

    def open_architecture_framework_workspace(self):
        """Kanıta bağlı DoDAF/NAF Mimari Çerçeve Stüdyosu'nu açar."""
        workspace = getattr(self, "mimari_cerceve_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh()
            workspace.focus()
            return

        try:
            self.mimari_cerceve_workspace = (
                mimari_cerceve_ui.ArchitectureFrameworkWorkspace(
                    master=self.master,
                    style=self.style,
                    flat_data_getter=lambda: self.flat_data,
                    traceability_getter=self._get_current_traceability_report,
                    project_name_getter=lambda: (
                        self.entry_widgets["proje_ismi"].get().strip() or "Proje"
                    ),
                    language_getter=lambda: self.lang,
                    palette_getter=self._hardware_palette,
                    on_close=self._on_architecture_framework_workspace_closed,
                    language_toggle_callback=self._toggle_lang,
                    theme_toggle_callback=self._toggle_theme,
                )
            )
            generation_state = getattr(self, "_architecture_generation_state", "ready")
            if generation_state == "running":
                self.mimari_cerceve_workspace.on_generation_started()
            elif generation_state == "failed":
                self.mimari_cerceve_workspace.on_generation_failed(
                    getattr(self, "_architecture_generation_detail", "")
                )
        except Exception as error:
            self.mimari_cerceve_workspace = None
            messagebox.showerror(
                self._t("Mimari Çerçeve Hatası", "Architecture Framework Error"),
                self._t(
                    f"Mimari Çerçeve Stüdyosu açılamadı: {error}",
                    f"Architecture Framework Studio could not be opened: {error}",
                ),
            )

    def _on_architecture_framework_workspace_closed(self):
        self.mimari_cerceve_workspace = None

    def _notify_architecture_sources_changed(self, requirement_ids=None):
        """Açık Mimari Stüdyo'da eski snapshot/yayım kapılarını geçersizler."""

        workspace = getattr(self, "mimari_cerceve_workspace", None)
        if not workspace or not workspace.exists:
            return
        stable_ids = (
            None if requirement_ids is None
            else tuple(sorted({str(item).strip().upper() for item in requirement_ids if str(item).strip()}))
        )
        try:
            workspace.on_sources_changed(stable_ids)
        except Exception as error:
            self.update_status_text(
                f"Mimari Çerçeve kaynak değişikliği uygulanamadı: {error}",
                is_error=True,
            )

    def _notify_architecture_source_mutation_started(self):
        """Kaynak değişmeden önce devam eden mimari yayımı senkron iptal eder.

        Bu kanca sohbet ve belge üretim worker'larından da çağrılır. Mimari
        çalışma alanındaki hedef metot Tk nesnesine dokunmaz; ayrıntılı
        stale/yenileme bildirimi mutasyon sonrası ana-thread akışında kalır.
        """

        workspace = getattr(self, "mimari_cerceve_workspace", None)
        hook = getattr(workspace, "on_source_mutation_started", None)
        if callable(hook):
            try:
                hook()
            except Exception as error:
                # Bu kanca worker thread'den gelebilir; hata yolunda dahi Tk
                # ``after``/widget çağrısı yapma. Sonraki ana-thread kaynak
                # bildirimi normal durum/uyarı akışını tamamlar.
                self._architecture_source_mutation_error = str(error)

    def _notify_architecture_generation_started(self):
        """Kısmi belge setinin eski izlenebilirlikle kullanılmasını engeller."""

        self._architecture_generation_state = "running"
        self._architecture_generation_detail = ""
        workspace = getattr(self, "mimari_cerceve_workspace", None)
        if workspace and workspace.exists:
            workspace.on_generation_started()

    def _notify_architecture_traceability_ready(self, requirement_ids=None):
        self._architecture_generation_state = "ready"
        self._architecture_generation_detail = ""
        workspace = getattr(self, "mimari_cerceve_workspace", None)
        if workspace and workspace.exists:
            stable_ids = (
                None if requirement_ids is None
                else tuple(sorted({
                    str(item).strip().upper()
                    for item in requirement_ids if str(item).strip()
                }))
            )
            workspace.on_traceability_ready(stable_ids)

    def _notify_architecture_generation_failed(self, detail=""):
        self._architecture_generation_state = "failed"
        self._architecture_generation_detail = str(detail or "")
        workspace = getattr(self, "mimari_cerceve_workspace", None)
        if workspace and workspace.exists:
            workspace.on_generation_failed(self._architecture_generation_detail)

    def _on_impact_analysis_workspace_closed(self):
        self.impact_analysis_workspace = None

    def _on_hardware_workspace_closed(self):
        self.hardware_workspace = None

    def _on_hardware_cards_workspace_closed(self):
        self.hardware_cards_workspace = None

    def _refresh_hardware_workspace(self):
        workspace = getattr(self, "hardware_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh()

    def _refresh_hardware_cards_workspace(self):
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.refresh()

    def _hardware_palette(self):
        return self._PALETTE["dark" if self.dark else "light"]

    def _get_current_hardware_catalog(self):
        """Açık projenin otomatik kataloğunu bellekten veya atomik JSON'dan yükler."""
        project_name = self.entry_widgets["proje_ismi"].get().strip()
        current = getattr(self, "last_hardware_catalog", None)
        if current and (
            not project_name or current.get("project_name") == project_name
        ):
            return dict(current)
        if not project_name:
            return None
        try:
            loaded = donanim_kartlari_algilama.load_hardware_catalog(project_name)
        except Exception as error:
            self.update_status_text(
                f"Donanım kataloğu yüklenemedi: {error}", is_error=True
            )
            return None
        if loaded:
            self.last_hardware_catalog = loaded.to_dict()
            return dict(self.last_hardware_catalog)
        return None

    def _prepare_hardware_catalog_visuals(self, project_name, catalog):
        """AI/kavramsal görseli kullanıcı onayı olmadan otomatik üretmez."""
        try:
            from config import HARDWARE_AUTO_IMAGE_GENERATION
        except Exception:
            HARDWARE_AUTO_IMAGE_GENERATION = False
        if not HARDWARE_AUTO_IMAGE_GENERATION:
            return 0
        raw_catalog = (
            catalog.to_dict() if hasattr(catalog, "to_dict")
            else dict(catalog or {})
        )
        if not raw_catalog.get("hardware_items"):
            return 0
        overrides = donanim_kartlari_yonetim.load_overrides(
            project_name, raw_catalog
        )
        view, _conflicts = donanim_kartlari_yonetim.apply_overrides(
            raw_catalog, overrides
        )
        pending = [
            item for item in view.get("hardware_items", [])
            if donanim_kartlari_gorsel.illustration_required(item)
        ]
        if not pending:
            return 0
        output_dir = (
            donanim_kartlari_yonetim.overrides_path(
                project_name, raw_catalog
            ).parent / "donanim_gorselleri"
        )
        records = donanim_kartlari_gorsel.generate_catalog_illustrations(
            view, output_dir
        )
        donanim_kartlari_yonetim.set_generated_visuals(overrides, records)
        donanim_kartlari_yonetim.save_overrides(
            project_name, overrides, raw_catalog
        )
        return len(records)

    def _open_hardware_comparison(self, payload):
        """Donanım Kartındaki alternatif ve teknik değerleri manuel analize aktarır."""
        self.open_impact_analysis_workspace()
        workspace = getattr(self, "impact_analysis_workspace", None)
        if not workspace or not workspace.exists:
            return
        try:
            workspace.prefill_hardware_comparison(payload)
        except Exception as error:
            messagebox.showerror(
                "Etki Analizi Aktarım Hatası",
                f"Donanım kartı Etki Analizine aktarılamadı: {error}",
                parent=workspace.window,
            )

    def _open_hardware_requirement(self, requirement_id):
        """Karttan seçilen gereksinimi değişiklik simülasyonu formunda açar."""
        self.open_impact_analysis_workspace()
        workspace = getattr(self, "impact_analysis_workspace", None)
        if not workspace or not workspace.exists:
            return
        try:
            workspace.prefill_requirement_simulation(requirement_id)
        except Exception as error:
            messagebox.showwarning(
                "Gereksinime Git",
                f"Gereksinim simülasyon formunda açılamadı: {error}",
                parent=workspace.window,
            )

    def _open_hardware_detail(self, hardware_reference):
        """Etki Analizi parça bağlantısını aynı Donanım Detay ekranında açar."""
        reference = str(hardware_reference or "").strip()
        report = self._get_current_traceability_report() or {}
        for node in report.get("nodes", []):
            if not isinstance(node, dict) or str(node.get("id", "")).strip() != reference:
                continue
            reference = str(node.get("title") or node.get("description") or reference).strip()
            break
        self.open_hardware_cards_workspace()
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.open_detailed_review(reference)

    def _on_hardware_impact_result(self, result):
        """Gereksinim simülasyonundaki parça etkilerini kart rozetlerine taşır."""
        self.last_hardware_impact_result = (
            result.to_dict() if hasattr(result, "to_dict") else dict(result or {})
        )
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.set_simulation_result(self.last_hardware_impact_result)

    def _start_hardware_datasheet_ingestion(self, datasheet_paths, hardware_id):
        """Datasheet okumasını Tk ana iş parçacığını durdurmadan başlatır."""
        project_name = self.entry_widgets["proje_ismi"].get().strip()
        if not project_name:
            messagebox.showwarning(
                "Datasheet Yükle",
                "Datasheet işlemeden önce proje adını girin.",
            )
            return
        self._hardware_catalog_generation_token += 1
        token = self._hardware_catalog_generation_token
        report = self._get_current_traceability_report() or {}
        flat_snapshot = {
            str(key): dict(value) for key, value in self.flat_data.items()
            if isinstance(value, dict)
        }
        hardware_snapshot = {
            str(key): dict(value) for key, value in self.hardware_data.items()
            if isinstance(value, dict)
        }
        source_paths = list(dict.fromkeys([
            *self.file_paths, *self.generated_document_paths,
        ]))
        threading.Thread(
            target=self._hardware_datasheet_worker,
            args=(
                token, project_name, report, flat_snapshot, hardware_snapshot,
                source_paths, tuple(datasheet_paths), hardware_id,
            ),
            daemon=True,
        ).start()

    def _hardware_datasheet_worker(
        self, token, project_name, report, flat_snapshot, hardware_snapshot,
        source_paths, datasheet_paths, hardware_id,
    ):
        try:
            previous = donanim_kartlari_algilama.load_hardware_catalog(project_name)
            catalog = donanim_kartlari_algilama.build_or_update_hardware_catalog(
                project_name,
                traceability_report=report,
                structured_hardware=hardware_snapshot,
                structured_records=flat_snapshot,
                source_paths=source_paths,
                datasheet_paths=datasheet_paths,
                persist=True,
                status_callback=self.update_status_text,
            )
            try:
                visual_count = self._prepare_hardware_catalog_visuals(
                    project_name, catalog
                )
            except Exception as visual_error:
                visual_count = 0
                self.update_status_text(
                    f"Donanım kartı görselleri hazırlanamadı: {visual_error}",
                    is_error=True,
                )
            change_summary = donanim_kartlari_yonetim.compare_catalogs(
                previous, catalog
            )
            status = {
                "status": "ready", "updated": catalog.updated,
                "hardware_count": len(catalog.hardware_items),
                "instance_count": len(catalog.product_instances),
                "conflict_count": len(catalog.conflicts),
                "visual_count": visual_count,
                "storage_path": catalog.storage_path,
                "target_hardware_id": hardware_id,
                "message": (
                    f"Datasheet işlendi; katalogda {len(catalog.hardware_items)} "
                    "donanım kartı hazır. AI görseli yalnızca kullanıcı onayıyla oluşturulur."
                ),
            }
            self.master.after(
                0, lambda: self._finish_hardware_catalog_refresh(
                    token, catalog, status, change_summary
                )
            )
        except Exception as error:
            self.master.after(
                0, lambda detail=str(error): self._finish_hardware_catalog_failure(
                    token, detail
                )
            )

    def _finish_hardware_catalog_refresh(
        self, token, catalog, status, change_summary
    ):
        if token != self._hardware_catalog_generation_token:
            return
        self.last_hardware_catalog = catalog.to_dict()
        status = dict(status)
        status["change_summary"] = dict(change_summary)
        self.last_hardware_catalog_status = status
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.on_catalog_ready(
                self.last_hardware_catalog, status, change_summary
            )
        self.update_status_text(status["message"], is_complete=True)

    def _finish_hardware_catalog_failure(self, token, detail):
        if token != self._hardware_catalog_generation_token:
            return
        workspace = getattr(self, "hardware_cards_workspace", None)
        if workspace and workspace.exists:
            workspace.set_loading(False, f"Datasheet işlenemedi: {detail}")
        messagebox.showwarning(
            "Datasheet İşleme Uyarısı",
            f"Datasheet katalogla birleştirilemedi: {detail}\nÖzgün dosya değiştirilmedi.",
        )

    def _get_current_traceability_report(self):
        """Açık proje için bellekteki veya kalıcı son izlenebilirlik haritasını döndürür."""
        project_name = self.entry_widgets["proje_ismi"].get().strip()
        if not project_name:
            return None
        current = getattr(self, "last_traceability_report", None)
        if current and current.get("project_name") == project_name:
            return current
        try:
            loaded = etki_analizi_izlenebilirlik.load_project_traceability(project_name)
        except Exception as error:
            self.update_status_text(
                f"İzlenebilirlik haritası yüklenemedi: {error}", is_error=True
            )
            return None
        if loaded:
            if etki_analizi_entegrasyon.overrides_path(loaded).exists():
                try:
                    loaded = etki_analizi_entegrasyon.apply_overrides(loaded)
                except Exception as error:
                    self.update_status_text(
                        f"İzlenebilirlik kullanıcı düzeltmeleri uygulanamadı: {error}",
                        is_error=True,
                    )
            self.last_traceability_report = loaded
        return loaded

    def _set_current_traceability_report(self, report):
        """Etki Analizi ekranındaki kullanıcı düzeltmelerini çalışma kopyasına alır."""
        self._notify_architecture_source_mutation_started()
        self.last_traceability_report = dict(report) if report else None
        if report:
            previous = getattr(self, "last_traceability_health", None) or {}
            self.last_traceability_health = (
                etki_analizi_entegrasyon.build_health_summary(
                    report,
                    {
                        "status": previous.get("rag_status", "not_run"),
                        "message": previous.get("rag_message", ""),
                    },
                )
            )
        self._notify_architecture_traceability_ready()

    def _get_impact_project_info(self):
        """Simülasyon üst durum şeridi için seçili proje/belge setini bildirir."""
        project_name = self.entry_widgets["proje_ismi"].get().strip()
        return {
            "project_name": project_name,
            "source_paths": tuple(self.file_paths),
            "generated_document_paths": tuple(self.generated_document_paths),
            "document_count": len(self.flat_data),
            "health": getattr(self, "last_traceability_health", None),
        }

    def _apply_approved_change_package(self, package, completion_callback, failure_callback):
        """Açık onaylı paketi arka planda yeni ve atomik belge sürümüne dönüştürür."""
        if not isinstance(package, etki_analizi_degisim_paketi.ChangePackage):
            failure_callback("Değişiklik paketi geçerli değil.")
            return
        project_name = self.entry_widgets["proje_ismi"].get().strip() or package.project_name
        flat_snapshot = {
            str(key): dict(value) for key, value in self.flat_data.items()
            if isinstance(value, dict)
        }
        hardware_snapshot = {
            str(key): dict(value) for key, value in self.hardware_data.items()
            if isinstance(value, dict)
        }
        source_paths = list(dict.fromkeys(
            [*self.file_paths, *self.generated_document_paths]
        ))
        sections = tuple(tuple(section) for section in self.VMODEL_SECTIONS)
        self.update_status_text(
            "Onaylanan değişiklikler geçici sürüm alanında hazırlanıyor; özgün belgeler korunuyor..."
        )

        def worker():
            def validator(new_flat, new_hardware, stage):
                post_report = etki_analizi_izlenebilirlik.build_traceability_map(
                    project_name=project_name,
                    flat_data=new_flat,
                    hardware_data=new_hardware,
                    source_paths=(),
                    document_sections=sections,
                    persist=False,
                    check_lm_studio=False,
                )
                request = etki_analizi_simulasyon.ChangeRequest.from_mapping(
                    package.change_request
                )
                target_id = (
                    request.requirement_id
                    or str((package.selected_item or {}).get("id") or "")
                )
                node_ids = {
                    str(node.get("id")) for node in post_report.get("nodes", [])
                    if isinstance(node, dict)
                }
                warnings = []
                if target_id and target_id in node_ids:
                    try:
                        if request.change_type == etki_analizi_simulasyon.CHANGE_REQUIREMENT_ADD:
                            request = etki_analizi_simulasyon.ChangeRequest(
                                requirement_id=target_id,
                                current_value=request.proposed_value,
                                proposed_value=request.proposed_value,
                                reason="Yeni gereksinimin V-Model kapanış kontrolü",
                                requested_by=request.requested_by,
                                change_type=etki_analizi_simulasyon.CHANGE_REQUIREMENT_TEXT,
                                assumptions=request.assumptions,
                                query=request.query,
                            )
                        post_result = etki_analizi_simulasyon.simulate_change(
                            post_report,
                            request,
                            selected_id=target_id,
                            use_existing_rag=False,
                            use_lm_studio=False,
                        ).to_dict()
                    except Exception as error:
                        post_result = {
                            "status": "failed",
                            "message": f"Son etki analizi çalıştırılamadı: {error}",
                            "summary": {"impact_count": 0},
                        }
                        warnings.append(post_result["message"])
                else:
                    post_result = {
                        "status": "completed",
                        "message": (
                            "Değişen gereksinim yeni izlenebilirlikte bulunmuyor; "
                            "kaldırma işlemi doğrulandı."
                        ),
                        "summary": {"impact_count": 0},
                    }
                closure = etki_analizi_degisim_paketi.compare_closure(
                    package, post_report, post_result
                )
                reports_dir = stage / "reports"
                pdf_path = reports_dir / f"{package.change_id}_Etki_Analizi.pdf"
                excel_path = reports_dir / f"{package.change_id}_Etki_Analizi.xlsx"
                etki_analizi_degisim_raporlama.export_change_package_pdf(
                    pdf_path,
                    package,
                    before_traceability=package.baseline_traceability,
                    after_traceability=post_report,
                    closure_summary=closure,
                )
                etki_analizi_degisim_raporlama.export_change_package_excel(
                    excel_path,
                    package,
                    before_traceability=package.baseline_traceability,
                    after_traceability=post_report,
                    closure_summary=closure,
                )
                return {
                    "post_traceability": post_report,
                    "post_simulation": post_result,
                    "closure_summary": closure,
                    "report_paths": {"pdf": pdf_path, "excel": excel_path},
                    "warnings": warnings,
                }

            try:
                result = etki_analizi_degisim_paketi.apply_approved_changes(
                    package,
                    flat_snapshot,
                    hardware_data=hardware_snapshot,
                    source_paths=source_paths,
                    validator=validator,
                )
            except Exception as error:
                self.master.after(
                    0, lambda detail=str(error): failure_callback(detail)
                )
                return
            try:
                result.post_traceability = (
                    etki_analizi_izlenebilirlik.persist_traceability_report(
                        result.post_traceability
                    )
                )
            except Exception as error:
                result.warnings.append(
                    f"Yeni sürüm doğrulandı; güncel izlenebilirlik işaretçisi yazılamadı: {error}"
                )
            try:
                rag_status = etki_analizi_entegrasyon.update_structured_rag_index(
                    result.post_traceability,
                    source_paths=result.created_documents,
                    force=True,
                )
            except Exception as error:
                rag_status = {
                    "status": "failed",
                    "message": f"RAG indeksi güncellenemedi: {error}",
                }
                result.warnings.append(rag_status["message"])
            health = etki_analizi_entegrasyon.build_health_summary(
                result.post_traceability, rag_status
            )
            self.master.after(
                0,
                lambda: self._finish_approved_change_package(
                    result, health, completion_callback
                ),
            )

        threading.Thread(target=worker, daemon=True).start()

    def _finish_approved_change_package(self, result, health, completion_callback):
        """Doğrulanmış sürümü tek noktada aktif eder ve açık ekranları yeniler."""
        self._notify_architecture_source_mutation_started()
        previous_requirement_ids = {
            str(key).strip().upper()
            for key, record in self.flat_data.items()
            if isinstance(record, dict) and record.get("type") in {"TID", "SGD", "STT"}
        }
        self.flat_data.clear()
        self.flat_data.update(result.new_flat_data)
        self.hardware_data.clear()
        self.hardware_data.update(result.new_hardware_data)
        ordered_types = [section[0] for section in self.VMODEL_SECTIONS]
        lines = []
        for document_type in ordered_types:
            for record in self.flat_data.values():
                if record.get("type") == document_type:
                    lines.append(
                        f"{record.get('ID', '')} | {record.get('content', '')}"
                    )
        self.last_generated_output = "\n".join(lines)
        self.raw_output_cache = self.last_generated_output
        self.last_traceability_report = result.post_traceability
        self.last_traceability_health = health
        current_requirement_ids = {
            str(key).strip().upper()
            for key, record in self.flat_data.items()
            if isinstance(record, dict) and record.get("type") in {"TID", "SGD", "STT"}
        }
        changed_requirement_ids = (
            previous_requirement_ids - current_requirement_ids
        ) | {
            str(item).strip().upper()
            for item in (*result.modified_item_ids, *result.added_item_ids)
            if str(item).strip()
        }
        self._notify_architecture_traceability_ready(changed_requirement_ids)
        for path in result.created_documents:
            if path not in self.generated_document_paths:
                self.generated_document_paths.append(path)
        self._refresh_hardware_workspace()
        self._refresh_hardware_cards_workspace()
        workspace = getattr(self, "impact_analysis_workspace", None)
        if workspace and workspace.exists:
            workspace.on_traceability_ready(result.post_traceability, health)
        self.update_status_text(
            f"{result.change_id}: v{result.new_version:04d} oluşturuldu; "
            f"{result.closure_summary.get('resolved_count', 0)} etki çözüldü, "
            f"{result.closure_summary.get('continuing_count', 0)} etki devam ediyor.",
            is_complete=True,
        )
        try:
            completion_callback(result)
        except tk.TclError:
            pass

    def _rescan_traceability_from_workspace(self, force=True):
        project_name = self.entry_widgets["proje_ismi"].get().strip()
        if not project_name:
            messagebox.showwarning(
                "İzlenebilirliği Yeniden Tara",
                "Önce proje adını girin ve belgeleri üretin.",
            )
            workspace = getattr(self, "impact_analysis_workspace", None)
            if workspace and workspace.exists:
                workspace.on_traceability_failed("Proje adı bulunamadı.")
            return
        if not self.flat_data and not self.file_paths:
            messagebox.showwarning(
                "İzlenebilirliği Yeniden Tara",
                "Taranacak üretilmiş belge verisi bulunamadı. Önce 'Dokümanları Üret' işlemini tamamlayın.",
            )
            workspace = getattr(self, "impact_analysis_workspace", None)
            if workspace and workspace.exists:
                workspace.on_traceability_failed("Taranacak belge seti bulunamadı.")
            return
        self._traceability_generation_token += 1
        self._start_traceability_build(project_name, force_rescan=bool(force))

    def _cancel_traceability_from_workspace(self):
        self._traceability_cancel_event.set()
        self._traceability_generation_token += 1
        self._notify_architecture_generation_failed("Traceability scan was cancelled.")
        self.update_status_text("Etki Analizi arka plan işlemi iptal edildi.", is_error=True)

    def _start_traceability_build(self, project_name, force_rescan=False):
        """Başarılı belge üretiminin yapılandırılmış verisini arka planda tarar."""
        token = self._traceability_generation_token
        self._traceability_cancel_event.set()
        cancel_event = threading.Event()
        self._traceability_cancel_event = cancel_event
        flat_snapshot = {
            str(key): dict(value)
            for key, value in self.flat_data.items()
            if isinstance(value, dict)
        }
        hardware_snapshot = {
            str(key): dict(value)
            for key, value in self.hardware_data.items()
            if isinstance(value, dict)
        }
        source_paths = list(dict.fromkeys([
            *self.file_paths,
            *self.generated_document_paths,
        ]))
        sections = tuple(tuple(section) for section in self.VMODEL_SECTIONS)
        self.update_status_text(
            "Etki analizi izlenebilirlik altyapısı arka planda hazırlanıyor..."
        )
        self.master.after(0, self._notify_traceability_started)
        threading.Thread(
            target=self._traceability_worker,
            args=(
                token,
                project_name,
                flat_snapshot,
                hardware_snapshot,
                source_paths,
                sections,
                force_rescan,
                cancel_event,
            ),
            daemon=True,
        ).start()

    def _traceability_worker(
        self,
        token,
        project_name,
        flat_snapshot,
        hardware_snapshot,
        source_paths,
        sections,
        force_rescan=False,
        cancel_event=None,
    ):
        """İzlenebilirliği üretir; hata belge üretiminin sonucunu değiştirmez."""
        cancel_event = cancel_event or threading.Event()
        try:
            if cancel_event.is_set():
                return
            report = etki_analizi_izlenebilirlik.build_traceability_map(
                project_name=project_name,
                flat_data=flat_snapshot,
                hardware_data=hardware_snapshot,
                source_paths=source_paths,
                document_sections=sections,
                status_callback=self.update_status_text,
            )
        except Exception as error:
            if token != self._traceability_generation_token or cancel_event.is_set():
                return
            message = (
                "Belgeler üretildi ancak Etki Analizi izlenebilirlik altyapısı "
                f"hazırlanamadı: {error}"
            )
            self.update_status_text(message, is_error=True)
            self.master.after(0, lambda detail=str(error): self._finish_traceability_failure(detail))
            return

        if token != self._traceability_generation_token or cancel_event.is_set():
            return
        try:
            report = etki_analizi_entegrasyon.apply_overrides(report)
        except Exception as error:
            self.update_status_text(
                f"Kullanıcı izlenebilirlik düzeltmeleri uygulanamadı: {error}",
                is_error=True,
            )
        try:
            rag_status = etki_analizi_entegrasyon.update_structured_rag_index(
                report,
                source_paths=source_paths,
                force=force_rescan,
                cancel_event=cancel_event,
                status_callback=self.update_status_text,
            )
        except Exception as error:
            rag_status = {
                "status": "failed",
                "updated": False,
                "message": f"RAG güncelleme uyarısı: {error}",
            }
        if token != self._traceability_generation_token or cancel_event.is_set():
            return
        health = etki_analizi_entegrasyon.build_health_summary(report, rag_status)
        catalog = None
        catalog_status = {"status": "unavailable", "message": "Donanım kataloğu oluşturulmadı."}
        try:
            previous_catalog = (
                donanim_kartlari_algilama.load_hardware_catalog(project_name)
                if report.get("storage_path") else None
            )
            catalog = donanim_kartlari_algilama.build_or_update_hardware_catalog(
                project_name,
                traceability_report=report,
                structured_hardware=hardware_snapshot,
                structured_records=flat_snapshot,
                source_paths=source_paths,
                persist=bool(report.get("storage_path")),
                status_callback=self.update_status_text,
            )
            try:
                visual_count = self._prepare_hardware_catalog_visuals(
                    project_name, catalog
                )
            except Exception as visual_error:
                visual_count = 0
                self.update_status_text(
                    f"Donanım kartı görselleri hazırlanamadı: {visual_error}",
                    is_error=True,
                )
            relation_count = sum(
                1 for item in catalog.product_tree
                if item.get("parent_instance_id") not in {None, "", "Veri bulunamadı"}
            )
            catalog_change_summary = donanim_kartlari_yonetim.compare_catalogs(
                previous_catalog, catalog
            )
            catalog_status = {
                "status": "ready",
                "updated": catalog.updated,
                "hardware_count": len(catalog.hardware_items),
                "instance_count": len(catalog.product_instances),
                "relation_count": relation_count,
                "conflict_count": len(catalog.conflicts),
                "visual_count": visual_count,
                "storage_path": catalog.storage_path,
                "change_summary": catalog_change_summary,
                "message": (
                    f"Donanım kataloğu hazır: {len(catalog.hardware_items)} kart, "
                    f"{len(catalog.product_instances)} kullanım yeri, "
                    "AI görsel üretimi kullanıcı onayı bekliyor."
                ),
            }
        except Exception as error:
            catalog_status = {
                "status": "failed",
                "message": (
                    "Etki analizi hazır; donanım kataloğu güncellenemedi: "
                    f"{error}"
                ),
            }
            self.update_status_text(catalog_status["message"], is_error=True)
        summary = report.get("summary", {})
        lm_status = report.get("capabilities", {}).get("lm_studio", {})
        ready_message = (
            "Etki analizi altyapısı hazır. "
            f"{summary.get('node_count', 0)} düğüm ve "
            f"{summary.get('edge_count', 0)} ilişki oluşturuldu. "
            f"{health.get('unlinked_count', 0)} bağlantısız ve "
            f"{health.get('unverified_count', 0)} doğrulama testi olmayan gereksinim bulundu. "
            f"{catalog_status.get('hardware_count', 0)} donanım kartı hazırlandı."
        )
        self.master.after(
            0,
            lambda: self._finish_traceability_success(
                token, report, health, ready_message, lm_status, rag_status,
                catalog, catalog_status,
            ),
        )

    def _notify_traceability_started(self):
        self._notify_architecture_generation_started()
        workspace = getattr(self, "impact_analysis_workspace", None)
        if workspace and workspace.exists:
            workspace.on_traceability_started()
        cards_workspace = getattr(self, "hardware_cards_workspace", None)
        if cards_workspace and cards_workspace.exists:
            cards_workspace.set_loading(
                True, "Belge seti ve donanım kataloğu arka planda taranıyor…"
            )

    def _finish_traceability_failure(self, detail):
        self._notify_architecture_generation_failed(detail)
        workspace = getattr(self, "impact_analysis_workspace", None)
        if workspace and workspace.exists:
            workspace.on_traceability_failed(
                f"İzlenebilirlik haritası oluşturulamadı: {detail}"
            )
        cards_workspace = getattr(self, "hardware_cards_workspace", None)
        if cards_workspace and cards_workspace.exists:
            cards_workspace.set_loading(
                False, f"İzlenebilirlik taraması tamamlanamadı: {detail}"
            )
        messagebox.showwarning(
            "Etki Analizi Altyapısı",
            "Belge üretimi başarıyla tamamlandı. İzlenebilirlik haritası "
            f"oluşturulamadı:\n{detail}",
        )

    def _finish_traceability_success(
        self, token, report, health, ready_message, lm_status, rag_status,
        hardware_catalog=None, hardware_catalog_status=None,
    ):
        if token != self._traceability_generation_token:
            return
        self._notify_architecture_source_mutation_started()
        self.last_traceability_report = report
        self.last_traceability_health = health
        self._notify_architecture_traceability_ready()
        self.last_hardware_catalog = (
            hardware_catalog.to_dict() if hardware_catalog is not None else None
        )
        self.last_hardware_catalog_status = hardware_catalog_status or {
            "status": "unavailable"
        }
        self.update_status_text(ready_message, is_complete=True)
        if lm_status.get("available") is False:
            self.update_status_text(lm_status.get("message", ""), is_error=True)
        if rag_status.get("status") in {"failed", "unavailable"}:
            self.update_status_text(rag_status.get("message", ""), is_error=True)
        workspace = getattr(self, "impact_analysis_workspace", None)
        if workspace and workspace.exists:
            workspace.on_traceability_ready(report, health)
        cards_workspace = getattr(self, "hardware_cards_workspace", None)
        if cards_workspace and cards_workspace.exists:
            cards_workspace.on_catalog_ready(
                self.last_hardware_catalog,
                self.last_hardware_catalog_status,
                self.last_hardware_catalog_status.get("change_summary", {}),
            )
        notice = ready_message
        if rag_status.get("status") in {"failed", "unavailable"}:
            notice += (
                "\n\nRAG indeksi güncellenemedi; grafik tabanlı analiz kullanılmaya devam edecek."
            )
        if self.last_hardware_catalog_status.get("status") == "failed":
            notice += "\n\n" + self.last_hardware_catalog_status.get("message", "")
        messagebox.showinfo("Etki Analizi Altyapısı", notice)


    def _invalidate_hardware_generation(self, message=""):
        """Eski arka plan sonucunun yeni belge verisini ezmesini önler."""
        self._hardware_generation_token += 1
        workspace = getattr(self, "hardware_workspace", None)
        if workspace and workspace.exists:
            workspace.set_generation_state(False, message)

    def start_hardware_generation(self):
        """SGD/STT kayıtlarından arka planda yapılandırılmış donanım önerileri üretir."""
        records = hardware_list_logic.eligible_requirement_records(self.flat_data)
        workspace = getattr(self, "hardware_workspace", None)
        if not records:
            message = self._t(
                "Donanım analizi için içerikli SGD veya STT kaydı bulunamadı. "
                "Önce bu dokümanlardan en az birini üretin.",
                "No populated SGD or STT record was found for hardware analysis. "
                "Generate at least one of these documents first.",
            )
            if workspace and workspace.exists:
                workspace.set_generation_state(False, message)
            messagebox.showwarning(
                self._t("Donanım Önerisi", "Hardware Suggestions"),
                message,
            )
            return

        self._hardware_generation_token += 1
        token = self._hardware_generation_token
        project_name = self.entry_widgets["proje_ismi"].get().strip() or "Proje"
        flat_snapshot = {
            str(key): dict(value)
            for key, value in self.flat_data.items()
            if isinstance(value, dict)
        }
        hardware_snapshot = {
            str(key): dict(value)
            for key, value in self.hardware_data.items()
            if isinstance(value, dict)
        }
        running_message = self._t(
            f"{len(records)} SGD/STT maddesi analiz ediliyor…",
            f"Analyzing {len(records)} SGD/STT records…",
        )
        if workspace and workspace.exists:
            workspace.set_generation_state(True, running_message)
        self.update_status_text(running_message)

        thread = threading.Thread(
            target=self._run_hardware_generation,
            args=(token, project_name, flat_snapshot, hardware_snapshot),
            daemon=True,
        )
        thread.start()

    def _run_hardware_generation(
        self,
        token,
        project_name,
        flat_snapshot,
        hardware_snapshot,
    ):
        try:
            result = hardware_generator_logic.run_generation_from_requirements(
                flat_data=flat_snapshot,
                project_name=project_name,
                existing_hardware=hardware_snapshot,
                status_callback=self.update_status_text,
            )
        except Exception as error:
            result = {
                "result": False,
                "message": f"Donanım önerisi üretim hatası: {error}",
            }
        try:
            self.master.after(0, lambda: self._finish_hardware_generation(token, result))
        except tk.TclError:
            pass

    def _finish_hardware_generation(self, token, result):
        if token != self._hardware_generation_token:
            return

        workspace = getattr(self, "hardware_workspace", None)
        if result.get("result"):
            self.hardware_data.clear()
            self.hardware_data.update(result.get("hardware_data", {}))
            suggestion_count = int(result.get("suggestion_count", 0))
            requirement_count = int(result.get("requirement_count", 0))
            message = self._t(
                f"{suggestion_count} donanım önerisi oluşturuldu; {requirement_count} gereksinim analiz edildi.",
                f"Created {suggestion_count} hardware suggestions; analyzed {requirement_count} requirements.",
            )
            if workspace and workspace.exists:
                workspace.refresh()
                workspace.set_generation_state(False, message)
            self.update_status_text(message, is_complete=True)
            return

        message = str(result.get("message") or self._t(
            "Donanım önerisi üretilemedi.",
            "Hardware suggestions could not be generated.",
        ))
        if workspace and workspace.exists:
            workspace.set_generation_state(False, message)
        self.update_status_text(message, is_error=True)
        messagebox.showerror(
            self._t("Donanım Önerisi Hatası", "Hardware Suggestion Error"),
            message,
        )

    def run_kalite_denetimi(self):
        """Üretilen maddeleri kalite kurallarına göre denetler ve ayrı bir rapor penceresi açar.
        Salt-okunur: hiçbir veriyi değiştirmez. (Ayrı özellik — kalite_denetci.py'ye bağlı.)"""
        if not self.flat_data:
            messagebox.showwarning("Uyarı", "Denetlenecek veri yok. Önce doküman üretin.")
            return
        try:
            import kalite_denetci
        except Exception as e:
            messagebox.showerror("Hata", f"Kalite denetçisi yüklenemedi: {e}")
            return
        win = tk.Toplevel(self.master)
        win.title("Gereksinim Kalite Raporu")
        win.geometry("670x620")

        baslik = ttk.Label(win, font=("Segoe UI", 12, "bold"), foreground=self.dark_blue)
        baslik.pack(anchor="w", padx=12, pady=(10, 4))

        frame = ttk.Frame(win)
        frame.pack(fill="both", expand=True, padx=10, pady=(0, 6))
        sb = ttk.Scrollbar(frame, orient="vertical")
        sb.pack(side="right", fill="y")
        txt = tk.Text(frame, wrap="word", font=("Consolas", 10),
                      yscrollcommand=sb.set, relief="solid", borderwidth=1)
        txt.pack(side="left", fill="both", expand=True)
        sb.config(command=txt.yview)

        def _yenile():
            rapor = kalite_denetci.denetle(self.flat_data)
            s = rapor["summary"]
            baslik.config(text=f"Kalite Puanı: %{s['kalite']}   "
                               f"(⚠️ {s['problemli']} sorunlu / {s['total']} madde, 📌 {s['dsb']} DSB)")
            txt.config(state=tk.NORMAL)
            txt.delete("1.0", tk.END)
            txt.insert("1.0", kalite_denetci.rapor_metni(rapor))
            txt.config(state=tk.DISABLED)

        ttk.Button(win, text="🔄 Yenile (Copilot düzeltmelerinden sonra bas)",
                   command=_yenile, style="primary.TButton").pack(pady=(0, 10))
        _yenile()

    def reset_app(self):
        cevap = messagebox.askyesno("Sıfırla", "Tüm seçimler, yüklenen dosyalar ve üretilen veriler silinecek.\nEmin misiniz?")
        if not cevap:
            return

        self._notify_architecture_source_mutation_started()
        architecture_workspace = getattr(self, "mimari_cerceve_workspace", None)
        if architecture_workspace and architecture_workspace.exists:
            architecture_workspace.close()

        self.file_paths = []
        self.template_file_path = None
        self.entry_widgets["proje_bilesenleri"].config(
            text=self.FILE_HINT if self._dnd_enabled else "Seçilmedi")
        self.entry_widgets["template_file"].config(text="Seçilmedi")
        self.entry_widgets["proje_ismi"].delete(0, tk.END)
        
        sayac_alanlari = ["teknik_ister", "sistem_gereksinimi", "sistem_tanimlama_testi"]
        for key in sayac_alanlari:
            self.entry_widgets[key].delete(0, tk.END)
            self.entry_widgets[key].insert(0, "0")

        for var in self.checkbox_vars.values():
            var.set(True)

        self.last_generated_output = ""
        self.raw_output_cache = ""
        self.tree_data.clear()
        self._traceability_generation_token += 1
        self._traceability_cancel_event.set()
        self._architecture_generation_state = "ready"
        self._architecture_generation_detail = ""
        self.flat_data.clear()
        self.hardware_data.clear()
        self.generated_document_paths.clear()
        self.last_hardware_catalog = None
        self.last_hardware_catalog_status = None
        self.last_hardware_impact_result = None
        self._invalidate_hardware_generation()
        self._refresh_hardware_workspace()
        self._refresh_hardware_cards_workspace()
        self.last_tid_list = []      
        self.last_sgd_list = []
        self.last_stt_list = []
        self.last_dgoygo_list = []
        self.last_sitet_list = []
        self.last_alt_sistem_test_list = []
        self.last_dtet_ytet_list = []

        self.output_text_area.config(state=tk.NORMAL)
        self.output_text_area.delete("1.0", tk.END)
        self.output_text_area.config(state=tk.DISABLED)

        self.create_docs_button.config(state=tk.NORMAL, text=self._t("Dokümanları Üret", "Generate Documents"), style="primary.TButton")
        self.download_docs_button.config(state=tk.NORMAL)

        messagebox.showinfo("Bilgi", "Uygulama başarıyla sıfırlandı.")

    def select_files(self):
        paths = filedialog.askopenfilenames(
            title="Girdi Dosyalarını Seç",
            filetypes=[("PDF ve TXT dosyaları", "*.pdf *.txt")]
        )
        if paths:
            self.file_paths = list(paths)
            names = [os.path.basename(p) for p in self.file_paths]
            display = ", ".join(names) if len(names) <= 3 else f"{len(names)} dosya seçildi"
            self.entry_widgets["proje_bilesenleri"].config(text=display)
            self.update_status_text(f"Seçilen dosya(lar): {display}", clear=True)

    def _register_drop_target(self, widget):
        """Bir widget'ı PDF/TXT sürükle-bırak hedefi yapar (tkinterdnd2 varsa)."""
        if not getattr(self, "_dnd_enabled", False):
            return
        try:
            from tkinterdnd2 import DND_FILES
            widget.drop_target_register(DND_FILES)
            widget.dnd_bind("<<Drop>>", self._on_files_dropped)
        except Exception as e:
            print(f"Sürükle-bırak kaydı başarısız: {e}")

    def _on_files_dropped(self, event):
        """OS'ten sürüklenip bırakılan dosyaları işler."""
        try:
            dropped = self.master.tk.splitlist(event.data)
        except Exception:
            dropped = [event.data]
        valid = [p for p in dropped if p.lower().endswith((".pdf", ".txt"))]
        if not valid:
            self.update_status_text("⚠️ Sadece PDF/TXT dosyaları sürükle-bırak yapılabilir.", is_error=True)
            return
        self.file_paths = list(valid)
        names = [os.path.basename(p) for p in self.file_paths]
        display = ", ".join(names) if len(names) <= 3 else f"{len(names)} dosya seçildi"
        self.entry_widgets["proje_bilesenleri"].config(text=display)
        self.update_status_text(f"Sürükle-bırak ile eklendi: {display}", clear=True)
        return event.action

    def select_template_file(self):
        path = filedialog.askopenfilename(
            title="Şablon Dosyasını Seç (docx)",
            filetypes=[("DOCX dosyaları", "*.docx"), ("Tüm Dosyalar", "*.*")]
        )
        if path:
            self.template_file_path = path
            self.entry_widgets["template_file"].config(text=os.path.basename(path))
            self.update_status_text(f"Seçilen şablon: {os.path.basename(path)}")
        else:
            self.template_file_path = None
            self.entry_widgets["template_file"].config(text="Seçilmedi")

    def update_status_text(self, message, is_error=False, is_complete=False, clear=False):
        def _inner():
            self.output_text_area.config(state=tk.NORMAL)
            if clear:
                self.output_text_area.delete("1.0", tk.END)

            tag = "msg"
            if is_error:
                tag = "error"
                self.output_text_area.tag_config(tag, foreground=self.style.colors.danger,
                                                 font=("Segoe UI", 10, "bold"))
                self.output_text_area.insert(tk.END, f"\n[HATA] {message}", tag)
            elif is_complete:
                tag = "complete"
                self.output_text_area.tag_config(tag, foreground=self.style.colors.success,
                                                 font=("Segoe UI", 10, "bold"))
                self.output_text_area.insert(tk.END, f"\n{message}", tag)
            else:
                self.output_text_area.insert(tk.END, f"\n{message}", tag)

            self.output_text_area.see(tk.END)
            self.output_text_area.config(state=tk.DISABLED)
        self.master.after(0, _inner)

    # V-Modelindeki doküman türleri: (flat_data type, başlık, bacak)
    # 'req' = sol bacak (gereksinim), 'test' = sağ bacak (doğrulama)
    VMODEL_SECTIONS = [
        ("TID",       "Kullanıcı Gereksinimi (User Requirement)",                "req"),
        ("KMTD",      "Kabul Testi (Acceptance Test)",               "test"),
        ("SGD",       "Sistem Gereksinimi (System Requirements)",            "req"),
        ("SITET",     "Sistem Testi (System Test)",                          "test"),
        ("STT",       "Alt Sistem Gereksinimleri (Subsystem Requirements)",  "req"),
        ("AST",       "Alt Sistem Testi (Subsystem Test)",                   "test"),
    ]

    def _register_pdf_font(self):
        """Türkçe karakterleri destekleyen bir font ailesi kaydeder; (normal, kalın) döner."""
        candidates = [
            ("VModelFont", "VModelFont-Bold",
             r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        ]
        for reg_name, bold_name, reg_path, bold_path in candidates:
            try:
                if os.path.exists(reg_path):
                    pdfmetrics.registerFont(TTFont(reg_name, reg_path))
                    if os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont(bold_name, bold_path))
                    else:
                        bold_name = reg_name
                    return reg_name, bold_name
            except Exception:
                continue
        # Yedek: reportlab ile gelen Bitstream Vera (Türkçe destekli)
        try:
            import reportlab
            fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
            pdfmetrics.registerFont(TTFont("VModelFont", os.path.join(fonts_dir, "Vera.ttf")))
            pdfmetrics.registerFont(TTFont("VModelFont-Bold", os.path.join(fonts_dir, "VeraBd.ttf")))
            return "VModelFont", "VModelFont-Bold"
        except Exception:
            return "Helvetica", "Helvetica-Bold"

    def _generate_vmodel_pdf(self, path, proje_ismi):
        """flat_data'yı kullanarak V-Model yapısında biçimli bir PDF üretir."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        reg, bold = self._register_pdf_font()

        title_style = ParagraphStyle("vm_title", fontName=bold, fontSize=20,
                                     textColor=colors.HexColor("#1F3864"), leading=24, spaceAfter=4)
        sub_style = ParagraphStyle("vm_sub", fontName=reg, fontSize=11,
                                   textColor=colors.HexColor("#555555"), leading=15)
        legend_style = ParagraphStyle("vm_legend", fontName=reg, fontSize=9,
                                      textColor=colors.HexColor("#555555"), leading=13)
        sec_style = ParagraphStyle("vm_sec", fontName=bold, fontSize=13,
                                   textColor=colors.white, leading=16)
        cell_style = ParagraphStyle("vm_cell", fontName=reg, fontSize=9, leading=12)
        cellid_style = ParagraphStyle("vm_cellid", fontName=bold, fontSize=9, leading=12,
                                      textColor=colors.HexColor("#1F3864"))
        head_style = ParagraphStyle("vm_head", fontName=bold, fontSize=9, leading=12,
                                    textColor=colors.white)

        # Renkler: sol bacak (gereksinim) mavi, sağ bacak (test) yeşil tonları
        REQ_COLOR = colors.HexColor("#1F3864")
        REQ_LIGHT = colors.HexColor("#D6E0F0")
        TEST_COLOR = colors.HexColor("#375623")
        TEST_LIGHT = colors.HexColor("#E2EFDA")

        story = []
        story.append(Paragraph("V-MODEL TEKNİK DOKÜMAN PAKETİ", title_style))
        story.append(Paragraph(f"Proje: <b>{proje_ismi}</b>", sub_style))
        story.append(Paragraph(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}", sub_style))
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1F3864")))
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            "<font color='#1F3864'>■</font> Sol Bacak: Gereksinim Dokümanları &nbsp;&nbsp;&nbsp; "
            "<font color='#375623'>■</font> Sağ Bacak: Test / Doğrulama Dokümanları", legend_style))
        story.append(Spacer(1, 12))

        # Sayfa genişliği ~ 170mm kullanılabilir
        col_widths = [26 * mm, 118 * mm, 26 * mm]
        any_section = False

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_section = True

            is_req = (leg == "req")
            head_color = REQ_COLOR if is_req else TEST_COLOR
            light_color = REQ_LIGHT if is_req else TEST_LIGHT
            bound_header = "Kaynak / Bağlı" if is_req else "Doğruladığı Madde"

            # Bölüm başlığı bandı
            sec_band = Table([[Paragraph(f"{title}  ({len(items)} madde)", sec_style)]],
                             colWidths=[sum(col_widths)])
            sec_band.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), head_color),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(sec_band)

            # Tablo verisi
            data = [[Paragraph("ID", head_style),
                     Paragraph("Açıklama", head_style),
                     Paragraph(bound_header, head_style)]]
            for d in items:
                data.append([
                    Paragraph(str(d.get("ID", "")), cellid_style),
                    Paragraph(str(d.get("content", "")), cell_style),
                    Paragraph(str(d.get("bound_to", "-")), cell_style),
                ])

            tbl = Table(data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), head_color),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_color]),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 16))

        if not any_section:
            # flat_data boşsa ham metni sığdır
            story.append(Paragraph("Üretilen doküman verisi bulunamadı. Ham çıktı:", sub_style))
            story.append(Spacer(1, 6))
            for line in (self.last_generated_output or "").split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), cell_style))

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=18 * mm, bottomMargin=18 * mm,
            title=f"{proje_ismi} - V-Model Doküman Paketi"
        )
        doc.build(story)

    # ================================================================== #
    #  DOKÜMAN COPILOT — sağ panel, tek madde revizyonu (grounding'li)    #
    # ================================================================== #

    # --- Kıdemli Sistem Mühendisi system prompt (paydaş kuralları 1-4) ---
    COPILOT_SYSTEM_PROMPT = (
        "ROL: Sen kıdemli bir Sistem Mühendisisin (INCOSE / sistem mühendisliği "
        "standartlarına hâkim). Görevin, sana verilen TEK bir gereksinim veya test "
        "maddesini kullanıcının isteğine göre yeniden yazmaktır. Aşağıdaki kurallara "
        "İSTİSNASIZ uyarsın:\n"
        "1) SADE ve SAF: Yanıtın YALNIZCA revize edilmiş madde metnidir. 'Tabii ki', "
        "'İşte güncellenmiş hâli' gibi laf kalabalığı, giriş/kapanış cümlesi, açıklama, "
        "yorum, başlık, ID, numara, madde işareti, yıldız veya tırnak KESİNLİKLE YOK. "
        "Tek paragraf, en fazla 2-3 cümle.\n"
        "2) BİRİMLER: Tüm sayısal değerleri metrik/teknik birimlerle ver "
        "(ms, s, Hz, kHz, MHz, dB, dBm, m, km, %, °C). Birimsiz çıplak sayı bırakma.\n"
        "3) TEKNİK DİL: Günlük dil kullanma; sistem mühendisliği terminolojisi kullan. "
        "Cümleler net, ölçülebilir ve doğrulanabilir olmalı ('sistem ... -malıdır/-melidir', "
        "'doğrulanmalıdır', 'karşılanmalıdır').\n"
        "4) DSB KURALI (SAYI UYDURMAK KESİNLİKLE YASAK): Bir sayısal değeri YALNIZCA şu iki "
        "durumda yazabilirsin: (a) kullanıcı o sayıyı isteğinde AÇIKÇA vermişse, ya da "
        "(b) sayı zaten MEVCUT METİN'de varsa. Kullanıcı 'bir değer/süre/eşik/mesafe ekle' "
        "diyor ama sayının kendisini SÖYLEMİYORSA, o sayıyı TAHMİN ETME — makul/tipik bir "
        "değer bile olsa UYDURMA. Bunun yerine değerin geçtiği yere harfi harfine 'DSB' yaz "
        "(DSB = Daha Sonra Belirlenecek), birimi koru. "
        "Örnek: kullanıcı 'tespit süresi ekle' dedi ama süreyi vermediyse → "
        "'Sistem, hedefleri DSB ms içinde tespit etmelidir.'\n"
        "ÖRNEKLER (uy):\n"
        "- MEVCUT: 'Sistem, hedefleri tespit etmelidir.' | İSTEK: 'tespit süresi ekle' "
        "(kullanıcı sayı VERMEDİ) → DOĞRU: 'Sistem, hedefleri DSB ms içinde tespit etmelidir.' | "
        "YANLIŞ: 'Sistem, hedefleri 100 ms içinde tespit etmelidir.' (100 UYDURMA, yasak!).\n"
        "- MEVCUT: 'Sistem çalışmalıdır.' | İSTEK: 'çalışma sıcaklığı -20 ile +55 °C olsun' "
        "(değer VERİLDİ) → DOĞRU: 'Sistem, -20 ile +55 °C sıcaklık aralığında çalışmalıdır.'\n"
        "ÇIKTI: Yalnızca revize edilmiş madde metni. Başka HİÇBİR ŞEY yazma."
    )

    # DSB ripple'ının yayılacağı test tipi maddeler
    TEST_TYPES = ("KMTD", "SITET", "AST")
    # Revize edilince bağlı testleri yeniden üretilecek gereksinim tipleri
    REQ_TYPES = ("TID", "SGD", "STT")

    def _create_chat_panel(self, master, target_bg):
        """Sağ tarafta üretilen maddeleri revize etmek için sohbet paneli."""
        panel = ttk.Frame(master, style="light", width=370)
        panel.pack(side="right", fill="y")
        panel.pack_propagate(False)

        head = ttk.Frame(panel, style="light")
        head.pack(fill="x", padx=10, pady=(12, 4))
        self._L(head, "🤖 Doküman Copilot", "🤖 Document Copilot", font=("Segoe UI", 13, "bold"),
                foreground=self.dark_blue, style="primary.TLabel").pack(side="left")

        self._L(panel,
                "Üretilen bir maddeyi revize ettir. Örn:\n«UR-004'ü daha teknik yaz, güvenlik standardı ekle»",
                "Revise a generated item. E.g.:\n«Make UR-004 more technical, add a security standard»",
                font=("Segoe UI", 8), foreground="#666", justify="left",
                wraplength=340, style="secondary.TLabel").pack(fill="x", padx=10, pady=(0, 6))

        hist_frame = ttk.Frame(panel, style="light")
        hist_frame.pack(fill="both", expand=True, padx=10)
        sb = ttk.Scrollbar(hist_frame, orient="vertical")
        sb.pack(side="right", fill="y")
        self.chat_history = tk.Text(hist_frame, wrap="word", relief="solid", borderwidth=1,
                                    state=tk.DISABLED, font=("Segoe UI", 9), yscrollcommand=sb.set)
        self.chat_history.pack(side="left", fill="both", expand=True)
        self._theme_texts.append(self.chat_history)
        sb.config(command=self.chat_history.yview)
        self.chat_history.tag_config("user", foreground="#0052cc", font=("Segoe UI", 9, "bold"))
        self.chat_history.tag_config("bot", foreground="#1a7f37")
        self.chat_history.tag_config("err", foreground="#c1121f")
        self.chat_history.tag_config("info", foreground="#666", font=("Segoe UI", 8, "italic"))

        entry_frame = ttk.Frame(panel, style="light")
        entry_frame.pack(fill="x", padx=10, pady=10)
        self.chat_entry = ttk.Entry(entry_frame, font=("Segoe UI", 10))
        self.chat_entry.pack(side="left", fill="x", expand=True)
        self.chat_entry.bind("<Return>", lambda e: self._chat_send())
        self.chat_send_btn = ttk.Button(entry_frame, style="primary.TButton",
                                        command=self._chat_send, width=8)
        self._reg_btn(self.chat_send_btn, "Gönder", "Send")
        self.chat_send_btn.config(text=self._t("Gönder", "Send"))
        self.chat_send_btn.pack(side="left", padx=(6, 0))

        # Karşılama metni (dil değişince yeniden yazılabilmesi için saklanır)
        self._chat_greeting = (
            "Merhaba! Önce dokümanları üret, sonra bir maddeyi bana revize ettir. Örn: «SR-002'yi daha ölçülebilir yap».",
            "Hello! First generate documents, then ask me to revise an item. E.g.: «Make SR-002 more measurable».")
        self._chat_has_convo = False   # gerçek bir sohbet başlayınca True
        self._chat_append(self._t(*self._chat_greeting), "info")

    def _chat_append(self, text, tag="bot"):
        def _inner():
            self.chat_history.config(state=tk.NORMAL)
            prefix = {"user": "\n👤 Sen:\n", "bot": "\n🤖 Copilot:\n", "err": "\n⚠️ "}.get(tag, "\n")
            self.chat_history.insert(tk.END, prefix + text + "\n", tag)
            self.chat_history.see(tk.END)
            self.chat_history.config(state=tk.DISABLED)
        self.master.after(0, _inner)

    def _find_target_id(self, msg):
        """Mesajdan flat_data'daki bir madde ID'sini yakalar (Türkçe/nokta duyarsız)."""
        def norm(s):
            s = s.upper()
            for a, b in (("İ", "I"), ("Ö", "O"), ("Ğ", "G"), ("Ü", "U"),
                         ("Ş", "S"), ("Ç", "C"), ("I", "I")):
                s = s.replace(a, b)
            return s
        nmsg = norm(msg)
        # Uzun ID'leri (DTET-YTET-001 gibi) önce dene ki kısa parçalar yanlış eşleşmesin
        for key in sorted(self.flat_data.keys(), key=len, reverse=True):
            if norm(key) in nmsg:
                return key
        return None

    def _chat_send(self):
        msg = self.chat_entry.get().strip()
        if not msg:
            return
        if not self.flat_data:
            self._chat_append("Henüz üretilmiş doküman yok. Önce 'Dokümanları Üret' ile üretim yap.", "err")
            return
        self.chat_entry.delete(0, tk.END)
        self._chat_has_convo = True   # artık dil değişiminde karşılamayı üzerine yazma
        self._chat_append(msg, "user")
        self.chat_send_btn.config(state=tk.DISABLED)
        threading.Thread(target=self._chat_worker, args=(msg,), daemon=True).start()

    def _chat_worker(self, msg):
        try:
            from llm_handler import call_gemma3_api
            target = self._find_target_id(msg)
            if not target:
                self._chat_append("Hangi maddeyi revize edeceğimi bulamadım. Lütfen ID yaz "
                                  "(ör: UR-004, SR-002).", "err")
                return
            old = self.flat_data[target].get("content", "")
            # "BÖL" komutu: atomik olmayan bir maddeyi ayrı gereksinimlere böl
            if any(w in msg.lower() for w in ("böl", "parçala", "ayrı madde")):
                self._split_requirement(target)
                return
            self._chat_append(f"{target} revize ediliyor...", "info")

            # DSB kararını PYTHON'da deterministik ver (4B model 'değer verildi mi' ayrımını
            # güvenilir yapamıyor). 3 durum: (1) değer isteniyor+sayı var → kullan,
            # (2) değer isteniyor+sayı yok → DSB, (3) sadece ton/ifade değişikliği → sayıları koru.
            import re as _re
            _VALUE_KW = ("değer", "sayı", "sayısal", "süre", "eşik", "sınır", "menzil", "mesafe",
                         "sıcaklık", "hız", "frekans", "oran", "gecikme", "tolerans", "doğruluk",
                         "hassasiyet", "kapasite", "aralık", "metrik", "ölçülebilir", "birim",
                         "limit", "seviye", "band", "bant", "voltaj", "akım", "güç", "boyut", "ağırlık")
            # ÖNEMLİ: "SR-001" gibi ID'lerin içindeki rakamları (001) "kullanıcı sayı verdi"
            # sanmamak için, sayı kontrolünden ÖNCE mesajdan ID'leri temizle.
            _msg_no_id = _re.sub(_re.escape(target), "", msg, flags=_re.I) if target else msg
            _msg_no_id = _re.sub(r"[A-Za-zÇĞİÖŞÜçğıöşü]{2,}(?:-[A-Za-zÇĞİÖŞÜçğıöşü]+)*-\d+", "", _msg_no_id)
            _has_number = bool(_re.search(r"\d", _msg_no_id))
            _asks_value = any(k in msg.lower() for k in _VALUE_KW)
            if _has_number:
                dsb_directive = ("NOT: Kullanıcı somut sayısal değer(ler) verdi. Bu değerleri AYNEN "
                                 "kullan. Bu istekte 'DSB' YAZMA.")
            elif _asks_value:
                dsb_directive = ("NOT: Kullanıcı bir sayısal değer/eşik istiyor ama sayının kendisini "
                                 "VERMEDİ. O değerin yerine MUTLAKA 'DSB' yaz (birimi koru) — sayı UYDURMA.")
            else:
                dsb_directive = ("NOT: Bu istek yalnızca ifade/ton değişikliğidir. Mevcut metindeki TÜM "
                                 "sayıları ve değerleri AYNEN KORU; yeni sayı ekleme, 'DSB' YAZMA.")
            prompt = (
                f"MADDE ID: {target}\n"
                f"MEVCUT METİN: \"{old}\"\n\n"
                f"KULLANICI İSTEĞİ: {msg}\n\n"
                f"{dsb_directive}\n\n"
                "Bu maddeyi yukarıdaki nota ve sistem mühendisliği kurallarına göre revize et. "
                "Yanıt olarak SADECE revize edilmiş madde metnini ver."
            )
            resp = call_gemma3_api(prompt, max_tokens=280, temperature=0.1,
                                   system_message=self.COPILOT_SYSTEM_PROMPT)
            if not resp or not resp.strip():
                self._chat_append(f"{target} için model cevap vermedi. LM Studio açık ve model yüklü mü?", "err")
                return
            new = self._clean_revision(resp)
            # Model, kullanıcı mesajındaki madde ID'sini gövdeye kopyalıyor
            # ("... AT-014 test senaryosunda karşılanmalıdır"). Madde metni kendi ID'sini
            # içermemeli (ID zaten ayrı sütunda) → hedef ID'yi metinden temizle.
            # Ek YALNIZCA kesme işaretiyle bitişikse silinir ("SR-002'nin"); aksi halde
            # sonraki kelimeden harf yenir ("AT-014 test" → "t")!
            _id_pat = _re.compile(r"\s*\b" + _re.escape(target) + r"\b(?:['’][a-zçğıöşü]{1,4})?\s*", _re.I)
            new = _id_pat.sub(" ", new)
            new = _re.sub(r"\s{2,}", " ", new).strip(" ,;:")
            if "DSB" in new.upper():
                new = text_cleanup.dsb_temizle(new)   # DSB varsa çelişen uydurma sayıları temizle
            self._apply_revision(target, old, new)

            # Ana konsolda da GÖRÜNÜR yap (kullanıcı genelde sol tarafa bakıyor)
            self.update_status_text(f"\n━━━ COPILOT · {target} GÜNCELLENDİ ━━━", is_complete=True)
            self.update_status_text(f"ESKİ: {old}")
            self.update_status_text(f"YENİ: {new}", is_complete=True)

            # --- DEĞİŞİKLİK ETKİ ANALİZİ: bir GEREKSİNİM değişince, ona bağlı test(ler)i
            #     yeni gereksinim metnine göre YENİDEN ÜRET (kullanıcı tercihi).
            #     Gereksinimde DSB varsa üretilen teste de DSB notu eklenir. ---
            extra = ""
            architecture_changed_ids = {target}
            if self.flat_data[target].get("type") in self.REQ_TYPES:
                self._chat_append(f"{target} değişti → bağlı test(ler) yeniden üretiliyor...", "info")
                affected = self._ripple_regenerate(target)
                if affected:
                    architecture_changed_ids.update(affected)
                    extra = ("\n🔗 Bağlı test(ler) yeni gereksinime göre güncellendi: "
                             + ", ".join(affected))
            self.master.after(
                0,
                lambda ids=tuple(sorted(architecture_changed_ids)):
                    self._notify_architecture_sources_changed(ids),
            )
            self._chat_append(f"✅ {target} güncellendi:\n{new}{extra}\n\n"
                              "(İndirdiğinde pdf/excel/html/word çıktısına yansır.)", "bot")
        except Exception as e:
            self._chat_append(f"Hata: {e}", "err")
        finally:
            self.master.after(0, lambda: self.chat_send_btn.config(state=tk.NORMAL))

    def _clean_revision(self, text):
        """4B modelin eklediği başlık/etiket/markdown/liste artıklarını temizler, tek paragraf yapar."""
        import re
        text = (text or "").strip().strip('"').strip()
        # markdown vurgu ve işaretlerini kaldır
        text = re.sub(r'[*_`#]+', '', text)
        # uydurma standart kodlarını (örn. "(GS-005)") her yerden temizle
        text = re.sub(r'\s*\((?:GS|STD|MIL-?STD|DS|TS)[-\s]?\d+\)', '', text, flags=re.I)

        label = re.compile(r'^(madde\s*id|madde|revizyon|revize|id|çıktı|cikti|cevap|output|sonuç|sonuc|not)\b[\s:–-]*', re.I)
        verb = re.compile(r'(malı|meli|olmalı|etmeli|dır|dir|dur|dür)\b', re.I)

        kept = []
        for ln in text.split("\n"):
            s = ln.strip().lstrip("-•·*").strip()
            if not s:
                continue
            low = s.lower().replace("̇", "")  # Türkçe İ combining-dot'u temizle
            if label.match(low):
                # etiketi kırp; geriye anlamlı bir gereksinim cümlesi kalıyorsa tut, yoksa (başlık) at
                stripped = label.sub('', s).strip()
                if verb.search(stripped.lower()) and len(stripped.split()) >= 4:
                    kept.append(stripped)
                continue
            kept.append(s)

        out = " ".join(kept)
        out = re.sub(r'\s+', ' ', out)
        out = re.sub(r'\s+([.,;:])', r'\1', out)   # " ." -> "."
        return out.strip(' -–•:')

    def _sync_item_text(self, item_id, old_text, new_text):
        """Bir maddenin yeni metnini ham çıktı metnine ve ilgili last_*_list'e yansıtır."""
        # 1) ham metin (txt çıktısı / konsol tutarlılığı)
        #    ID'ye BAĞLI satırı değiştir ("ID | metin" formatı). Kör replace kullanılırsa,
        #    iki maddenin metni birebir AYNI olduğunda ikisi birden değişirdi → yanlış madde bozulur.
        if old_text and (self.last_generated_output or ""):
            eski_satir = f"{item_id} | {old_text}"
            yeni_satir = f"{item_id} | {new_text}"
            if eski_satir in self.last_generated_output:
                self.last_generated_output = self.last_generated_output.replace(eski_satir, yeni_satir)
            elif old_text in self.last_generated_output:
                # yedek yol: sadece İLK eşleşme (diğer maddelere bulaşmasın)
                self.last_generated_output = self.last_generated_output.replace(old_text, new_text, 1)
        # 2) ilgili last_*_list (sınıflandırma tutarlılığı) — best effort
        t = self.flat_data.get(item_id, {}).get("type", "")
        list_map = {
            "TID":       (self.last_tid_list, "TID_ID", "TID_Aciklama"),
            "SGD":       (self.last_sgd_list, "SGD_ID", "SGD_Aciklama"),
            "STT":       (self.last_stt_list, "STT_ID", "STT_Aciklama"),
            "SITET":     (self.last_sitet_list, "SITET_ID", "SITET_Aciklama"),
            "AST":       (self.last_alt_sistem_test_list, "AST_ID", "AST_Aciklama"),
        }
        if t in list_map:
            lst, idk, txtk = list_map[t]
            for it in lst:
                if it.get(idk) == item_id:
                    it[txtk] = new_text
                    break

    def _apply_revision(self, target, old, new):
        """Revizyonu flat_data + ham metin + ilgili last_*_list üzerinde uygular."""
        self._notify_architecture_source_mutation_started()
        self.flat_data[target]["content"] = new     # tüm çıktıların (pdf/excel/html/docx) kaynağı
        self._sync_item_text(target, old, new)
        self.update_status_text(f"[Copilot] {target} güncellendi.", is_complete=True)

    def _ripple_regenerate(self, requirement_id):
        """
        DEĞİŞİKLİK ETKİ ANALİZİ (TAM KASKAD): Bir gereksinim revize edilince, ona bağlı
        HEM alt gereksinimleri HEM testleri yeni metne göre yeniden üretir. Alt gereksinimler
        için işlem özyinelemeli olarak aşağı iner (UR→SR→SSR ve her birinin testi).
        Üst maddede 'DSB' varsa, türeyen tüm maddelerde de ilgili değer DSB olur (uydurma yok).
        """
        self._notify_architecture_source_mutation_started()
        child_req_gen = {
            "TID": sgd_generator_logic.generate_sgd_from_ur,
            "SGD": stt_generator_logic.generate_subsystem_req_from_sgd,
        }
        test_gen = {
            "TID": kmtd_generator_logic.generate_kmtd_from_tid,
            "SGD": sitet_generator_logic.generate_sitet_from_sgd,
            "STT": alt_sistem_test_logic.generate_subsystem_test,
        }
        dsb_kural = (" [KURAL: Yukarıdaki maddede 'DSB' geçen değer belirsizdir; türeteceğin "
                     "maddede de o değer için sayı UYDURMA, yerine 'DSB' yaz.]")
        note = (" (Not: Üst gereksinimde DSB bulunduğu için bu maddede de ilgili değer DSB'dir.)")
        affected, visited = [], set()

        def kaskad(parent_id):
            if parent_id in visited:
                return
            visited.add(parent_id)
            parent = self.flat_data.get(parent_id, {})
            ptype = parent.get("type")
            ptext = parent.get("content", "")
            dsb = "DSB" in (ptext or "").upper()
            girdi = (ptext + dsb_kural) if dsb else ptext
            cocuklar = [(k, v) for k, v in list(self.flat_data.items())
                        if v.get("bound_to") == parent_id]
            for iid, it in cocuklar:
                itype = it.get("type")
                if itype in self.TEST_TYPES:
                    gen_fn = test_gen.get(ptype)
                elif itype in self.REQ_TYPES:
                    gen_fn = child_req_gen.get(ptype)
                else:
                    gen_fn = None
                if not gen_fn:
                    continue
                try:
                    raw = gen_fn(girdi, "Proje")
                except Exception as e:
                    self.update_status_text(f"[Copilot] {iid} yeniden üretilemedi: {e}", is_error=True)
                    continue
                if not raw:
                    continue
                yeni = text_cleanup.temizle(raw, test=(itype in self.TEST_TYPES))
                if not yeni:
                    continue
                if dsb:
                    # DSB ile çelişen uydurma sayı/örnekleri temizle ('(örneğin 100g)', '100g DSB'...)
                    yeni = text_cleanup.dsb_temizle(yeni)
                    if "DSB" not in yeni.upper():
                        # Model DSB yerine sayı uydurdu → çelişkili not yerine sayıları DSB yap
                        donusmus = text_cleanup.sayilari_dsb_yap(yeni)
                        yeni = donusmus if "DSB" in donusmus.upper() else (yeni.rstrip() + note).strip()
                eski = it.get("content", "")
                it["content"] = yeni
                self._sync_item_text(iid, eski, yeni)
                affected.append(iid)
                if itype in self.REQ_TYPES:   # alt gereksinimin de altını güncelle
                    kaskad(iid)

        kaskad(requirement_id)
        if affected:
            self.update_status_text(
                f"[Copilot] {requirement_id} değişti → bağlı madde(ler) güncellendi: "
                f"{', '.join(affected)}", is_complete=True)
        return affected

    def _split_requirement(self, target):
        """Atomik olmayan bir maddeyi ('...ve...ve...') ayrı gereksinimlere böler.
        İlk parça orijinali günceller; diğerleri yeni ID'lerle (SR-002b, SR-002c...) eklenir,
        aynı üst maddeye bağlanır. Her yeni gereksinim parçasına bağlı bir test de üretilir."""
        from llm_handler import call_gemma3_api
        item = self.flat_data.get(target, {})
        content = item.get("content", "")
        itype = item.get("type", "")
        bound = item.get("bound_to", "")
        self._chat_append(f"{target} atomik parçalara bölünüyor...", "info")

        prompt = (
            "Aşağıdaki gereksinim maddesi birden fazla gereksinim içeriyor (atomik değil). "
            "Bunu, her biri TEK ve bağımsız bir gereksinim olan AYRI cümlelere böl.\n\n"
            f"MADDE: \"{content}\"\n\n"
            "Kurallar:\n"
            "- Her satıra yalnızca 1 gereksinim yaz; numara, işaret, etiket KOYMA.\n"
            "- Her biri tam bir 'sistem ... -malıdır' cümlesi olsun.\n"
            "- Orijinal anlamı KORU; yeni gereksinim uydurma, hiçbirini atlama.\n"
        )
        resp = call_gemma3_api(prompt, max_tokens=300, temperature=0.1,
                               system_message=self.COPILOT_SYSTEM_PROMPT)
        parts = []
        for ln in (resp or "").split("\n"):
            p = text_cleanup.temizle(ln)
            if p and len(p.split()) >= 3 and p not in parts:
                parts.append(p)
        if len(parts) < 2:
            self._chat_append(f"{target} bölünemedi (tek gereksinim gibi görünüyor).", "err")
            return

        test_gen = {
            "TID": kmtd_generator_logic.generate_kmtd_from_tid,
            "SGD": sitet_generator_logic.generate_sitet_from_sgd,
            "STT": alt_sistem_test_logic.generate_subsystem_test,
        }
        test_tip = {"TID": "KMTD", "SGD": "SITET", "STT": "AST"}
        test_pre = {"TID": "AT", "SGD": "SITET", "STT": "SST"}

        prefix, _, numpart = target.rpartition("-")
        olusan = [target]

        # 1) İlk parça orijinali günceller + (gereksinimse) testini senkronla
        self._apply_revision(target, content, parts[0])
        architecture_changed_ids = {target}
        if itype in self.REQ_TYPES:
            ripple_changed_ids = self._ripple_regenerate(target) or ()
            architecture_changed_ids.update(
                item_id for item_id in ripple_changed_ids
                if self.flat_data.get(item_id, {}).get("type") in self.REQ_TYPES
            )

        # 2) Diğer parçalar → yeni maddeler (+ bağlı test)
        harf = "bcdefghi"
        for i, p in enumerate(parts[1:]):
            son = harf[i] if i < len(harf) else str(i + 2)
            yid = f"{prefix}-{numpart}{son}"
            while yid in self.flat_data:
                yid += "x"
            self.flat_data[yid] = {"type": itype, "ID": yid, "content": p, "bound_to": bound}
            self.last_generated_output += f"\n{yid} | {p}"
            olusan.append(yid)
            architecture_changed_ids.add(yid)
            gen = test_gen.get(itype)
            if gen:
                try:
                    tt = text_cleanup.temizle(gen(p, "Proje") or "", test=True)
                except Exception:
                    tt = ""
                if tt:
                    tid_ = f"{test_pre[itype]}-{numpart}{son}"
                    while tid_ in self.flat_data:
                        tid_ += "x"
                    self.flat_data[tid_] = {"type": test_tip[itype], "ID": tid_,
                                            "content": tt, "bound_to": yid}
                    self.last_generated_output += f"\n{tid_} | {tt}"

        self.update_status_text(f"\n━━━ COPILOT · {target} BÖLÜNDÜ ({len(parts)} parça) ━━━",
                                is_complete=True)
        for pid in olusan:
            self.update_status_text(f"{pid}: {self.flat_data[pid]['content']}")
        self._chat_append(
            f"✅ {target} {len(parts)} atomik gereksinime bölündü: " + ", ".join(olusan)
            + "\n(Her yeni parçaya bağlı test de üretildi. İndirdiğinde çıktılara yansır.)", "bot")
        self.master.after(
            0,
            lambda ids=tuple(sorted(architecture_changed_ids)):
                self._notify_architecture_sources_changed(ids),
        )

    def _ripple_dsb(self, requirement_id):
        """
        Paydaş kuralı 5 — İZLENEBİLİRLİK YAYILIMI (ripple effect):
        Bir gereksinim maddesi 'DSB' içerecek şekilde revize edildiğinde, ona BAĞLI olan
        test maddelerini (KMTD/SITET/AST/DTET-YTET) bulur ve onların test kriterine de
        'DSB' notu ekler. Bağ ilişkisi: test maddesinin bound_to == gereksinim ID'si.
        Etkilenen test ID'lerinin listesini döndürür.
        """
        note = " (Not: İlgili gereksinimde DSB bulunduğu için test kriteri de DSB'dir.)"
        affected = []
        for item_id, d in list(self.flat_data.items()):
            if d.get("type") in self.TEST_TYPES and d.get("bound_to") == requirement_id:
                old_c = d.get("content", "")
                if "DSB" in old_c.upper():
                    continue  # zaten DSB'li → tekrar ekleme
                new_c = (old_c.rstrip() + note).strip()
                d["content"] = new_c
                self._sync_item_text(item_id, old_c, new_c)
                affected.append(item_id)
        if affected:
            self.update_status_text(
                f"[Copilot] DSB yayılımı → {requirement_id} bağlı test(ler)i güncellendi: "
                f"{', '.join(affected)}", is_complete=True)
        return affected

    # ------------------------------------------------------------------ #
    #  EXCEL / HTML / WORD çıktıları (hepsi flat_data + VMODEL_SECTIONS)  #
    # ------------------------------------------------------------------ #
    def _export_excel(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile çok sayfalı, biçimli Excel üretir."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def style_header(ws, ncols, color="1F3864"):
            fill = PatternFill("solid", fgColor=color)
            for c in range(1, ncols + 1):
                cell = ws.cell(row=1, column=c)
                cell.fill = fill
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

        wb = Workbook()
        # --- Kapak ---
        info = wb.active
        info.title = "Proje Bilgisi"
        info.append(["V-MODEL İZLENEBİLİRLİK RAPORU"])
        info.append(["Proje", proje_ismi])
        info.append(["Tarih", datetime.now().strftime("%d.%m.%Y %H:%M")])
        info.append(["Toplam Madde", len(self.flat_data)])
        info["A1"].font = Font(bold=True, size=14, color="1F3864")
        info.column_dimensions["A"].width = 20
        info.column_dimensions["B"].width = 50

        leg_tr = {"req": "Gereksinim", "test": "Test"}

        # --- İzlenebilirlik Matrisi ---
        ws = wb.create_sheet("Izlenebilirlik Matrisi")
        ws.append(["ID", "Tür", "Bacak", "Açıklama", "Bağlı Olduğu"])
        for type_key, title, leg in self.VMODEL_SECTIONS:
            for d in self.flat_data.values():
                if d.get("type") == type_key:
                    ws.append([d.get("ID", ""), type_key, leg_tr.get(leg, leg),
                               d.get("content", ""), d.get("bound_to", "Yok")])
        style_header(ws, 5)
        for col, w in zip("ABCDE", [16, 12, 12, 80, 22]):
            ws.column_dimensions[col].width = w
        for r in range(2, ws.max_row + 1):
            ws.cell(row=r, column=4).alignment = Alignment(wrap_text=True, vertical="top")

        # --- Her tür için ayrı sayfa ---
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            safe = type_key.replace("/", "-")[:31]
            sh = wb.create_sheet(safe)
            sh.append([f"{type_key} ID", "Açıklama", "Bağlı Olduğu"])
            for d in items:
                sh.append([d.get("ID", ""), d.get("content", ""), d.get("bound_to", "Yok")])
            style_header(sh, 3, "1F3864" if leg == "req" else "375623")
            for col, w in zip("ABC", [16, 90, 22]):
                sh.column_dimensions[col].width = w
            for r in range(2, sh.max_row + 1):
                sh.cell(row=r, column=2).alignment = Alignment(wrap_text=True, vertical="top")

        wb.save(path)

    def _export_html(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile bağımsız, biçimli HTML raporu üretir."""
        import html as _html
        esc = lambda s: _html.escape(str(s))
        leg_color = {"req": "#1F3864", "test": "#375623"}
        leg_light = {"req": "#D6E0F0", "test": "#E2EFDA"}

        p = []
        p.append("<!DOCTYPE html><html lang='tr'><head><meta charset='utf-8'>")
        p.append(f"<title>{esc(proje_ismi)} - V-Model Doküman Paketi</title>")
        p.append("<style>"
                 "body{font-family:'Segoe UI',Arial,sans-serif;margin:24px;color:#222;background:#f7f9fc}"
                 "h1{color:#1F3864;margin-bottom:0}.sub{color:#666;margin:2px 0 14px}"
                 "table{border-collapse:collapse;width:100%;margin-bottom:22px;background:#fff;"
                 "box-shadow:0 1px 3px rgba(0,0,0,.08)}"
                 "th,td{border:1px solid #ccc;padding:7px 9px;text-align:left;vertical-align:top;font-size:14px}"
                 ".sec{color:#fff;padding:8px 12px;font-weight:bold;font-size:15px;margin-top:8px;border-radius:3px 3px 0 0}"
                 "td.id{font-weight:bold;white-space:nowrap}"
                 ".legend span{display:inline-block;margin-right:18px}"
                 ".chip{display:inline-block;width:12px;height:12px;border-radius:2px;margin-right:5px;vertical-align:middle}"
                 "</style></head><body>")
        p.append("<h1>V-Model Teknik Doküman Paketi</h1>")
        p.append(f"<div class='sub'>Proje: <b>{esc(proje_ismi)}</b> &nbsp;|&nbsp; "
                 f"{datetime.now().strftime('%d.%m.%Y %H:%M')} &nbsp;|&nbsp; Toplam {len(self.flat_data)} madde</div>")
        p.append("<div class='legend'>"
                 "<span><span class='chip' style='background:#1F3864'></span>Sol Bacak: Gereksinim</span>"
                 "<span><span class='chip' style='background:#375623'></span>Sağ Bacak: Test / Doğrulama</span></div><br>")

        any_sec = False
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_sec = True
            col = leg_color.get(leg, "#333"); light = leg_light.get(leg, "#eee")
            bound_h = "Kaynak / Bağlı" if leg == "req" else "Doğruladığı Madde"
            p.append(f"<div class='sec' style='background:{col}'>{esc(title)} ({len(items)} madde)</div>")
            p.append(f"<table><tr style='background:{col};color:#fff'>"
                     f"<th>ID</th><th>Açıklama</th><th>{bound_h}</th></tr>")
            for i, d in enumerate(items):
                bg = "#ffffff" if i % 2 == 0 else light
                p.append(f"<tr style='background:{bg}'><td class='id'>{esc(d.get('ID',''))}</td>"
                         f"<td>{esc(d.get('content',''))}</td><td>{esc(d.get('bound_to','-'))}</td></tr>")
            p.append("</table>")
        if not any_sec:
            p.append("<p><i>Üretilen doküman verisi bulunamadı.</i></p>")
        p.append("</body></html>")

        with open(path, "w", encoding="utf-8") as f:
            f.write("".join(p))

    def _export_docx(self, path, proje_ismi, template_path=None):
        """Kurumsal .docx şablonunu doldurur (veya sıfırdan Word üretir). Başarı=True."""
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Eksik Kütüphane",
                "Word (.docx) çıktısı için 'python-docx' gerekli.\n\n"
                "Terminalde şunu çalıştırın:\n    pip install python-docx")
            return False

        doc = Document(template_path) if template_path else Document()

        if template_path:
            mapping = {"{{PROJE_ADI}}": proje_ismi, "{{TARIH}}": datetime.now().strftime("%d.%m.%Y")}
            for para in doc.paragraphs:
                for k, v in mapping.items():
                    if k in para.text:
                        para.text = para.text.replace(k, v)
            doc.add_page_break()
        else:
            doc.add_heading("V-MODEL TEKNİK DOKÜMAN PAKETİ", level=0)
            doc.add_paragraph(f"Proje: {proje_ismi}")
            doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            doc.add_heading(title, level=2)
            t = doc.add_table(rows=1, cols=3)
            try:
                t.style = "Light Grid Accent 1"
            except Exception:
                t.style = "Table Grid"
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = "ID", "Açıklama", "Bağlı Olduğu"
            for d in items:
                c = t.add_row().cells
                c[0].text = str(d.get("ID", ""))
                c[1].text = str(d.get("content", ""))
                c[2].text = str(d.get("bound_to", "Yok"))
        doc.save(path)
        return True

    def _export_doors_csv(self, path, proje_ismi):
        """V-Model maddelerini IBM DOORS'a aktarılabilir UTF-8 CSV olarak üretir."""
        leg_names = {"req": "Requirement", "test": "Test"}

        with open(path, "w", encoding="utf-8-sig", newline="") as csv_file:
            writer = csv.writer(
                csv_file,
                dialect="excel",
                quoting=csv.QUOTE_ALL,
                lineterminator="\r\n",
            )
            writer.writerow([
                "Object Identifier",
                "Object Text",
                "Requirement Type",
                "V-Model Leg",
                "Linked Object",
                "Project",
            ])

            for type_key, _title, leg in self.VMODEL_SECTIONS:
                for item in self.flat_data.values():
                    if item.get("type") != type_key:
                        continue
                    writer.writerow([
                        item.get("ID", ""),
                        item.get("content", ""),
                        type_key,
                        leg_names.get(leg, leg),
                        item.get("bound_to", "Yok"),
                        proje_ismi,
                    ])

    def download_docs(self):
        fmt = self.format_combo.get().lower()
        if fmt != "şablon" and not self.last_generated_output and not self.flat_data:
            messagebox.showwarning("Uyarı", "İndirilecek içerik yok. Önce doküman üretin.")
            return

        proje_ismi = self.entry_widgets["proje_ismi"].get().strip().replace(" ", "_") or "yapay_zeka_dokumanlar"

        if fmt == "excel":
            ext, ftypes = ".xlsx", [("Excel Workbook", "*.xlsx")]
        elif fmt == "doors":
            ext, ftypes = ".csv", [("IBM DOORS CSV", "*.csv")]
        elif fmt in ("şablon", "docx"):
            ext, ftypes = ".docx", [("Word Document", "*.docx")]
        else:
            ext, ftypes = f".{fmt}", [(f"{fmt.upper()} files", f"*.{fmt}")]

        init_name = f"{proje_ismi}_Dokumanlar{ext}"
        path = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=ftypes + [("All files", "*.*")],
            initialfile=init_name
        )
        if not path:
            return

        try:
            if fmt == "txt":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.last_generated_output or "")

            elif fmt == "pdf":
                self._generate_vmodel_pdf(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "excel":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Excel'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_excel(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "html":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "HTML'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_html(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "doors":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "DOORS'a aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_doors_csv(
                    path,
                    self.entry_widgets["proje_ismi"].get().strip() or "Proje"
                )

            elif fmt in ("şablon", "docx"):
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Word'e yerleştirilecek veri yok. Önce doküman üretin.")
                    return
                tmpl = self.template_file_path if fmt == "şablon" else None
                if fmt == "şablon" and not tmpl:
                    messagebox.showerror("Hata", "Şablon dosyası seçilmedi (.docx).")
                    return
                ok = self._export_docx(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje", tmpl)
                if not ok:
                    return

            resolved_path = os.path.abspath(path)
            if resolved_path not in self.generated_document_paths:
                self.generated_document_paths.append(resolved_path)
            messagebox.showinfo("Başarılı", f"Dosya kaydedildi:\n{path}")
        except Exception as e:
            messagebox.showerror("Hata", f"Kaydetme hatası: {e}")
            self.update_status_text(f"Kaydetme hatası: {e}", is_error=True)

    def start_generation(self):
        if not self.file_paths:
            messagebox.showerror("Hata", "Girdi dosyalarını seçin.")
            return

        try:
            proje_ismi = self.entry_widgets["proje_ismi"].get().strip()
            if not proje_ismi:
                raise ValueError("Proje İsmi boş bırakılamaz.")

            doc_counts = {
                "max_tids": int(self.entry_widgets["teknik_ister"].get() or 0),
                "max_sgds": int(self.entry_widgets["sistem_gereksinimi"].get() or 0),
                "max_stts": int(self.entry_widgets["sistem_tanimlama_testi"].get() or 0),
            }

            doc_flags = {
                "generate_kmtd": self.checkbox_vars["generate_kmtd"].get(),
                "generate_sitet": self.checkbox_vars["generate_sitet"].get(),
                "generate_alt_sistem_testi": self.checkbox_vars["generate_alt_sistem_testi"].get(),
            }

            if sum(v > 0 for v in doc_counts.values()) == 0 and not any(doc_flags.values()):
                raise ValueError("En az bir doküman sayısı veya test seçilmeli.")

        except ValueError as e:
            messagebox.showerror("Hata", f"Geçersiz giriş: {e}")
            return

        # Eski mimari yayımı, kaynak sözlüğü temizlenmeden önce iptal
        # edilmelidir; aksi halde publish worker eski snapshot'ı yeni kaynak
        # revizyonunun ``latest`` sürümü olarak commit edebilir.
        self._notify_architecture_generation_started()
        self.last_generated_output = ""
        self._traceability_generation_token += 1
        self._traceability_cancel_event.set()
        self.tree_data.clear()
        self.flat_data.clear()
        # Yeni ``flat_data`` worker tarafından parça parça doldurulur. Bu
        # sürede önceki rapor yeni kaynak setine aitmiş gibi kullanılamaz.
        self.last_traceability_report = None
        self.last_traceability_health = None
        self.hardware_data.clear()
        self.generated_document_paths.clear()
        self.last_hardware_catalog = None
        self.last_hardware_catalog_status = None
        self.last_hardware_impact_result = None
        self._invalidate_hardware_generation()
        self._refresh_hardware_workspace()
        self._refresh_hardware_cards_workspace()
        self.update_status_text("--- YAPAY ZEKA ÜRETİMİ BAŞLADI ---\n", clear=True)

        self.create_docs_button.config(state=tk.DISABLED, text=self._t("İŞLENİYOR...", "PROCESSING..."), style="success.TButton")
        self.download_docs_button.config(state=tk.DISABLED)

        thread = threading.Thread(
            target=self.run_ai_process,
            args=(self.file_paths, doc_counts, doc_flags,
                  self.format_combo.get().lower(), proje_ismi)
        )
        thread.start()

    def run_ai_process(self, file_paths, doc_counts, doc_flags, output_format, proje_ismi):
        # Bu metot test/entegrasyon tarafından ``start_generation`` dışında
        # doğrudan çağrılsa bile ilk kaynak yazımından önce yayım kapansın.
        self._notify_architecture_source_mutation_started()
        total_start_time = time.time()

        # NOT: Alt Sistem Gereksinimleri (max_stts) artık kaynak dosyadan (chunk)
        # değil, üretilen SGD listesinden türetildiği için bu kontrole dahil değildir.
        kaynak_dokuman_gerekli = (
            doc_counts["max_tids"] > 0 or
            doc_counts["max_sgds"] > 0
        )

        all_chunks = None
        sorted_indices = None

        if kaynak_dokuman_gerekli:
            all_chunks, sorted_indices = pre_process_files(file_paths, self.update_status_text)
            
            if not all_chunks:
                self.update_status_text("Hata: Kaynak dosyalardan veri alınamadı, işlem durduruluyor.", is_error=True)
                self.master.after(
                    0,
                    lambda: self._notify_architecture_generation_failed(
                        "Kaynak dosyalardan veri alınamadı."
                    ),
                )
                self.master.after(0, lambda: self._reset_buttons_state())
                return

        try:
            if doc_counts.get("max_tids", 0) > 0:
                self.update_status_text("Kullanıcı Gereksinimi üretimi başlıyor...")
                result = tid_generator_logic.run_generation_logic(
                    file_paths=None,
                    max_tids=doc_counts["max_tids"],
                    output_format=output_format,
                    project_name=proje_ismi,
                    status_callback=self.update_status_text,
                    precomputed_chunks=all_chunks,
                    precomputed_indices=sorted_indices
                )
                if result.get("result"):
                    self.last_tid_list = result.get("tid_list", [])
                    output = f"\n--- KULLANICI GEREKSİNİMİ (User Requirement) --- {proje_ismi} ---\n\n"
                    for item in self.last_tid_list:
                        output += f"{item['TID_ID']} | {item['TID_Aciklama']}\n"
                        self.flat_data[item['TID_ID']] = {
                            'type': 'TID', 'bound_to': 'Yok',
                            'ID': item['TID_ID'], 'content': item['TID_Aciklama']
                        }
                    self.last_generated_output += output
                    self.update_status_text(output, is_complete=True)

            if doc_flags["generate_kmtd"]:
                if self.last_tid_list:
                    self.update_status_text("KMTD üretimi başlıyor...")
                    try:
                        result = kmtd_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_tid_list,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            kmtd_list = result.get("kmtd_list", [])
                            output = f"\n\n--- KABUL MUAYENE TESTİ (Acceptance Test) --- {proje_ismi} ---\n\n"
                            for item in kmtd_list:
                                output += f"{item['KMTD_ID']} | {item['KMTD_Aciklama']}\n"
                                self.flat_data[item['KMTD_ID']] = {
                                    'type': 'KMTD', 'bound_to': item['Bound_TID'],
                                    'ID': item['KMTD_ID'], 'content': item['KMTD_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"KMTD Hatası: {e}", is_error=True)

            if doc_counts.get("max_sgds", 0) > 0:
                self.update_status_text("SGD üretimi başlıyor...")
                if self.last_tid_list:
                    # İZLENEBİLİRLİK: her Kullanıcı Gereksiniminden (UR) türeyen SGD (Bound_TID)
                    result = sgd_generator_logic.run_generation_from_requirements(
                        requirement_list=self.last_tid_list,
                        max_sgds=doc_counts["max_sgds"],
                        project_name=proje_ismi,
                        status_callback=self.update_status_text
                    )
                else:
                    # UR yoksa eski yöntem: doğrudan kaynak dosyadan (genel bağ)
                    result = sgd_generator_logic.run_generation_logic(
                        file_paths=None,
                        max_sgds=doc_counts["max_sgds"],
                        output_format=output_format,
                        project_name=proje_ismi,
                        status_callback=self.update_status_text,
                        precomputed_chunks=all_chunks,
                        precomputed_indices=sorted_indices
                    )
                if result.get("result"):
                    self.last_sgd_list = result.get("sgd_list", [])
                    output = f"\n\n--- SİSTEM GEREKSİNİMİ (System Requirements) --- {proje_ismi} ---\n\n"
                    for item in self.last_sgd_list:
                        output += f"{item['SGD_ID']} | {item['SGD_Aciklama']}\n"
                        self.flat_data[item['SGD_ID']] = {
                            'type': 'SGD', 'bound_to': item.get('Bound_TID', 'TID-Genel'),
                            'ID': item['SGD_ID'], 'content': item['SGD_Aciklama']
                        }
                    self.last_generated_output += output
                    self.update_status_text(output, is_complete=True)

            if doc_flags["generate_sitet"]:
                if self.last_sgd_list:
                    self.update_status_text("SITET üretimi başlıyor...")
                    try:
                        if 'sitet_generator_logic' in sys.modules:
                            result = sitet_generator_logic.run_generation_from_requirements(
                                requirement_list=self.last_sgd_list,
                                project_name=proje_ismi,
                                status_callback=self.update_status_text
                            )
                            if result.get("result"):
                                self.last_sitet_list = result.get("sitet_list", [])
                                output = f"\n\n--- Sistem Testi (System Test) LİSTESİ --- {proje_ismi} ---\n\n"
                                for item in self.last_sitet_list:
                                    output += f"{item['SITET_ID']} | {item['SITET_Aciklama']}\n"
                                    self.flat_data[item['SITET_ID']] = {
                                        'type': 'SITET', 'bound_to': item.get('Bound_SGD', 'SGD'),
                                        'ID': item['SITET_ID'], 'content': item['SITET_Aciklama']
                                    }
                                self.last_generated_output += output
                                self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"SITET Hatası: {e}", is_error=True)

            if doc_counts.get("max_stts", 0) > 0:
                self.update_status_text("Alt Sistem Gereksinimleri üretimi başlıyor...")
                if not self.last_sgd_list:
                    self.update_status_text(
                        "Alt Sistem Gereksinimleri için önce Sistem Gereksinimi (SGD) üretmelisiniz.",
                        is_error=True
                    )
                else:
                    try:
                        result = stt_generator_logic.run_generation_from_requirements(
                            requirement_list=self.last_sgd_list,
                            max_stts=doc_counts["max_stts"],
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            self.last_stt_list = result.get("stt_list", [])
                            output = f"\n\n--- ALT SİSTEM GEREKSİNİMLERİ (Subsystem Requirements) --- {proje_ismi} ---\n\n"
                            for item in self.last_stt_list:
                                output += f"{item['STT_ID']} | {item['STT_Aciklama']}\n"
                                self.flat_data[item['STT_ID']] = {
                                    'type': 'STT', 'bound_to': item.get('Bound_SGD', 'SGD'),
                                    'ID': item['STT_ID'], 'content': item['STT_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"Alt Sistem Gereksinimleri Hatası: {e}", is_error=True)

            if doc_flags["generate_alt_sistem_testi"]:
                if self.last_stt_list:
                    self.update_status_text("Alt Sistem Testi üretimi başlıyor...")
                    try:
                        result = alt_sistem_test_logic.run_generation_from_requirements(
                            requirement_list=self.last_stt_list,
                            project_name=proje_ismi,
                            status_callback=self.update_status_text
                        )
                        if result.get("result"):
                            self.last_alt_sistem_test_list = result.get("ast_list", [])
                            output = f"\n\n--- ALT SİSTEM TESTİ (Subsystem Test) LİSTESİ --- {proje_ismi} ---\n\n"
                            for item in self.last_alt_sistem_test_list:
                                output += f"{item['AST_ID']} | {item['AST_Aciklama']}\n"
                                self.flat_data[item['AST_ID']] = {
                                    'type': 'AST', 'bound_to': item.get('Bound_STT', 'ASG'),
                                    'ID': item['AST_ID'], 'content': item['AST_Aciklama']
                                }
                            self.last_generated_output += output
                            self.update_status_text(output, is_complete=True)
                    except Exception as e:
                        self.update_status_text(f"Alt Sistem Testi Hatası: {e}", is_error=True)

            # --- ÜRETİM SONRASI DSB DÜZELTMESİ (merkezi) ---
            # DSB SADECE 'değer belli değil' işaretidir; standart/sıfat/kap adı DEĞİLDİR.
            # (a) Model yanlış kullandıysa temizle. (b) Gereksinimde DSB varsa, ona bağlı test
            #     de o değeri DSB olarak taşısın (uydurma sayı → DSB).
            try:
                # (1) HER maddede: etiket/numara/markdown artıklarını temizle ('DONANIM/YAZILIM:',
                #     baştaki '1.' gibi — text_cleanup kullanmayan generator'lar için tek noktadan)
                #     + DSB yanlış kullanımını düzelt.
                for _d in self.flat_data.values():
                    _orig = _d.get("content", "")
                    if not _orig:
                        continue
                    _is_test = _d.get("type") in self.TEST_TYPES
                    _new = text_cleanup.temizle(_orig, test=_is_test)
                    if "DSB" in (_new or "").upper():
                        _new = text_cleanup.dsb_temizle(_new)
                    if _new and _new != _orig:
                        _d["content"] = _new
                        if _orig in (self.last_generated_output or ""):
                            self.last_generated_output = self.last_generated_output.replace(_orig, _new)
                # (2) DSB ZİNCİRİ: üst gereksinimde DSB varsa, ona bağlı testte uydurma sayı → DSB
                for _d in self.flat_data.values():
                    if _d.get("type") in self.TEST_TYPES:
                        _p = self.flat_data.get(_d.get("bound_to"))
                        _orig = _d.get("content", "")
                        if _p and "DSB" in (_p.get("content", "")).upper() and "DSB" not in _orig.upper():
                            _new = text_cleanup.sayilari_dsb_yap(_orig)
                            if _new and _new != _orig:
                                _d["content"] = _new
                                if _orig in (self.last_generated_output or ""):
                                    self.last_generated_output = self.last_generated_output.replace(_orig, _new)
            except Exception as _e:
                self.update_status_text(f"Temizlik/DSB düzeltme uyarısı: {_e}", is_error=True)

            total_end_time = time.time()
            total_duration = total_end_time - total_start_time
            minutes = int(total_duration // 60)
            seconds = int(total_duration % 60)
            
            if not self.last_generated_output.strip():
                self.update_status_text("Hiçbir doküman üretilemedi.", is_error=True)
                self.master.after(
                    0,
                    lambda: self._notify_architecture_generation_failed(
                        "Hiçbir doküman üretilemedi."
                    ),
                )
            else:
                final_msg = f"Toplam Geçen Süre: {minutes} dakika {seconds} saniye\n--- YAPAY ZEKA ÜRETİMİ TAMAMLANDI ---\n"
                self.update_status_text(final_msg, is_complete=True)
                self.raw_output_cache = self.last_generated_output  
                self._start_traceability_build(proje_ismi)

        except Exception as e:
            self.update_status_text(f"KRİTİK HATA: {e}", is_error=True)
            self.master.after(
                0,
                lambda detail=str(e): self._notify_architecture_generation_failed(detail),
            )
            traceback.print_exc()
        finally:
            self.master.after(0, lambda: self._reset_buttons_state())

    def _reset_buttons_state(self):
        self.create_docs_button.config(state=tk.NORMAL, text=self._t("Dokümanları Üret", "Generate Documents"), style="primary.TButton")
        self.download_docs_button.config(state=tk.NORMAL)
        self.reset_button.config(state=tk.NORMAL)

if __name__ == "__main__":
    try:
        from ctypes import windll
        windll.shcore.SetProcessDpiAwareness(1)
    except Exception:
        pass

    prepare_process_identity()
    root = ttk.Window()
    app = TIDGeneratorApp(root)
    root.mainloop()
