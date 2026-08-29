# -*- coding: utf-8 -*-
"""Üretilen V-Model verilerinden kalıcı izlenebilirlik haritası oluşturur.

Modül, uygulamanın yapılandırılmış ``flat_data`` kayıtlarını birincil kaynak
olarak kullanır. Yapılandırılmış veri yoksa mevcut bağımlılıklarla PDF, DOCX,
XLSX ve TXT dosyalarından gerçek kimlikleri geri okumayı dener. Hiçbir belgeyi
veya Etki Analizi hesabını değiştirmez.
"""

from __future__ import annotations

from copy import deepcopy
from datetime import datetime, timezone
from difflib import SequenceMatcher
import hashlib
import json
import os
from pathlib import Path
import re
import tempfile
import unicodedata
from typing import Any, Callable, Iterable, Mapping, Sequence


SCHEMA_VERSION = "1.0"
# Faz 7'de bu dosya proje kokunden core/ alt paketine tasindi; __file__.parent
# artik core/'u gosterdigi icin proje kokune cikmak icin bir ust dizine
# (.parent.parent) cikiyoruz - davranis (outputs/traceability'yi proje
# kokunde bulmak) tasimadan onceki haliyle ayni kalsin diye.
DEFAULT_OUTPUT_ROOT = Path(__file__).resolve().parent.parent / "outputs" / "traceability"

CONFIDENCE_EXACT = "Kesin"
CONFIDENCE_SUGGESTED = "Önerilen bağlantı"
CONFIDENCE_INFERRED = "Çıkarım"

RELATION_LABELS_TR = {
    "derives_from": "türetilmiştir",
    "satisfies": "karşılar",
    "allocated_to": "tahsis edilmiştir",
    "depends_on": "bağımlıdır",
    "interfaces_with": "arayüzü vardır",
    "implemented_by": "uygulanır",
    "verified_by": "doğrulanır",
    "validated_by": "geçerlenir",
    "conflicts_with": "çakışır",
    "documented_in": "belgelenmiştir",
}

# Bunlar Arayüz.py içindeki gerçek VMODEL_SECTIONS ve üretici çıktı türleridir.
DOCUMENT_TYPE_DEFINITIONS: dict[str, dict[str, str]] = {
    "TID": {
        "document_title": "Kullanıcı Gereksinimi (User Requirement)",
        "node_type": "Müşteri/paydaş gereksinimi",
        "v_model_level": "Müşteri gereksinimi",
        "leg": "sol",
        "role": "requirement",
    },
    "SGD": {
        "document_title": "Sistem Gereksinimi (System Requirements)",
        "node_type": "Sistem gereksinimi",
        "v_model_level": "Sistem gereksinimi",
        "leg": "sol",
        "role": "requirement",
    },
    "STT": {
        "document_title": "Alt Sistem Gereksinimleri (Subsystem Requirements)",
        "node_type": "Alt sistem gereksinimi",
        "v_model_level": "Alt sistem gereksinimi",
        "leg": "sol",
        "role": "requirement",
    },
    "KMTD": {
        "document_title": "Kabul Testi (Acceptance Test)",
        "node_type": "Müşteri kabul/geçerleme testi",
        "v_model_level": "Müşteri kabulü ve geçerleme",
        "leg": "sağ",
        "role": "test",
    },
    "SITET": {
        "document_title": "Sistem Testi (System Test)",
        "node_type": "Sistem doğrulama testi",
        "v_model_level": "Sistem doğrulaması",
        "leg": "sağ",
        "role": "test",
    },
    "AST": {
        "document_title": "Alt Sistem Testi (Subsystem Test)",
        "node_type": "Entegrasyon testi",
        "v_model_level": "Alt sistem entegrasyon testi",
        "leg": "sağ",
        "role": "test",
    },
}

_TYPE_ALIASES = {
    "TID": "TID",
    "UR": "TID",
    "SGD": "SGD",
    "SR": "SGD",
    "STT": "STT",
    "SSR": "STT",
    "KMTD": "KMTD",
    "AT": "KMTD",
    "KABUL TESTI": "KMTD",
    "KABUL MUAYENE": "KMTD",
    "SITET": "SITET",
    "SISTEM TESTI": "SITET",
    "AST": "AST",
    "SST": "AST",
    "ALT SISTEM TESTI": "AST",
}

_ID_PATTERN = re.compile(
    r"\b(?:SITET|SSR|SST|UR|SR|AT)-?\s*\d{1,8}\b",
    re.IGNORECASE,
)
_GENERIC_ID_PATTERN = re.compile(
    r"\b[A-ZÇĞİÖŞÜ]{2,12}(?:-[A-ZÇĞİÖŞÜ]{2,12})?-?\s*\d{1,8}\b",
    re.IGNORECASE,
)
_NO_PARENT_MARKERS = {
    "", "yok", "none", "null", "-", "genel", "tid-genel", "sgd", "asg",
}
_STOPWORDS = {
    "ve", "veya", "ile", "bir", "bu", "şu", "icin", "için", "olarak",
    "sistem", "gereksinim", "gereksinimi", "test", "testi", "kontrol",
    "edilmelidir", "olmalidir", "olmalıdır", "saglanmalidir", "sağlanmalıdır",
    "dogrulanmalidir", "doğrulanmalıdır", "ilgili", "alt", "ust", "üst",
}
_PARAMETER_RE = re.compile(
    r"(?P<raw>(?:<=|>=|<|>|±)?\s*[+-]?\d+(?:[.,]\d+)?\s*"
    r"(?P<unit>%|ms|sn|saniye|dakika|dk|saat|°c|c|v|a|ma|w|kw|hz|khz|mhz|ghz|"
    r"bit/s|kbit/s|mbit/s|gbit/s|kb|mb|gb|tb|mm|cm|km|kg|db|adet|rpm))\b",
    re.IGNORECASE,
)


class TraceabilityError(ValueError):
    """İzlenebilirlik girdisi veya kalıcı kayıt hatası."""


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _ascii_fold(value: Any) -> str:
    text = unicodedata.normalize("NFKD", _clean(value))
    return "".join(char for char in text if not unicodedata.combining(char))


def normalize_identifier(value: Any) -> str:
    """Farklı yazılmış aynı kimliği eşleştirmek için karşılaştırma anahtarı üretir."""
    text = _ascii_fold(value).upper()
    text = text.replace("–", "-").replace("—", "-").replace("_", "-")
    text = re.sub(r"\s+", "", text)
    match = re.fullmatch(r"([A-Z]{2,12}(?:-[A-Z]{2,12})?)-?(\d{1,8})", text)
    if not match:
        return text
    return f"{match.group(1)}-{match.group(2)}"


def extract_identifiers(text: Any) -> list[str]:
    """Metindeki kimlikleri sırasını koruyarak ve kopyaları ayıklayarak döndürür."""
    raw_text = _clean(text)
    matches = _GENERIC_ID_PATTERN.findall(raw_text)
    seen: set[str] = set()
    identifiers: list[str] = []
    for raw in matches:
        normalized = normalize_identifier(raw)
        if normalized and normalized not in seen:
            identifiers.append(normalized)
            seen.add(normalized)
    return identifiers


def _type_from_identifier(identifier: str) -> str | None:
    prefix = normalize_identifier(identifier).split("-", 1)[0]
    return {
        "UR": "TID",
        "SR": "SGD",
        "SSR": "STT",
        "AT": "KMTD",
        "SITET": "SITET",
        "SST": "AST",
    }.get(prefix)


def _normalize_type(value: Any, identifier: str = "") -> str | None:
    raw = _ascii_fold(value).upper().replace("_", " ").replace("-", " ")
    raw = re.sub(r"\s+", " ", raw).strip()
    direct = _TYPE_ALIASES.get(raw)
    if direct:
        return direct
    return _type_from_identifier(identifier)


def extract_technical_parameters(text: Any) -> list[dict[str, Any]]:
    """Yalnızca metinde açıkça bulunan sayı-birim çiftlerini çıkarır."""
    parameters: list[dict[str, Any]] = []
    for match in _PARAMETER_RE.finditer(_clean(text)):
        raw = _clean(match.group("raw"))
        unit = match.group("unit")
        number_match = re.search(r"[+-]?\d+(?:[.,]\d+)?", raw)
        if not number_match:
            continue
        value = float(number_match.group(0).replace(",", "."))
        parameters.append({
            "raw": raw,
            "value": value,
            "unit": unit,
        })
    return parameters


def _tokens(text: Any) -> set[str]:
    normalized = re.sub(r"[^a-z0-9çğıöşü]+", " ", _ascii_fold(text).casefold())
    return {
        token for token in normalized.split()
        if len(token) > 2 and token not in _STOPWORDS and not token.isdigit()
    }


def semantic_similarity(left: Any, right: Any) -> float:
    """Harici model olmadan, denetlenebilir metinsel benzerlik skoru üretir."""
    left_text = _ascii_fold(left).casefold()
    right_text = _ascii_fold(right).casefold()
    if not left_text or not right_text:
        return 0.0
    left_tokens = _tokens(left_text)
    right_tokens = _tokens(right_text)
    overlap = (
        len(left_tokens & right_tokens) / min(len(left_tokens), len(right_tokens))
        if left_tokens and right_tokens
        else 0.0
    )
    sequence = SequenceMatcher(None, left_text, right_text).ratio()
    return round((overlap * 0.68) + (sequence * 0.32), 4)


def check_lm_studio_status(
    base_url: str | None = None,
    timeout: float = 1.5,
    request_get: Callable[..., Any] | None = None,
) -> dict[str, Any]:
    """LM Studio durumunu kısa zaman aşımıyla kontrol eder; taramayı asla durdurmaz."""
    if base_url is None:
        try:
            from config import LMSTUDIO_BASE_URL

            base_url = LMSTUDIO_BASE_URL
        except Exception:
            base_url = "http://localhost:1234/v1"
    try:
        if request_get is None:
            import requests

            request_get = requests.get
        response = request_get(f"{base_url.rstrip('/')}/models", timeout=timeout)
        if getattr(response, "status_code", None) == 200:
            return {
                "available": True,
                "message": "LM Studio bağlantısı hazır.",
                "endpoint": base_url,
            }
        return {
            "available": False,
            "message": (
                "LM Studio yanıt vermedi veya model sunucusu hazır değil; "
                "kesin ve metinsel bağlantılarla devam edildi."
            ),
            "endpoint": base_url,
        }
    except Exception as error:
        return {
            "available": False,
            "message": (
                "LM Studio kapalı veya erişilemiyor; kesin ve metinsel "
                f"bağlantılarla devam edildi ({error})."
            ),
            "endpoint": base_url,
        }


def _record_from_line(
    line: str,
    source: str,
    location: str,
) -> dict[str, Any] | None:
    matches = _ID_PATTERN.findall(line)
    if not matches:
        return None
    raw_id = matches[0]
    identifier = normalize_identifier(raw_id)
    item_type = _type_from_identifier(identifier)
    if not item_type:
        return None
    parts = [part.strip() for part in re.split(r"\s*\|\s*|\t+", line) if part.strip()]
    content = ""
    bound_to = ""
    for part in parts:
        if normalize_identifier(part) == identifier:
            continue
        if not content:
            content = part
        elif not bound_to and extract_identifiers(part):
            bound_to = extract_identifiers(part)[0]
    if not content:
        content = _clean(line.replace(raw_id, "", 1).strip(" |-:"))
    return {
        "ID": raw_id,
        "type": item_type,
        "content": content,
        "bound_to": bound_to,
        "source_document": source,
        "source_section": location,
    }


def _read_pdf(path: Path) -> list[tuple[str, str]]:
    try:
        import pdf_extraction

        extracted = pdf_extraction.extract_pdf_to_txt(str(path))
        if extracted is None:
            raise TraceabilityError("PDF içeriği çıkarılamadı veya dosya geçersiz.")
        text = extracted
    except Exception as error:
        raise TraceabilityError(f"PDF okunamadı: {error}") from error
    current_page = ""
    lines: list[tuple[str, str]] = []
    for line in text.splitlines():
        page_match = re.match(r"---\s*(?:Page|Sayfa)\s*(\d+)\s*---", line, re.I)
        if page_match:
            current_page = f"Sayfa {page_match.group(1)}"
            continue
        lines.append((line, current_page))
    return lines


def _read_docx(path: Path) -> list[tuple[str, str]]:
    try:
        from docx import Document

        document = Document(str(path))
    except Exception as error:
        raise TraceabilityError(f"Word belgesi okunamadı: {error}") from error
    rows: list[tuple[str, str]] = [
        (paragraph.text, "Paragraf") for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table_index, table in enumerate(document.tables, start=1):
        for row_index, row in enumerate(table.rows, start=1):
            rows.append((
                " | ".join(cell.text.strip() for cell in row.cells),
                f"Tablo {table_index}, Satır {row_index}",
            ))
    return rows


def _read_xlsx(path: Path) -> list[tuple[str, str]]:
    try:
        from openpyxl import load_workbook

        workbook = load_workbook(str(path), read_only=True, data_only=True)
    except Exception as error:
        raise TraceabilityError(f"Excel belgesi okunamadı: {error}") from error
    rows: list[tuple[str, str]] = []
    try:
        for sheet in workbook.worksheets:
            for row_index, values in enumerate(sheet.iter_rows(values_only=True), start=1):
                text = " | ".join(_clean(value) for value in values if value is not None)
                if text:
                    rows.append((text, f"{sheet.title}!{row_index}"))
    finally:
        workbook.close()
    return rows


def _read_text(path: Path) -> list[tuple[str, str]]:
    try:
        text = path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        text = path.read_text(encoding="utf-8-sig")
    except Exception as error:
        raise TraceabilityError(f"Metin belgesi okunamadı: {error}") from error
    return [(line, f"Satır {index}") for index, line in enumerate(text.splitlines(), start=1)]


def read_document_lines(
    path: str | os.PathLike[str],
) -> list[tuple[str, str]]:
    """Desteklenen bir belgeyi ``(metin, konum)`` satırları olarak okur.

    Donanım kartları gibi ikincil çıkarıcıların PDF/Word/Excel okuma kodunu
    kopyalamaması için izlenebilirlik okuyucularının salt-okunur ortak
    girişidir. Kaynak dosyada hiçbir değişiklik yapmaz.
    """
    document_path = Path(path)
    readers = {
        ".pdf": _read_pdf,
        ".docx": _read_docx,
        ".xlsx": _read_xlsx,
        ".xlsm": _read_xlsx,
        ".txt": _read_text,
    }
    reader = readers.get(document_path.suffix.lower())
    if reader is None:
        raise TraceabilityError(
            f"Desteklenmeyen belge biçimi: {document_path.suffix or 'uzantısız dosya'}"
        )
    if not document_path.exists():
        raise TraceabilityError(f"Belge bulunamadı: {document_path}")
    return reader(document_path)


def read_document_records(
    paths: Iterable[str | os.PathLike[str]],
    status_callback: Callable[[str], None] | None = None,
) -> tuple[dict[str, dict[str, Any]], list[dict[str, Any]]]:
    """Yapılandırılmış veri yoksa desteklenen dosyalardan maddeleri geri okur."""
    records: dict[str, dict[str, Any]] = {}
    documents: list[dict[str, Any]] = []
    readers = {
        ".pdf": _read_pdf,
        ".docx": _read_docx,
        ".xlsx": _read_xlsx,
        ".xlsm": _read_xlsx,
        ".txt": _read_text,
    }
    for raw_path in paths:
        path = Path(raw_path)
        document_info = {
            "name": path.name,
            "path": str(path.resolve()) if path.exists() else str(path),
            "format": path.suffix.lower().lstrip("."),
            "source_kind": "physical_file",
            "status": "ok",
            "record_count": 0,
        }
        reader = readers.get(path.suffix.lower())
        if reader is None:
            document_info["status"] = "unsupported"
            document_info["error"] = "Desteklenmeyen belge biçimi."
            documents.append(document_info)
            continue
        try:
            if status_callback:
                status_callback(f"İzlenebilirlik için okunuyor: {path.name}")
            rows = reader(path)
            for line, location in rows:
                record = _record_from_line(line, path.name, location)
                if not record:
                    continue
                key = normalize_identifier(record["ID"])
                if key not in records:
                    records[key] = record
                    document_info["record_count"] += 1
        except Exception as error:
            document_info["status"] = "error"
            document_info["error"] = str(error)
        documents.append(document_info)
    return records, documents


def _project_identity(project_name: str) -> tuple[str, str]:
    cleaned_name = _clean(project_name)
    if not cleaned_name:
        raise TraceabilityError("İzlenebilirlik için proje adı gerekli.")
    slug = re.sub(r"[^a-z0-9]+", "-", _ascii_fold(cleaned_name).casefold()).strip("-")
    slug = slug[:64] or "proje"
    digest = hashlib.sha256(cleaned_name.casefold().encode("utf-8")).hexdigest()[:8]
    return f"{slug}-{digest}", slug


def project_identity(project_name: str) -> tuple[str, str]:
    """Diğer proje-bazlı veri katmanları için ortak ve kararlı proje kimliği."""
    return _project_identity(project_name)


def _edge_id(source_id: str, target_id: str, relationship_type: str) -> str:
    payload = f"{source_id}|{relationship_type}|{target_id}".encode("utf-8")
    return "EDGE-" + hashlib.sha256(payload).hexdigest()[:16].upper()


def _add_edge(
    edges: list[dict[str, Any]],
    seen: set[tuple[str, str, str]],
    *,
    source_id: str,
    target_id: str,
    relationship_type: str,
    confidence_level: str,
    confidence: float,
    evidence_text: str,
    source_document: str,
    derivation_method: str,
) -> None:
    key = (source_id, target_id, relationship_type)
    if key in seen:
        return
    seen.add(key)
    edges.append({
        "id": _edge_id(*key),
        "source_id": source_id,
        "target_id": target_id,
        "relationship_type": relationship_type,
        "relationship_label_tr": RELATION_LABELS_TR[relationship_type],
        "confidence_level": confidence_level,
        "confidence": round(float(confidence), 4),
        "evidence_text": evidence_text,
        "source_document": source_document,
        "derivation_method": derivation_method,
    })


def _node_from_record(
    identifier: str,
    raw: Mapping[str, Any],
    item_type: str,
) -> dict[str, Any]:
    definition = DOCUMENT_TYPE_DEFINITIONS[item_type]
    description = _clean(raw.get("content") or raw.get("description"))
    original_id = _clean(raw.get("ID") or identifier)
    return {
        "id": original_id,
        "canonical_id": normalize_identifier(original_id),
        "aliases": [original_id],
        "node_type": definition["node_type"],
        "document_type": item_type,
        "title": _clean(raw.get("title")) or original_id,
        "description": description,
        "v_model_level": definition["v_model_level"],
        "v_model_leg": definition["leg"],
        "version": raw.get("version"),
        "status": _clean(raw.get("status")) or "Üretildi",
        "source_document": _clean(raw.get("source_document")) or definition["document_title"],
        "source_section": _clean(raw.get("source_section")) or None,
        "evidence_text": description,
        "confidence_level": CONFIDENCE_EXACT,
        "confidence": 1.0,
        "technical_parameters": extract_technical_parameters(description),
    }


def _hardware_nodes(
    hardware_data: Mapping[str, Mapping[str, Any]] | None,
) -> tuple[list[dict[str, Any]], list[dict[str, str]]]:
    nodes: list[dict[str, Any]] = []
    links: list[dict[str, str]] = []
    for fallback_id, raw in (hardware_data or {}).items():
        if not isinstance(raw, Mapping):
            continue
        item_id = _clean(raw.get("ID") or fallback_id)
        if not item_id:
            continue
        description = _clean(raw.get("description") or raw.get("name"))
        specifications = raw.get("specifications") if isinstance(raw.get("specifications"), Mapping) else {}
        parameters = [
            {"name": _clean(name), "value": _clean(value)}
            for name, value in specifications.items()
            if _clean(name)
        ]
        nodes.append({
            "id": item_id,
            "canonical_id": normalize_identifier(item_id),
            "aliases": [item_id],
            "node_type": "Parça/bileşen",
            "document_type": "HARDWARE",
            "title": description or item_id,
            "description": description,
            "v_model_level": "Parça veya yazılım uygulaması",
            "v_model_leg": "sol",
            "version": raw.get("version"),
            "status": _clean(raw.get("status")) or None,
            "source_document": "Akıllı Donanım Listesi",
            "source_section": None,
            "evidence_text": _clean(raw.get("source_excerpt") or raw.get("rationale") or description),
            "confidence_level": CONFIDENCE_EXACT,
            "confidence": raw.get("confidence") if raw.get("confidence") is not None else 1.0,
            "technical_parameters": parameters,
            "risk": _clean(raw.get("risk")) or None,
        })
        linked = raw.get("linked_requirements") or raw.get("requirement_ids") or []
        if isinstance(linked, str):
            linked = re.split(r"[,;\s]+", linked)
        for requirement_id in linked:
            normalized = normalize_identifier(requirement_id)
            if normalized:
                links.append({"requirement_id": normalized, "hardware_id": item_id})
    return nodes, links


def _expected_parent_type(item_type: str) -> str | None:
    return {
        "SGD": "TID",
        "STT": "SGD",
        "KMTD": "TID",
        "SITET": "SGD",
        "AST": "STT",
    }.get(item_type)


def _semantic_relation_type(item_type: str) -> str:
    if item_type == "KMTD":
        return "validated_by"
    if item_type in {"SITET", "AST"}:
        return "verified_by"
    return "derives_from"


def _relation_direction(item_type: str, item_id: str, parent_id: str) -> tuple[str, str]:
    if item_type in {"KMTD", "SITET", "AST"}:
        return parent_id, item_id
    return item_id, parent_id


def _atomic_write_json(path: Path, payload: Mapping[str, Any]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_name = ""
    try:
        with tempfile.NamedTemporaryFile(
            mode="w",
            encoding="utf-8",
            dir=path.parent,
            prefix=f".{path.name}.",
            suffix=".tmp",
            delete=False,
        ) as temporary:
            json.dump(payload, temporary, ensure_ascii=False, indent=2)
            temporary.write("\n")
            temporary.flush()
            os.fsync(temporary.fileno())
            temporary_name = temporary.name
        os.replace(temporary_name, path)
    finally:
        if temporary_name and os.path.exists(temporary_name):
            os.unlink(temporary_name)


def atomic_write_json(
    path: str | os.PathLike[str], payload: Mapping[str, Any]
) -> None:
    """JSON verisini hedef dosyaya atomik olarak yazar."""
    _atomic_write_json(Path(path), payload)


def persist_traceability_report(
    report: dict[str, Any],
    output_root: str | os.PathLike[str] | None = None,
) -> dict[str, Any]:
    """Raporu atomik olarak hem sürüm dosyasına hem güncel dosyaya yazar."""
    root = Path(output_root) if output_root else DEFAULT_OUTPUT_ROOT
    project_dir = root / report["project_id"]
    latest_path = project_dir / "traceability.json"
    revision = 1
    if latest_path.exists():
        try:
            previous = json.loads(latest_path.read_text(encoding="utf-8"))
            revision = int(previous.get("revision", 0)) + 1
        except Exception:
            revision = 1
    report["revision"] = revision
    version_path = project_dir / f"traceability.v{revision:04d}.json"
    _atomic_write_json(version_path, report)
    _atomic_write_json(latest_path, report)
    report["storage_path"] = str(latest_path.resolve())
    report["version_path"] = str(version_path.resolve())
    return report


def load_project_traceability(
    project_name: str,
    output_root: str | os.PathLike[str] | None = None,
) -> dict[str, Any] | None:
    """Bir projenin son atomik izlenebilirlik kaydını salt okunur yükler."""
    project_id, _ = _project_identity(project_name)
    root = Path(output_root) if output_root else DEFAULT_OUTPUT_ROOT
    latest_path = root / project_id / "traceability.json"
    if not latest_path.exists():
        return None
    try:
        report = json.loads(latest_path.read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        raise TraceabilityError(
            f"Projenin izlenebilirlik kaydı okunamadı: {error}"
        ) from error
    if not isinstance(report, dict) or not isinstance(report.get("nodes"), list) or not isinstance(report.get("edges"), list):
        raise TraceabilityError(
            "Projenin izlenebilirlik kaydı geçerli nodes/edges listeleri içermiyor."
        )
    report["storage_path"] = str(latest_path.resolve())
    revision = int(report.get("revision", 0) or 0)
    if revision > 0:
        report["version_path"] = str(
            (latest_path.parent / f"traceability.v{revision:04d}.json").resolve()
        )
    return report


def build_traceability_map(
    project_name: str,
    flat_data: Mapping[str, Mapping[str, Any]] | None = None,
    *,
    hardware_data: Mapping[str, Mapping[str, Any]] | None = None,
    source_paths: Sequence[str | os.PathLike[str]] | None = None,
    document_sections: Sequence[Sequence[str]] | None = None,
    output_root: str | os.PathLike[str] | None = None,
    persist: bool = True,
    check_lm_studio: bool = True,
    request_get: Callable[..., Any] | None = None,
    now: datetime | None = None,
    status_callback: Callable[[str], None] | None = None,
) -> dict[str, Any]:
    """V-Model kayıtlarından izlenebilirlik haritasını oluşturur ve saklar."""
    project_id, _ = _project_identity(project_name)
    generated_at = (now or datetime.now(timezone.utc)).astimezone().isoformat(timespec="seconds")
    records: dict[str, Mapping[str, Any]] = {
        str(key): deepcopy(value)
        for key, value in (flat_data or {}).items()
        if isinstance(value, Mapping)
    }
    physical_documents: list[dict[str, Any]] = []
    if not records and source_paths:
        if status_callback:
            status_callback("Yapılandırılmış veri bulunamadı; belgeler geri okunuyor...")
        records, physical_documents = read_document_records(source_paths, status_callback)

    section_titles: dict[str, str] = {}
    for section in document_sections or ():
        if len(section) >= 2:
            section_titles[_clean(section[0]).upper()] = _clean(section[1])

    nodes: list[dict[str, Any]] = []
    edges: list[dict[str, Any]] = []
    conflicts: list[dict[str, Any]] = []
    missing_information: list[dict[str, Any]] = []
    source_documents: list[dict[str, Any]] = list(physical_documents)
    seen_edges: set[tuple[str, str, str]] = set()
    nodes_by_canonical: dict[str, dict[str, Any]] = {}
    raw_by_canonical: dict[str, Mapping[str, Any]] = {}
    ids_by_type: dict[str, list[str]] = {key: [] for key in DOCUMENT_TYPE_DEFINITIONS}

    if status_callback:
        status_callback("Yapılandırılmış V-Model kayıtları taranıyor...")

    for fallback_id, raw in records.items():
        raw_id = _clean(raw.get("ID") or fallback_id)
        canonical = normalize_identifier(raw_id)
        item_type = _normalize_type(raw.get("type"), raw_id)
        if not raw_id or not canonical:
            missing_information.append({
                "type": "missing_identifier",
                "item_id": raw_id or None,
                "message": "Kimliği olmayan kayıt düğüme dönüştürülmedi.",
            })
            continue
        if item_type not in DOCUMENT_TYPE_DEFINITIONS:
            missing_information.append({
                "type": "unsupported_document_type",
                "item_id": raw_id,
                "message": f"Projede eşlemesi olmayan belge türü: {_clean(raw.get('type')) or 'boş'}.",
            })
            continue
        node = _node_from_record(raw_id, raw, item_type)
        existing = nodes_by_canonical.get(canonical)
        if existing:
            if raw_id not in existing["aliases"]:
                existing["aliases"].append(raw_id)
            existing_description = _clean(existing.get("description")).casefold()
            new_description = _clean(node.get("description")).casefold()
            if new_description and existing_description and new_description != existing_description:
                conflicts.append({
                    "type": "duplicate_identifier_conflict",
                    "item_id": existing["id"],
                    "aliases": sorted(set(existing["aliases"] + [raw_id])),
                    "message": "Aynı kimlik farklı açıklamalarla bulundu.",
                    "existing_evidence": existing["description"],
                    "conflicting_evidence": node["description"],
                })
            if existing.get("document_type") != item_type:
                conflicts.append({
                    "type": "identifier_type_conflict",
                    "item_id": existing["id"],
                    "message": "Aynı kimlik farklı belge türlerinde bulundu.",
                    "types": [existing.get("document_type"), item_type],
                })
            continue
        nodes_by_canonical[canonical] = node
        raw_by_canonical[canonical] = raw
        ids_by_type[item_type].append(canonical)

    nodes.extend(nodes_by_canonical.values())

    # Üretilen her gerçek belge türü için bir teknik belge düğümü oluştur.
    for item_type, identifiers in ids_by_type.items():
        if not identifiers:
            continue
        definition = DOCUMENT_TYPE_DEFINITIONS[item_type]
        document_title = section_titles.get(item_type) or definition["document_title"]
        document_id = f"DOC-{item_type}"
        nodes.append({
            "id": document_id,
            "canonical_id": document_id,
            "aliases": [document_id],
            "node_type": "Teknik belge",
            "document_type": item_type,
            "title": document_title,
            "description": f"{len(identifiers)} yapılandırılmış madde içerir.",
            "v_model_level": definition["v_model_level"],
            "v_model_leg": definition["leg"],
            "version": None,
            "status": "Üretildi",
            "source_document": document_title,
            "source_section": None,
            "evidence_text": None,
            "confidence_level": CONFIDENCE_EXACT,
            "confidence": 1.0,
            "technical_parameters": [],
        })
        source_documents.append({
            "document_id": document_id,
            "document_type": item_type,
            "name": document_title,
            "format": "structured-python",
            "source_kind": "generated_structured_data",
            "v_model_leg": definition["leg"],
            "item_count": len(identifiers),
            "status": "ok",
        })
        for canonical in identifiers:
            item_node = nodes_by_canonical[canonical]
            _add_edge(
                edges,
                seen_edges,
                source_id=item_node["id"],
                target_id=document_id,
                relationship_type="documented_in",
                confidence_level=CONFIDENCE_EXACT,
                confidence=1.0,
                evidence_text=f"{item_node['id']} yapılandırılmış {item_type} kaydından alındı.",
                source_document=document_title,
                derivation_method="structured_document_membership",
            )

    unresolved: list[tuple[str, str, Mapping[str, Any]]] = []
    for canonical, node in nodes_by_canonical.items():
        item_type = node["document_type"]
        expected_parent_type = _expected_parent_type(item_type)
        if not expected_parent_type:
            continue
        raw = raw_by_canonical[canonical]
        bound_text = _clean(raw.get("bound_to") or raw.get("bound") or raw.get("parent_id"))
        normalized_bound = normalize_identifier(bound_text)
        bound_candidates = extract_identifiers(bound_text)
        if normalized_bound and normalized_bound not in bound_candidates:
            bound_candidates.insert(0, normalized_bound)
        bound_candidates = [
            item for item in bound_candidates
            if item.casefold() not in _NO_PARENT_MARKERS
        ]
        matched_parent = next(
            (candidate for candidate in bound_candidates if candidate in nodes_by_canonical),
            None,
        )
        if matched_parent:
            parent_node = nodes_by_canonical[matched_parent]
            relation_type = _semantic_relation_type(item_type)
            source_id, target_id = _relation_direction(item_type, node["id"], parent_node["id"])
            _add_edge(
                edges,
                seen_edges,
                source_id=source_id,
                target_id=target_id,
                relationship_type=relation_type,
                confidence_level=CONFIDENCE_EXACT,
                confidence=1.0,
                evidence_text=f"{node['id']} kaydındaki açık bound_to bağlantısı: {bound_text}",
                source_document=node["source_document"],
                derivation_method="structured_bound_to",
            )
        else:
            if bound_text and bound_text.casefold() not in _NO_PARENT_MARKERS:
                missing_information.append({
                    "type": "missing_bound_target",
                    "item_id": node["id"],
                    "message": f"Bağlantı hedefi bulunamadı: {bound_text}.",
                })
            unresolved.append((canonical, expected_parent_type, raw))

    # Açık bağlantı yoksa yalnızca güçlü metin benzerliğini öneri olarak kaydet.
    for canonical, parent_type, raw in unresolved:
        node = nodes_by_canonical[canonical]
        candidates = ids_by_type.get(parent_type, [])
        scored = sorted(
            (
                (semantic_similarity(node["description"], nodes_by_canonical[parent]["description"]), parent)
                for parent in candidates
            ),
            reverse=True,
        )
        if not scored or scored[0][0] < 0.48:
            continue
        score, parent_canonical = scored[0]
        parent_node = nodes_by_canonical[parent_canonical]
        relation_type = _semantic_relation_type(node["document_type"])
        source_id, target_id = _relation_direction(
            node["document_type"], node["id"], parent_node["id"]
        )
        _add_edge(
            edges,
            seen_edges,
            source_id=source_id,
            target_id=target_id,
            relationship_type=relation_type,
            confidence_level=CONFIDENCE_SUGGESTED,
            confidence=score,
            evidence_text=(
                f"Metin benzerliği: {node['id']} ↔ {parent_node['id']} "
                f"(skor {score:.2f})."
            ),
            source_document=node["source_document"],
            derivation_method="text_similarity",
        )

    hardware_nodes, hardware_links = _hardware_nodes(hardware_data)
    if hardware_nodes:
        hardware_document_id = "DOC-HARDWARE"
        nodes.extend(hardware_nodes)
        nodes.append({
            "id": hardware_document_id,
            "canonical_id": hardware_document_id,
            "aliases": [hardware_document_id],
            "node_type": "Teknik belge",
            "document_type": "HARDWARE",
            "title": "Akıllı Donanım Listesi",
            "description": f"{len(hardware_nodes)} parça/bileşen kaydı içerir.",
            "v_model_level": "Parça veya yazılım uygulaması",
            "v_model_leg": "sol",
            "version": None,
            "status": "Üretildi",
            "source_document": "Akıllı Donanım Listesi",
            "source_section": None,
            "evidence_text": None,
            "confidence_level": CONFIDENCE_EXACT,
            "confidence": 1.0,
            "technical_parameters": [],
        })
        source_documents.append({
            "document_id": hardware_document_id,
            "document_type": "HARDWARE",
            "name": "Akıllı Donanım Listesi",
            "format": "structured-python",
            "source_kind": "generated_structured_data",
            "v_model_leg": "sol",
            "item_count": len(hardware_nodes),
            "status": "ok",
        })
        for hardware_node in hardware_nodes:
            _add_edge(
                edges,
                seen_edges,
                source_id=hardware_node["id"],
                target_id=hardware_document_id,
                relationship_type="documented_in",
                confidence_level=CONFIDENCE_EXACT,
                confidence=1.0,
                evidence_text="Parça kaydı Akıllı Donanım Listesinde bulunuyor.",
                source_document="Akıllı Donanım Listesi",
                derivation_method="structured_document_membership",
            )
        for link in hardware_links:
            requirement = nodes_by_canonical.get(link["requirement_id"])
            if requirement:
                _add_edge(
                    edges,
                    seen_edges,
                    source_id=requirement["id"],
                    target_id=link["hardware_id"],
                    relationship_type="allocated_to",
                    confidence_level=CONFIDENCE_EXACT,
                    confidence=1.0,
                    evidence_text="Donanım kaydındaki açık linked_requirements bağlantısı.",
                    source_document="Akıllı Donanım Listesi",
                    derivation_method="structured_hardware_link",
                )
            else:
                missing_information.append({
                    "type": "missing_hardware_requirement",
                    "item_id": link["hardware_id"],
                    "message": f"Donanım bağlantısındaki gereksinim bulunamadı: {link['requirement_id']}.",
                })

    node_ids = {node["id"] for node in nodes}
    substantive_edges = [
        edge for edge in edges if edge["relationship_type"] != "documented_in"
    ]
    requirement_nodes = [
        node for node in nodes
        if node.get("node_type") != "Teknik belge"
        and node.get("document_type") in {"TID", "SGD", "STT"}
    ]
    connected_ids = {
        endpoint
        for edge in substantive_edges
        for endpoint in (edge["source_id"], edge["target_id"])
        if endpoint in node_ids
    }
    verified_ids = {
        edge["source_id"] for edge in substantive_edges
        if edge["relationship_type"] in {"verified_by", "validated_by"}
    }
    unlinked_requirements = sorted(
        node["id"] for node in requirement_nodes if node["id"] not in connected_ids
    )
    unverified_requirements = sorted(
        node["id"] for node in requirement_nodes if node["id"] not in verified_ids
    )

    present_types = {node.get("node_type") for node in nodes}
    absent_structured_types = [
        node_type for node_type in (
            "Fonksiyon", "Yazılım birimi", "Mekanik arayüz", "Elektriksel arayüz",
            "Yazılımsal arayüz", "Tasarım kararı", "Risk", "Doğrulama kriteri",
            "Birim testi",
        )
        if node_type not in present_types
    ]
    if absent_structured_types:
        missing_information.append({
            "type": "unavailable_structured_node_types",
            "item_id": None,
            "message": (
                "Üretilen belge yapısında ayrı kimlikli kayıt bulunmadığı için "
                "uydurulmadan boş bırakılan düğüm türleri: " + ", ".join(absent_structured_types) + "."
            ),
        })

    lm_status = (
        check_lm_studio_status(request_get=request_get)
        if check_lm_studio
        else {
            "available": None,
            "message": "LM Studio kontrolü devre dışı; model çıkarımı yapılmadı.",
            "endpoint": None,
        }
    )
    if lm_status.get("available") is False:
        missing_information.append({
            "type": "lm_studio_unavailable",
            "item_id": None,
            "message": lm_status["message"],
        })

    confidence_counts = {
        CONFIDENCE_EXACT: sum(1 for edge in edges if edge["confidence_level"] == CONFIDENCE_EXACT),
        CONFIDENCE_SUGGESTED: sum(
            1 for edge in edges if edge["confidence_level"] == CONFIDENCE_SUGGESTED
        ),
        CONFIDENCE_INFERRED: sum(
            1 for edge in edges if edge["confidence_level"] == CONFIDENCE_INFERRED
        ),
    }
    report: dict[str, Any] = {
        "schema_version": SCHEMA_VERSION,
        "revision": 0,
        "project_id": project_id,
        "project_name": _clean(project_name),
        "generated_at": generated_at,
        "source_documents": source_documents,
        "nodes": sorted(nodes, key=lambda item: (item.get("document_type", ""), item["id"])),
        "edges": sorted(
            edges,
            key=lambda item: (
                item["relationship_type"], item["source_id"], item["target_id"]
            ),
        ),
        "unlinked_requirements": unlinked_requirements,
        "unverified_requirements": unverified_requirements,
        "conflicts": conflicts,
        "missing_information": missing_information,
        "capabilities": {
            "primary_source": "structured-python" if flat_data else "document-fallback",
            "exact_links": "bound_to ve linked_requirements",
            "semantic_links": "denetlenebilir metin benzerliği",
            "model_inference": "Bu aşamada otomatik model çıkarımı yapılmadı.",
            "lm_studio": lm_status,
        },
        "summary": {
            "node_count": len(nodes),
            "edge_count": len(edges),
            "source_document_count": len(source_documents),
            "unlinked_requirement_count": len(unlinked_requirements),
            "unverified_requirement_count": len(unverified_requirements),
            "conflict_count": len(conflicts),
            "confidence_counts": confidence_counts,
        },
    }
    if status_callback:
        status_callback(
            f"İzlenebilirlik haritası oluşturuldu: {len(nodes)} düğüm, {len(edges)} ilişki."
        )
    if persist:
        persist_traceability_report(report, output_root)
    return report


__all__ = [
    "atomic_write_json",
    "CONFIDENCE_EXACT",
    "CONFIDENCE_INFERRED",
    "CONFIDENCE_SUGGESTED",
    "DOCUMENT_TYPE_DEFINITIONS",
    "TraceabilityError",
    "build_traceability_map",
    "check_lm_studio_status",
    "extract_identifiers",
    "extract_technical_parameters",
    "load_project_traceability",
    "normalize_identifier",
    "persist_traceability_report",
    "project_identity",
    "read_document_lines",
    "read_document_records",
    "semantic_similarity",
]
