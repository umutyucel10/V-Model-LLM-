# -*- coding: utf-8 -*-
"""RAG (Retrieval-Augmented Generation) handler for IEEE 15288 traceability system"""

import os
import shutil
import threading
from typing import List, Dict, Optional, Tuple

from langchain_community.document_loaders import PyPDFDirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from langchain_chroma import Chroma
try:
    from langchain_community.vectorstores.utils import filter_complex_metadata
except ImportError:  # taşınmış sürümler için yedek
    from langchain_core.vectorstores.utils import filter_complex_metadata

from config import LMSTUDIO_BASE_URL, LMSTUDIO_API_KEY

import warnings
warnings.filterwarnings("ignore")

# Faz 9 performans bulgusu: paylasilan bir requests.Session(), tekil
# requests.get()/post() cagrilarina gore olculebilir sekilde (bu makinede
# ~10x) daha hizli - bkz. llm/handler.py'deki ayni not. embed_query() bir
# belge parcalama isleminde YUZLERCE kez ardisik cagrilabildigi icin burada
# etkisi ozellikle buyuk. Modulun kendi "ust seviyede requests'e sert
# bagimlilik olmasin" tercihini korumak icin session tembel (lazy) kuruluyor.
_session = None


def _get_session():
    global _session
    if _session is None:
        import requests
        _session = requests.Session()
    return _session

# Try to import PyMuPDF for better PDF processing
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("⚠️ PyMuPDF not available - advanced PDF processing will be skipped")

SIMPLE_DB_PATH = "rag_simple_db.txt"

class RAGHandler:
    """RAG handler for managing knowledge base and context retrieval"""

    def __init__(
        self, chroma_path: str = "rag_chroma", data_path: str = "rag_documents",
        *, initialize_embeddings: bool = True,
    ):
        self.chroma_path = chroma_path
        self.data_path = data_path
        self.embeddings = None
        self.db = None
        self.last_build_error = ""
        self._embeddings_initialization_attempted = False
        self._lock = threading.Lock()
        if initialize_embeddings:
            self.ensure_embeddings_initialized()
        else:
            self._update_db_path()

    # --------------------------- PDF PROCESSING --------------------------- #

    def _process_pdf_with_pymupdf(self, file_path: str) -> Optional[str]:
        """Process PDF using PyMuPDF for better text extraction"""
        if not PYMUPDF_AVAILABLE:
            print("⚠️ PyMuPDF not available - falling back to standard extraction")
            return None

        try:
            doc = fitz.open(file_path)
            text_content = []
            for page_num in range(doc.page_count):
                page = doc[page_num]
                page_text = page.get_text() or ""
                # Fallbacks for poor extraction
                if len(page_text.strip()) < 50:
                    page_text = page.get_text("text", sort=True) or ""
                if len(page_text.strip()) < 50:
                    blocks = page.get_text("dict")
                    block_texts = [
                        span.get("text", "")
                        for block in blocks.get("blocks", [])
                        if "lines" in block
                        for line in block["lines"]
                        for span in line["spans"]
                        if span.get("text", "").strip()
                    ]
                    page_text = " ".join(block_texts)
                if page_text.strip():
                    text_content.append(f"\n--- page {page_num + 1} ---\n{page_text}\n")
            doc.close()
            return "".join(text_content) if text_content else None
        except Exception as e:
            print(f"❌ PyMuPDF processing failed: {e}")
            return None

    # --------------------------- EMBEDDINGS --------------------------- #

    def _update_db_path(self):
        """Update database path based on embedding type to avoid dimension conflicts"""
        if self.embeddings and hasattr(self.embeddings, "base_url"):
            # LM Studio embeddings - use different path
            self.chroma_path = "rag_chroma_lms"
            print(f"🔄 Using LM Studio database path: {self.chroma_path}")
        else:
            # Other embeddings - use original path
            self.chroma_path = "rag_chroma"

    def _initialize_embeddings(self):
        """Initialize embedding function for vector database"""
        try:
            print(f"🔗 Testing LM Studio connection to {LMSTUDIO_BASE_URL}/models")
            headers = {"Authorization": f"Bearer {LMSTUDIO_API_KEY}"}
            response = _get_session().get(f"{LMSTUDIO_BASE_URL}/models", headers=headers, timeout=5)
            print(f"🔍 Response status: {response.status_code}")

            if response.status_code == 200:
                print("🔗 LM Studio detected, creating custom embeddings...")

                class LMStudioEmbeddings:
                    def __init__(self, base_url=LMSTUDIO_BASE_URL):
                        self.base_url = base_url
                        self.headers = {
                            "Content-Type": "application/json",
                            "Authorization": f"Bearer {LMSTUDIO_API_KEY}",
                        }

                    def embed_documents(self, texts: List[str]) -> List[List[float]]:
                        # Ensure one request per text, no silent drops (length must match)
                        vectors: List[List[float]] = []
                        for t in texts:
                            v = self.embed_query(t)
                            if v is None:
                                raise RuntimeError("Embedding failed for one or more documents.")
                            vectors.append(v)
                        return vectors

                    def embed_query(self, text: str) -> Optional[List[float]]:
                        try:
                            payload = {
                                "input": text,
                                "model": "text-embedding-nomic-embed-text-v1.5",
                            }
                            response = _get_session().post(
                                f"{self.base_url}/embeddings",
                                headers=self.headers,
                                json=payload,
                                timeout=90,
                            )
                            if response.status_code == 200:
                                result = response.json()
                                data = result.get("data") or []
                                if data and "embedding" in data[0]:
                                    return data[0]["embedding"]
                            print(f"❌ LM Studio embedding error: {response.status_code} - {response.text}")
                            return None
                        except Exception as e:
                            print(f"❌ LM Studio embedding request failed: {e}")
                            return None

                self.embeddings = LMStudioEmbeddings()
                # Ağır model yüklemesini uygulama başlangıcında tetiklememek için
                # deneme embedding'i gönderilmez. İlk gerçek indeksleme isteği hem
                # bağlantıyı doğrular hem de hata hâlinde basit metin aramasına düşer.
                print("✅ LM Studio embedding istemcisi hazır (tembel model yükleme)")
            else:
                print(f"❌ LM Studio not responding: {response.status_code}")
                self.embeddings = None

        except Exception as e:
            print(f"❌ LM Studio embeddings failed: {e}")
            print("RAG will work with simple text search only")
            self.embeddings = None

    def _get_lock(self) -> threading.Lock:
        # Eski test/entegrasyon nesneleri __new__ ile hazırlanmış olabilir,
        # bu yüzden kilit burada tembel (lazy) kuruluyor.
        lock = getattr(self, "_lock", None)
        if lock is None:
            lock = threading.Lock()
            self._lock = lock
        return lock

    def ensure_embeddings_initialized(self) -> bool:
        """Embedding istemcisini yalnızca RAG gerçekten kullanılacağı zaman hazırla."""
        if not hasattr(self, "_embeddings_initialization_attempted"):
            # Eski test/entegrasyon nesneleri __new__ ile hazırlanmış olabilir.
            return self.embeddings is not None
        if self._embeddings_initialization_attempted:
            return self.embeddings is not None
        # Faz 11: paylaşılan rag_handler tekil nesnesi birden fazla iş
        # parçacığından (belge üretimi + kopilot sohbeti aynı anda) çağrılabilir.
        # Kilit olmadan iki iş parçacığı bu kontrolü aynı anda geçip
        # _initialize_embeddings()'i (ağ isteği içerir) tekrar tekrar
        # çalıştırabilir. Çift kontrollü kilitleme (double-checked locking)
        # bunu önler.
        with self._get_lock():
            if self._embeddings_initialization_attempted:
                return self.embeddings is not None
            self._embeddings_initialization_attempted = True
            self._initialize_embeddings()
            self._update_db_path()
        return self.embeddings is not None

    # --------------------------- FS SETUP & LOADING --------------------------- #

    def setup_knowledge_base(self):
        """Setup the knowledge base directory structure"""
        if not os.path.exists(self.data_path):
            os.makedirs(self.data_path)
            print(f"📁 Created RAG documents directory: {self.data_path}")

        subdirs = ["ieee_standards", "radar_docs", "technical_specs", "existing_requirements"]
        for subdir in subdirs:
            subdir_path = os.path.join(self.data_path, subdir)
            if not os.path.exists(subdir_path):
                os.makedirs(subdir_path)
                print(f"📁 Created subdirectory: {subdir}")

        readme_path = os.path.join(self.data_path, "README.md")
        if not os.path.exists(readme_path):
            readme_content = """# RAG Knowledge Base

Bu klasör RAG (Retrieval-Augmented Generation) sistemi için dökümanları içerir.

## Klasör Yapısı:
- **ieee_standards/**: IEEE 15288 ve diğer standart dökümanları
- **radar_docs/**: Pasif radar sistemi ile ilgili teknik dökümanlar
- **technical_specs/**: Teknik spesifikasyonlar ve gereksinimler
- **existing_requirements/**: Mevcut proje gereksinimleri ve örnekler

## Desteklenen Dosya Formatları:
- PDF (.pdf) (çalıştığı onaylanan tek format)
- Text (.txt)
- Word (.docx)
- Markdown (.md)

## Kullanım:
1. İlgili klasörlere dökümanları kopyalayın
2. `build_knowledge_base()` fonksiyonunu çağırın
3. RAG sistemi otomatik olarak dökümanları işleyecek ve vector database oluşturacak
"""
            with open(readme_path, "w", encoding="utf-8") as f:
                f.write(readme_content)
            print(f"📄 Created README: {readme_path}")

    def _load_text_like_documents(self) -> List[Document]:
        """TXT/Markdown/DOCX belgelerini ek bağımlılık olmadan, dosya bazında yükle."""
        documents: List[Document] = []
        supported = {".txt", ".md", ".docx"}
        for root, _, files in os.walk(self.data_path):
            for filename in sorted(files):
                file_path = os.path.join(root, filename)
                suffix = os.path.splitext(filename)[1].lower()
                if suffix not in supported:
                    continue
                try:
                    if suffix in {".txt", ".md"}:
                        with open(file_path, "r", encoding="utf-8", errors="replace") as handle:
                            content = handle.read()
                    else:
                        from docx import Document as WordDocument
                        word = WordDocument(file_path)
                        parts = [
                            paragraph.text for paragraph in word.paragraphs
                            if paragraph.text.strip()
                        ]
                        for table in word.tables:
                            for row in table.rows:
                                values = [cell.text.strip() for cell in row.cells]
                                if any(values):
                                    parts.append(" | ".join(values))
                        content = "\n".join(parts)
                    if content.strip():
                        documents.append(Document(
                            page_content=content,
                            metadata={"source": file_path},
                        ))
                except Exception as error:
                    print(f"⚠️ Belge okunamadı ({file_path}): {error}")
        return documents

    def load_documents(self) -> List[Document]:
        """Load documents from the knowledge base directory."""
        if not os.path.exists(self.data_path):
            print(f"⚠️ Data path does not exist: {self.data_path}")
            return []

        documents: List[Document] = []

        try:
            # PDF files
            pdf_files = [
                os.path.join(root, file)
                for root, _, files in os.walk(self.data_path)
                for file in files
                if file.lower().endswith(".pdf")
            ]
            if pdf_files:
                print(f"🔧 Processing {len(pdf_files)} PDF files...")
                for pdf_file in pdf_files:
                    extracted_text = self._process_pdf_with_pymupdf(pdf_file)
                    if extracted_text:
                        documents.append(
                            Document(page_content=extracted_text, metadata={"source": pdf_file})
                        )
                    else:
                        try:
                            # Fallback loader for the directory of this PDF
                            fallback_loader = PyPDFDirectoryLoader(os.path.dirname(pdf_file))
                            fallback_docs = [
                                doc for doc in fallback_loader.load()
                                if doc.metadata.get("source") == pdf_file
                            ]
                            documents.extend(fallback_docs)
                        except Exception as e:
                            print(f"⚠️ Failed to process {pdf_file}: {e}")

            # Other file types
            documents.extend(self._load_text_like_documents())

        except Exception as e:
            print(f"❌ Error loading documents: {e}")

        print(f"📚 Total documents loaded: {len(documents)}")
        return documents

    # --------------------------- CHUNKING --------------------------- #

    def split_documents(
        self, documents: List[Document], chunk_size: int = 1200, chunk_overlap: int = 200
    ) -> List[Document]:
        """Split documents into chunks for better retrieval with word-aware splitting"""
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False,
            separators=[
                "\n\n",  # Paragraph breaks
                "\n",    # Line breaks
                ". ",    # Sentence endings
                "! ",
                "? ",
                "; ",
                ", ",
                " ",     # Word breaks (spaces)
                "",      # Character breaks (last resort)
            ],
        )

        chunks = text_splitter.split_documents(documents)
        print(f"✂️ Split {len(documents)} documents into {len(chunks)} chunks (size: {chunk_size}, overlap: {chunk_overlap})")
        print("📝 Used word-aware splitting to preserve word boundaries")
        return chunks

    def calculate_chunk_ids(self, chunks: List[Document]) -> List[Document]:
        """Calculate unique IDs for document chunks"""
        for i, chunk in enumerate(chunks):
            source = chunk.metadata.get("source", "unknown")
            filename = os.path.basename(source)
            chunk_id = f"{filename}:chunk_{i:04d}"
            chunk.metadata["id"] = chunk_id
            chunk.metadata["chunk_index"] = i

        print(f"🆔 Assigned IDs to {len(chunks)} chunks")
        return chunks

    # --------------------------- BUILD / LOAD DB --------------------------- #

    def _load_existing_database(self):
        """Load existing vector database"""
        if self.embeddings:
            try:
                self.db = Chroma(
                    persist_directory=self.chroma_path, embedding_function=self.embeddings
                )
                print(f"✅ Loaded existing knowledge base from {self.chroma_path}")
            except Exception as e:
                print(f"❌ Error loading existing database: {e}")
                self.db = None

    def ensure_database_loaded(self) -> bool:
        """Ensure database is loaded if it exists"""
        self.ensure_embeddings_initialized()
        if not self.db and self.embeddings and os.path.exists(self.chroma_path):
            # Faz 11: aynı çift kontrollü kilitleme, eşzamanlı çağrıların
            # aynı anda birden fazla Chroma bağlantısı açıp birbirinin
            # üzerine yazmasını (bağlantı sızıntısı) önlemek için.
            with self._get_lock():
                if not self.db:
                    print(f"🔄 Loading knowledge base from {self.chroma_path}")
                    self._load_existing_database()
        return self.db is not None

    def close_database(self):
        """Close the database connection"""
        if not self.db:
            return
        try:
            if hasattr(self.db, "_client") and self.db._client:
                try:
                    self.db._client.reset()
                except Exception:
                    pass
                try:
                    self.db._client.close()
                except Exception:
                    pass
            self.db = None
            import gc
            gc.collect()
            print("🔒 Database connection closed")
        except Exception as e:
            print(f"⚠️ Warning closing database: {e}")
            self.db = None

    def _build_simple_knowledge_base(self, chunks: List[Document], force_rebuild: bool = False):
        """Build a simple text-based knowledge base without embeddings"""
        if os.path.exists(SIMPLE_DB_PATH) and not force_rebuild:
            print("ℹ️ Simple knowledge base already exists")
            return

        print("🔧 Building simple text-based knowledge base...")
        with open(SIMPLE_DB_PATH, "w", encoding="utf-8") as f:
            for chunk in chunks:
                f.write(f"ID: {chunk.metadata.get('id', 'unknown')}\n")
                f.write(f"SOURCE: {chunk.metadata.get('source', 'unknown')}\n")
                f.write(f"CONTENT: {chunk.page_content}\n")
                f.write("---CHUNK_SEPARATOR---\n")

        self.simple_db_path = SIMPLE_DB_PATH
        print(f"✅ Simple knowledge base built: {len(chunks)} chunks stored")

    def build_knowledge_base(self, force_rebuild: bool = False, chunk_size: int = 1200, chunk_overlap: int = 200) -> bool:
        """Build or update the knowledge base."""
        print("📚 Building knowledge base...")
        self.last_build_error = ""
        self.ensure_embeddings_initialized()

        # If vector DB exists and no rebuild requested, avoid re-processing documents
        if self.embeddings and os.path.exists(self.chroma_path) and not force_rebuild:
            if self.ensure_database_loaded():
                print(f"ℹ️ Knowledge base already exists at {self.chroma_path}")
                return True

        documents = self.load_documents()
        if not documents:
            print("⚠️ No documents found.")
            self.last_build_error = (
                "RAG belge klasöründe okunabilir TXT, Markdown, DOCX veya PDF içeriği bulunamadı."
            )
            return False

        chunks = self.split_documents(documents, chunk_size, chunk_overlap)
        chunks = self.calculate_chunk_ids(chunks)

        # Vector DB path
        if self.embeddings:
            db_exists = os.path.exists(self.chroma_path)
            if force_rebuild and db_exists:
                self.close_database()
                try:
                    shutil.rmtree(self.chroma_path)
                    print(f"🗑️ Removed existing database: {self.chroma_path}")
                except Exception as e:
                    print(f"❌ Error removing database: {e}")
                    self.last_build_error = f"Eski vektör indeksi temizlenemedi: {e}"
                    return False
            try:
                self.db = Chroma.from_documents(
                    filter_complex_metadata(chunks),
                    self.embeddings,
                    persist_directory=self.chroma_path,
                )
                print(f"✅ Vector DB built: {len(chunks)} chunks.")
                return True
            except Exception as e:
                print(f"❌ Vector DB error: {e}")
                # Fall back to simple DB

        # Simple text DB fallback
        try:
            self._build_simple_knowledge_base(chunks, force_rebuild)
            return True
        except Exception as e:
            print(f"❌ Error building simple knowledge base: {e}")
            self.last_build_error = f"Basit metin indeksi oluşturulamadı: {e}"
            return False

    # --------------------------- QUERYING --------------------------- #

    def query_knowledge_base(self, query: str, k: int = 2) -> List[Tuple[Document, float]]:
        """Query the knowledge base for relevant document chunks"""
        self.ensure_embeddings_initialized()
        if self.db:
            try:
                results = self.db.similarity_search_with_relevance_scores(query, k=k)
                # Filter: keep higher-score chunks first; relax threshold if necessary
                score_threshold = 0.5
                filtered = [(doc, score) for doc, score in results if score >= score_threshold]
                if not filtered:
                    score_threshold = 0.3
                    filtered = [(doc, score) for doc, score in results if score >= score_threshold]
                print(f"🔍 Found {len(filtered)} relevant chunks (score ≥ {score_threshold})")
                return filtered
            except Exception as e:
                print(f"❌ Vector search failed: {e}")

        # Fallback simple search
        return self._simple_text_search(query, k)

    def _simple_text_search(self, query: str, k: int = 5) -> List[Tuple[Document, float]]:
        """Simple text-based search when vector search is not available"""
        if not os.path.exists(SIMPLE_DB_PATH):
            print("❌ Simple knowledge base not found")
            return []

        try:
            print(f"🔍 Using simple text search for: '{query}'")
            query_words = query.lower().split()
            results: List[Tuple[Document, float]] = []

            with open(SIMPLE_DB_PATH, "r", encoding="utf-8") as f:
                chunks = f.read().split("---CHUNK_SEPARATOR---")

            for chunk in chunks:
                if not chunk.strip():
                    continue
                chunk_id, source = "", ""
                content_lines: List[str] = []
                reading_content = False
                for line in chunk.strip().split("\n"):
                    if not reading_content and line.startswith("ID: "):
                        chunk_id = line[4:]
                    elif not reading_content and line.startswith("SOURCE: "):
                        source = line[8:]
                    elif not reading_content and line.startswith("CONTENT: "):
                        reading_content = True
                        content_lines.append(line[9:])
                    elif reading_content:
                        content_lines.append(line)

                chunk_content = "\n".join(content_lines).strip()

                if chunk_content:
                    content_lower = chunk_content.lower()
                    score = sum(content_lower.count(word) for word in query_words) / max(len(query_words), 1)
                    if score > 0:
                        doc = Document(page_content=chunk_content, metadata={"id": chunk_id, "source": source})
                        results.append((doc, score))

            results.sort(key=lambda x: x[1], reverse=True)
            print(f"🔍 Found {len(results[:k])} relevant documents (simple search)")
            return results[:k]
        except Exception as e:
            print(f"❌ Simple text search failed: {e}")
            return []

    # --------------------------- CONTEXT & INFO --------------------------- #

    def get_enhanced_context(self, tid_description: str, requirement_type: str = "SGD") -> str:
        """Get enhanced context for requirement generation using RAG"""
        try:
            self.ensure_database_loaded()

            search_queries = [
                f"{tid_description} {requirement_type}",
                f"system requirements {tid_description}",
                f"IEEE 15288 {requirement_type}",
                f"radar system {tid_description}",
            ]

            all_context: List[str] = []
            seen_content = set()

            for q in search_queries:
                print(f"🔍 Running query: {q}\n\n")
                try:
                    results = self.query_knowledge_base(q, k=2)
                    for doc, score in results:
                        content = doc.page_content.strip()
                        if score > 0.5 and content not in seen_content and len(content) > 50:
                            all_context.append(f"[Relevance: {score:.2f}] {content}")
                            seen_content.add(content)
                except Exception as e:
                    print(f"⚠️ Query failed for '{q}': {e}")

            if all_context:
                enhanced_context = "\n\n---\n\n".join(all_context[:2])  # Limit to top 2 results
                print(f"🔧 Enhanced context generated: {len(enhanced_context)} characters")
                return enhanced_context

            print("ℹ️ No relevant context found in knowledge base")
            return ""
        except Exception as e:
            print(f"❌ Enhanced context generation failed: {e}")
            return ""

    def get_database_info(self) -> Dict:
        """Get information about the knowledge base"""
        info: Dict = {
            "chroma_path": self.chroma_path,
            "data_path": self.data_path,
            "database_exists": os.path.exists(self.chroma_path),
            "documents_folder_exists": os.path.exists(self.data_path),
            "embeddings_initialized": self.embeddings is not None,
            "database_loaded": self.db is not None,
            "simple_db_exists": os.path.exists(SIMPLE_DB_PATH),
        }

        if self.db:
            try:
                collection_info = self.db.get(include=[])
                info["total_chunks"] = len(collection_info["ids"])
                info["search_method"] = "vector"
            except Exception:
                info["total_chunks"] = "unknown"
                info["search_method"] = "vector (error)"
        elif os.path.exists(SIMPLE_DB_PATH):
            try:
                with open(SIMPLE_DB_PATH, "r", encoding="utf-8") as f:
                    content = f.read()
                chunks = content.split("---CHUNK_SEPARATOR---")
                info["total_chunks"] = len([c for c in chunks if c.strip()])
                info["search_method"] = "simple_text"
            except Exception:
                info["total_chunks"] = "unknown"
                info["search_method"] = "simple_text (error)"
        else:
            info["total_chunks"] = 0
            info["search_method"] = "none"

        return info

    def add_document_to_knowledge_base(self, file_path: str, category: str = "technical_specs") -> bool:
        """Add a single document to the existing knowledge base"""
        if not os.path.exists(file_path):
            print(f"❌ File not found: {file_path}")
            return False

        filename = os.path.basename(file_path)
        target_dir = os.path.join(self.data_path, category)
        target_path = os.path.join(target_dir, filename)

        if not os.path.exists(target_dir):
            os.makedirs(target_dir)

        try:
            shutil.copy2(file_path, target_path)
            print(f"📄 Added document to knowledge base: {filename}")
            print("💡 Suggestion: Run build_knowledge_base(force_rebuild=True) to update the vector database")
            return True
        except Exception as e:
            print(f"❌ Error adding document: {e}")
            return False
        
# --------------------------- MODULE-LEVEL FUNCTIONS --------------------------- #

# Global RAG handler instance
rag_handler = RAGHandler(initialize_embeddings=False)

def initialize_rag_system() -> RAGHandler:
    """Initialize the RAG system"""
    print("🚀 Initializing RAG system...")
    rag_handler.ensure_embeddings_initialized()
    rag_handler.setup_knowledge_base()
    return rag_handler

def build_rag_knowledge_base(force_rebuild: bool = False) -> bool:
    """Build the RAG knowledge base"""
    return rag_handler.build_knowledge_base(force_rebuild=force_rebuild)

def get_rag_enhanced_context(tid_description: str, requirement_type: str = "SGD") -> str:
    """Get RAG-enhanced context for LLM generation"""
    return rag_handler.get_enhanced_context(tid_description, requirement_type)

def get_rag_info() -> Dict:
    """Get RAG system information"""
    return rag_handler.get_database_info()
