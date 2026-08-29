# -*- coding: utf-8 -*-
"""Etki Analizi için güvenli değişiklik paketi ve atomik sürümleme.

Bu modül özgün belge verisini hiçbir zaman yerinde değiştirmez. Onaylanan
öneriler bir çalışma kopyasına uygulanır, yeni sürüm geçici klasörde üretilip
doğrulanır ve ancak bütün adımlar başarılı olduğunda atomik olarak yayımlanır.
"""

from __future__ import annotations

from copy import deepcopy
from dataclasses import asdict, dataclass, field, replace
from datetime import datetime, timezone
from html import escape
import hashlib
import json
import os
from pathlib import Path
import re
import shutil
import tempfile
import threading
from typing import Any, Callable, Mapping, Sequence

from etki_analizi_izlenebilirlik import DOCUMENT_TYPE_DEFINITIONS
import etki_analizi_simulasyon as simulation


# Faz 7'de bu dosya proje kokunden etki_analizi/ alt paketine tasindi;
# __file__.parent artik alt paketi gosterdigi icin proje kokune cikmak icin
# bir ust dizine (.parent.parent) cikiyoruz - davranis (outputs/
# change_control'u proje kokunde bulmak) tasimadan onceki haliyle ayni
# kalsin diye.
DEFAULT_OUTPUT_ROOT = Path(__file__).resolve().parent.parent / "outputs" / "change_control"

DECISION_ACCEPT = "Kabul et"
DECISION_REJECT = "Reddet"
DECISION_EDIT = "Düzenle"
DECISION_DEFER = "Ertele"
DECISIONS = (DECISION_ACCEPT, DECISION_REJECT, DECISION_EDIT, DECISION_DEFER)

CATEGORY_MAIN = "Değişen ana gereksinim"
CATEGORY_CUSTOMER = "Müşteri gereksinimleri"
CATEGORY_SYSTEM = "Sistem gereksinimleri"
CATEGORY_SUBSYSTEM = "Alt sistem gereksinimleri"
CATEGORY_DESIGN = "Tasarım ve arayüzler"
CATEGORY_VERIFICATION = "Doğrulama kriterleri"
CATEGORY_TEST = "Test prosedürleri"
CATEGORY_NEW_TEST = "Yeni test önerileri"
CATEGORY_ACCEPTANCE = "Kabul kriterleri"
CATEGORY_RISK = "Risk azaltma faaliyetleri"
CATEGORY_QUESTION = "Açık sorular"
CATEGORY_ASSUMPTION = "Varsayımlar"

RECORD_DOCUMENT_TYPES = {"TID", "SGD", "STT", "KMTD", "SITET", "AST"}
_LOCK = threading.Lock()


class ChangePackageError(RuntimeError):
    """Kullanıcıya gösterilebilecek paket/sürümleme hatası."""


def _now(value: datetime | None = None) -> datetime:
    return value or datetime.now(timezone.utc).astimezone()


def _clean(value: Any) -> str:
    return " ".join(str(value or "").split())


def _safe(value: Any) -> str:
    text = "".join(
        char if char.isalnum() or char in "-_" else "-" for char in _clean(value)
    ).strip("-")
    return text or "proje"


def _json_ready(value: Any) -> Any:
    if hasattr(value, "__dataclass_fields__"):
        return _json_ready(asdict(value))
    if isinstance(value, Mapping):
        return {str(key): _json_ready(item) for key, item in value.items()}
    if isinstance(value, (list, tuple, set)):
        return [_json_ready(item) for item in value]
    if isinstance(value, datetime):
        return value.isoformat(timespec="seconds")
    return value


def _atomic_text(path: Path, content: str) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    temporary_name = ""
    try:
        with tempfile.NamedTemporaryFile(
            mode="w", encoding="utf-8", dir=path.parent,
            prefix=f".{path.name}.", suffix=".tmp", delete=False,
        ) as temporary:
            temporary.write(content)
            temporary.flush()
            os.fsync(temporary.fileno())
            temporary_name = temporary.name
        os.replace(temporary_name, path)
    finally:
        if temporary_name and os.path.exists(temporary_name):
            os.unlink(temporary_name)


def _atomic_json(path: Path, value: Mapping[str, Any]) -> None:
    _atomic_text(
        path, json.dumps(_json_ready(value), ensure_ascii=False, indent=2) + "\n"
    )


def _digest(value: Any) -> str:
    encoded = json.dumps(
        _json_ready(value), ensure_ascii=False, sort_keys=True, separators=(",", ":")
    ).encode("utf-8")
    return hashlib.sha256(encoded).hexdigest()


@dataclass
class UpdateProposal:
    proposal_id: str
    category: str
    document_name: str
    section: str
    requirement_id: str
    current_text: str
    proposed_text: str
    rationale: str
    impact_path: str
    risk_level: str
    decision: str = DECISION_DEFER
    target_kind: str = "structured_record"
    node_type: str = ""
    document_type: str = ""
    edited: bool = False
    assigned_id: str = ""

    @classmethod
    def from_mapping(cls, value: Mapping[str, Any]) -> "UpdateProposal":
        return cls(
            proposal_id=_clean(value.get("proposal_id")),
            category=_clean(value.get("category")),
            document_name=_clean(value.get("document_name")),
            section=_clean(value.get("section")),
            requirement_id=_clean(value.get("requirement_id")),
            current_text=str(value.get("current_text") or ""),
            proposed_text=str(value.get("proposed_text") or ""),
            rationale=_clean(value.get("rationale")),
            impact_path=_clean(value.get("impact_path")),
            risk_level=_clean(value.get("risk_level")) or "Belirsiz",
            decision=_clean(value.get("decision")) or DECISION_DEFER,
            target_kind=_clean(value.get("target_kind")) or "structured_record",
            node_type=_clean(value.get("node_type")),
            document_type=_clean(value.get("document_type")),
            edited=bool(value.get("edited")),
            assigned_id=_clean(value.get("assigned_id")),
        )

    def to_dict(self) -> dict[str, Any]:
        return _json_ready(self)


@dataclass
class ChangePackage:
    change_id: str
    project_id: str
    project_name: str
    created_at: str
    change_request: dict[str, Any]
    request_summary: str
    selected_item: dict[str, Any] | None
    proposals: list[UpdateProposal]
    affected_documents: list[dict[str, Any]]
    engineering_ideas: list[dict[str, Any]]
    risks: list[dict[str, Any]]
    open_questions: list[str]
    assumptions: list[str]
    simulation_snapshot: dict[str, Any]
    baseline_traceability: dict[str, Any]
    baseline_fingerprint: str
    status: str = "Taslak"
    approval_confirmed: bool = False
    approval_actor: str = ""
    approval_at: str = ""
    approval_digest: str = ""
    storage_path: str = ""

    @classmethod
    def from_mapping(cls, value: Mapping[str, Any]) -> "ChangePackage":
        return cls(
            change_id=_clean(value.get("change_id")),
            project_id=_clean(value.get("project_id")),
            project_name=_clean(value.get("project_name")),
            created_at=_clean(value.get("created_at")),
            change_request=dict(value.get("change_request") or {}),
            request_summary=_clean(value.get("request_summary")),
            selected_item=(
                dict(value.get("selected_item"))
                if isinstance(value.get("selected_item"), Mapping) else None
            ),
            proposals=[
                UpdateProposal.from_mapping(item)
                for item in value.get("proposals", [])
                if isinstance(item, Mapping)
            ],
            affected_documents=[
                dict(item) for item in value.get("affected_documents", [])
                if isinstance(item, Mapping)
            ],
            engineering_ideas=[
                dict(item) for item in value.get("engineering_ideas", [])
                if isinstance(item, Mapping)
            ],
            risks=[
                dict(item) for item in value.get("risks", [])
                if isinstance(item, Mapping)
            ],
            open_questions=[_clean(item) for item in value.get("open_questions", [])],
            assumptions=[_clean(item) for item in value.get("assumptions", [])],
            simulation_snapshot=dict(value.get("simulation_snapshot") or {}),
            baseline_traceability=dict(value.get("baseline_traceability") or {}),
            baseline_fingerprint=_clean(value.get("baseline_fingerprint")),
            status=_clean(value.get("status")) or "Taslak",
            approval_confirmed=bool(value.get("approval_confirmed")),
            approval_actor=_clean(value.get("approval_actor")),
            approval_at=_clean(value.get("approval_at")),
            approval_digest=_clean(value.get("approval_digest")),
            storage_path=_clean(value.get("storage_path")),
        )

    def to_dict(self) -> dict[str, Any]:
        return _json_ready(self)


@dataclass
class ApplicationResult:
    change_id: str
    project_id: str
    previous_version: int
    new_version: int
    status: str
    version_directory: str
    backup_directory: str
    created_documents: list[str]
    modified_item_ids: list[str]
    added_item_ids: list[str]
    non_document_actions: list[str]
    new_flat_data: dict[str, dict[str, Any]]
    new_hardware_data: dict[str, dict[str, Any]]
    post_traceability: dict[str, Any] = field(default_factory=dict)
    post_simulation: dict[str, Any] = field(default_factory=dict)
    closure_summary: dict[str, Any] = field(default_factory=dict)
    report_paths: dict[str, str] = field(default_factory=dict)
    warnings: list[str] = field(default_factory=list)

    def to_dict(self) -> dict[str, Any]:
        return _json_ready(self)


def _category(node_type: str, selected: bool = False) -> str:
    if selected and node_type in {
        "Müşteri/paydaş gereksinimi", "Sistem gereksinimi", "Alt sistem gereksinimi"
    }:
        return CATEGORY_MAIN
    return {
        "Müşteri/paydaş gereksinimi": CATEGORY_CUSTOMER,
        "Sistem gereksinimi": CATEGORY_SYSTEM,
        "Alt sistem gereksinimi": CATEGORY_SUBSYSTEM,
        "Fonksiyon": CATEGORY_DESIGN,
        "Parça/bileşen": CATEGORY_DESIGN,
        "Yazılım birimi": CATEGORY_DESIGN,
        "Mekanik arayüz": CATEGORY_DESIGN,
        "Elektriksel arayüz": CATEGORY_DESIGN,
        "Yazılımsal arayüz": CATEGORY_DESIGN,
        "Tasarım kararı": CATEGORY_DESIGN,
        "Doğrulama kriteri": CATEGORY_VERIFICATION,
        "Birim testi": CATEGORY_TEST,
        "Entegrasyon testi": CATEGORY_TEST,
        "Sistem doğrulama testi": CATEGORY_TEST,
        "Müşteri kabul/geçerleme testi": CATEGORY_ACCEPTANCE,
    }.get(node_type, CATEGORY_DESIGN)


def _replace_explicit_value(text: str, old: Any, new: Any) -> tuple[str, bool]:
    old_text, new_text = _clean(old), _clean(new)
    if not text or not old_text or not new_text:
        return text, False
    pattern = re.compile(re.escape(old_text), re.IGNORECASE)
    replaced, count = pattern.subn(new_text, text, count=1)
    if count:
        return replaced, True
    numeric = re.fullmatch(
        r"(?P<number>[+-]?\d+(?:[.,]\d+)?)\s*(?P<unit>[^\d\s]+)?", old_text
    )
    proposed = re.fullmatch(
        r"(?P<number>[+-]?\d+(?:[.,]\d+)?)\s*(?P<unit>[^\d\s]+)?", new_text
    )
    if not numeric or not proposed:
        return text, False
    old_number = re.escape(numeric.group("number")).replace(r"\,", "[,.]")
    old_unit = re.escape(numeric.group("unit") or "")
    flexible = re.compile(
        rf"(?<!\d){old_number}\s*{old_unit}(?![\w])", re.IGNORECASE
    )
    replaced, count = flexible.subn(new_text, text, count=1)
    return replaced, bool(count)


def _draft_text(
    node: Mapping[str, Any], request: simulation.ChangeRequest, selected: bool
) -> tuple[str, str]:
    current = _clean(node.get("description"))
    if selected:
        if request.change_type == simulation.CHANGE_REQUIREMENT_REMOVE:
            return "", "Ana gereksinimin kaldırılması kullanıcı tarafından önerildi."
        if request.change_type != simulation.CHANGE_REQUIREMENT_TEXT:
            replaced, changed = _replace_explicit_value(
                current, request.current_value, request.proposed_value
            )
            if changed:
                return replaced, (
                    "Değişiklik isteğindeki açık mevcut/yeni değer ana kaynak metinde değiştirildi."
                )
            return current, (
                "Ana kaynak metinde açık değer eşleşmesi bulunamadı. Güvenli uygulama için "
                "öneri kullanıcı tarafından 'Düzenle' seçeneğiyle tamamlanmalıdır."
            )
        return _clean(request.proposed_value), "Değişiklik isteğindeki açık yeni içerik kullanıldı."
    replaced, changed = _replace_explicit_value(
        current, request.current_value, request.proposed_value
    )
    if changed:
        return replaced, "Değişiklik isteğindeki açık mevcut/yeni değer kaynak metinde değiştirildi."
    source_id = request.requirement_id or "değişen gereksinim"
    proposed_value = _clean(request.proposed_value)
    compact_value = proposed_value if len(proposed_value) <= 80 else ""
    node_type = _clean(node.get("node_type"))
    if node_type in {
        "Müşteri/paydaş gereksinimi", "Sistem gereksinimi", "Alt sistem gereksinimi"
    }:
        suffix = (
            f" Bu gereksinimin {source_id} için önerilen {compact_value} "
            "değişikliğiyle uyumu doğrulanmalıdır."
            if compact_value else
            f" Bu gereksinimin {source_id} değişikliğiyle uyumu doğrulanmalıdır."
        )
    elif node_type in {
        "Doğrulama kriteri", "Birim testi", "Entegrasyon testi",
        "Sistem doğrulama testi", "Müşteri kabul/geçerleme testi",
    }:
        suffix = (
            f" Test, {source_id} için önerilen {compact_value} değerini "
            "doğrulayacak şekilde güncellenmelidir."
            if compact_value else
            f" Test, {source_id} değişikliğini doğrulayacak şekilde güncellenmelidir."
        )
    else:
        suffix = (
            f" Bu öğe, {source_id} değişikliğine uyum açısından yeniden "
            "değerlendirilmelidir."
        )
    return (current + suffix).strip(), (
        "Açık metin eşleşmesi bulunmadığı için kaynak değer uydurulmadan "
        "mühendislik inceleme cümlesi eklendi."
    )


def _proposal_id(change_id: str, category: str, item_id: str, path: str) -> str:
    suffix = hashlib.sha256(
        f"{change_id}|{category}|{item_id}|{path}".encode("utf-8")
    ).hexdigest()[:10].upper()
    return f"UPD-{suffix}"


def _node_source(node: Mapping[str, Any]) -> tuple[str, str]:
    return (
        _clean(node.get("source_document")) or "Belge adı belirtilmemiş",
        _clean(
            node.get("source_section")
            or node.get("section")
            or node.get("page_section")
        ) or "Bölüm belirtilmemiş",
    )


def _target_kind(node: Mapping[str, Any]) -> str:
    document_type = _clean(node.get("document_type"))
    if document_type in RECORD_DOCUMENT_TYPES:
        return "structured_record"
    if document_type == "HARDWARE":
        return "hardware_record"
    return "reference_only"


def _new_requirement_document_type(identifier: str) -> str:
    folded = _clean(identifier).upper().replace("_", "-")
    if folded.startswith(("TID-", "UR-", "CR-", "CUST-")):
        return "TID"
    if folded.startswith(("SGD-", "SR-", "SYS-")):
        return "SGD"
    if folded.startswith(("STT-", "SSR-", "SUB-")):
        return "STT"
    return ""


def build_change_package(
    result: simulation.SimulationResult,
    traceability: Mapping[str, Any],
    *,
    now: datetime | None = None,
) -> ChangePackage:
    """Simülasyon sonucundan kaynaksız değer uydurmayan değişiklik taslakları üretir."""
    if result.status != "completed":
        raise ChangePackageError(
            "Değişiklik paketi için tamamlanmış bir simülasyon sonucu gereklidir."
        )
    report = deepcopy(dict(traceability))
    nodes = {
        _clean(node.get("id")): dict(node)
        for node in report.get("nodes", [])
        if isinstance(node, Mapping) and _clean(node.get("id"))
    }
    request = result.change_request
    created = _now(now)
    seed = {
        "project": report.get("project_id"), "request": request.to_dict(),
        "generated": created.isoformat(timespec="seconds"),
    }
    change_id = (
        f"CR-{created.strftime('%Y%m%d-%H%M%S')}-"
        f"{_digest(seed)[:6].upper()}"
    )
    selected_id = _clean((result.selected_item or {}).get("id"))
    proposals: list[UpdateProposal] = []
    seen: set[tuple[str, str]] = set()

    candidates: list[tuple[dict[str, Any], str, str, int, bool]] = []
    if request.change_type == simulation.CHANGE_REQUIREMENT_ADD:
        identifier = _clean(request.requirement_id)
        document_type = _new_requirement_document_type(identifier)
        definition = DOCUMENT_TYPE_DEFINITIONS.get(document_type, {})
        target_kind = "new_requirement" if identifier and document_type else "governance_action"
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, CATEGORY_MAIN, identifier or "YENI", identifier),
            category=CATEGORY_MAIN,
            document_name=_clean(definition.get("document_title")) or "Belge türü kullanıcı tarafından seçilmeli",
            section="Yeni gereksinim",
            requirement_id=identifier,
            current_text="",
            proposed_text=_clean(request.proposed_value),
            rationale=(
                "Değişiklik isteğinde kullanıcı tarafından verilen yeni gereksinim kimliği ve metni kullanıldı."
                if target_kind == "new_requirement" else
                "Gereksinim kimliğinden güvenli bir V-Model belge türü belirlenemedi; uygulamadan önce belge türü netleştirilmelidir."
            ),
            impact_path=identifier or "Yeni gereksinim · tahsis yolu bekleniyor",
            risk_level=_clean(result.summary.get("overall_impact_level")) or "Belirsiz",
            target_kind=target_kind,
            node_type=_clean(definition.get("node_type")) or "Yeni gereksinim",
            document_type=document_type,
        ))
    if result.selected_item:
        selected = dict(result.selected_item)
        candidates.append((
            selected, selected_id, selected_id,
            int(result.summary.get("overall_impact_score", 0) or 0), True,
        ))
    for impact in result.impacts:
        node = nodes.get(impact.item_id)
        if node and _clean(node.get("node_type")) != "Teknik belge":
            candidates.append((
                node, impact.item_id, impact.traceability_path.display_path,
                impact.impact_score, False,
            ))

    for node, item_id, path, score, selected in candidates:
        key = (_clean(node.get("node_type")), item_id)
        if key in seen:
            continue
        seen.add(key)
        category = _category(_clean(node.get("node_type")), selected)
        document_name, section = _node_source(node)
        proposed, drafting_reason = _draft_text(node, request, selected)
        impact = next((item for item in result.impacts if item.item_id == item_id), None)
        rationale = (
            impact.rationale if impact else "Değişiklik başlangıç noktasıdır."
        )
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, category, item_id, path),
            category=category,
            document_name=document_name,
            section=section,
            requirement_id=item_id,
            current_text=_clean(node.get("description")),
            proposed_text=proposed,
            rationale=f"{rationale} {drafting_reason}".strip(),
            impact_path=path,
            risk_level=(
                impact.impact_level if impact
                else _clean(result.summary.get("overall_impact_level")) or "Belirsiz"
            ),
            target_kind=_target_kind(node),
            node_type=_clean(node.get("node_type")),
            document_type=_clean(node.get("document_type")),
        ))

    for index, action in enumerate(
        result.categorized_impacts.get("new_or_updated_tests", [])
    ):
        proposed = _clean(action.get("required_action") or action.get("reason"))
        if not proposed:
            continue
        item_id = _clean(action.get("test_id")) or f"YENİ-TEST-{index + 1:03d}"
        raw_path = action.get("path")
        path = (
            _clean(raw_path.get("display_path"))
            if isinstance(raw_path, Mapping) else _clean(raw_path)
        ) or selected_id
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, CATEGORY_NEW_TEST, item_id, path),
            category=CATEGORY_NEW_TEST,
            document_name="Test Güncelleme Planı",
            section="Yeni veya güncellenecek testler",
            requirement_id=item_id,
            current_text=_clean(action.get("reason")),
            proposed_text=proposed,
            rationale=_clean(action.get("reason")) or "Değişiklik için yeni doğrulama kapsamı gereklidir.",
            impact_path=path,
            risk_level="Yüksek",
            target_kind="new_test" if not action.get("test_id") else "reference_only",
            node_type="Doğrulama testi",
            document_type="",
        ))

    risk_rows = [item.to_dict() for item in result.risks]
    for index, risk in enumerate(risk_rows):
        item_id = f"RISK-ACTION-{index + 1:03d}"
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, CATEGORY_RISK, item_id, selected_id),
            category=CATEGORY_RISK,
            document_name="Risk ve Aksiyon Listesi",
            section=_clean(risk.get("category")) or "Risk",
            requirement_id=item_id,
            current_text=_clean(risk.get("rationale")),
            proposed_text=(
                f"{_clean(risk.get('rationale'))} için sorumlu, hedef tarih ve "
                "doğrulama kanıtı tanımlanmalıdır."
            ),
            rationale="Simülasyondaki deterministik risk sonucu için azaltma faaliyeti taslağıdır.",
            impact_path=" → ".join(risk.get("impacted_items") or ()) or selected_id,
            risk_level=_clean(risk.get("impact_level")) or "Belirsiz",
            target_kind="governance_action",
            node_type="Risk",
        ))

    open_questions = list(dict.fromkeys(
        [_clean(item) for item in result.warnings if _clean(item)]
        + [
            _clean(item.get("message"))
            for item in report.get("missing_information", [])
            if isinstance(item, Mapping) and _clean(item.get("message"))
        ]
    ))
    for index, question in enumerate(open_questions):
        item_id = f"OPEN-{index + 1:03d}"
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, CATEGORY_QUESTION, item_id, selected_id),
            category=CATEGORY_QUESTION,
            document_name="Değişiklik İsteği Kaydı",
            section="Açık sorular",
            requirement_id=item_id,
            current_text=question,
            proposed_text="Yanıt ve sorumlu kullanıcı tarafından tanımlanmalıdır.",
            rationale="Eksik veya düşük güvenli bilgi kapanıştan önce netleştirilmelidir.",
            impact_path=selected_id,
            risk_level="Belirsiz",
            target_kind="governance_action",
            node_type="Açık soru",
        ))

    assumptions = list(request.assumptions)
    for index, assumption in enumerate(assumptions):
        item_id = f"ASM-{index + 1:03d}"
        proposals.append(UpdateProposal(
            proposal_id=_proposal_id(change_id, CATEGORY_ASSUMPTION, item_id, selected_id),
            category=CATEGORY_ASSUMPTION,
            document_name="Değişiklik İsteği Kaydı",
            section="Varsayımlar",
            requirement_id=item_id,
            current_text=_clean(assumption),
            proposed_text=_clean(assumption),
            rationale="Simülasyonda kullanılan varsayım denetim kaydına alınmıştır.",
            impact_path=selected_id,
            risk_level="Belirsiz",
            target_kind="governance_action",
            node_type="Varsayım",
        ))

    affected_documents: dict[str, dict[str, Any]] = {}
    for proposal in proposals:
        if proposal.document_name not in affected_documents:
            affected_documents[proposal.document_name] = {
                "document_name": proposal.document_name,
                "proposal_count": 0,
                "categories": [],
                "sections": [],
            }
        row = affected_documents[proposal.document_name]
        row["proposal_count"] += 1
        if proposal.category not in row["categories"]:
            row["categories"].append(proposal.category)
        if proposal.section not in row["sections"]:
            row["sections"].append(proposal.section)

    project_id = _clean(report.get("project_id")) or _safe(report.get("project_name"))
    package = ChangePackage(
        change_id=change_id,
        project_id=project_id,
        project_name=_clean(report.get("project_name")) or "Proje",
        created_at=created.isoformat(timespec="seconds"),
        change_request=request.to_dict(),
        request_summary=(
            f"{request.requirement_id or selected_id or 'Yeni gereksinim'} · "
            f"{request.change_type} · {_clean(request.reason)}"
        ),
        selected_item=deepcopy(result.selected_item),
        proposals=proposals,
        affected_documents=sorted(
            affected_documents.values(), key=lambda item: item["document_name"]
        ),
        engineering_ideas=[
            item.to_dict() for item in result.engineering_suggestions
        ],
        risks=risk_rows,
        open_questions=open_questions,
        assumptions=assumptions,
        simulation_snapshot=result.to_dict(),
        baseline_traceability=report,
        baseline_fingerprint=_digest({
            "nodes": report.get("nodes", []), "edges": report.get("edges", [])
        }),
    )
    return package


def package_decision_digest(package: ChangePackage) -> str:
    return _digest({
        "change_id": package.change_id,
        "proposals": [
            {
                "id": item.proposal_id, "decision": item.decision,
                "proposed": item.proposed_text, "current": item.current_text,
            }
            for item in package.proposals
        ],
    })


def update_proposal(
    package: ChangePackage,
    proposal_id: str,
    decision: str,
    *,
    proposed_text: str | None = None,
) -> UpdateProposal:
    """Bir kullanıcı kararını uygular ve önceki toplu onayı geçersiz kılar."""
    if decision not in DECISIONS:
        raise ChangePackageError("Geçersiz kullanıcı kararı.")
    proposal = next(
        (item for item in package.proposals if item.proposal_id == proposal_id), None
    )
    if proposal is None:
        raise ChangePackageError("Güncellenecek öneri bulunamadı.")
    if decision == DECISION_EDIT:
        if proposed_text is None or not _clean(proposed_text):
            raise ChangePackageError("Düzenlenen öneri metni boş bırakılamaz.")
        proposal.proposed_text = str(proposed_text).strip()
        proposal.edited = True
    proposal.decision = decision
    package.approval_confirmed = False
    package.approval_actor = ""
    package.approval_at = ""
    package.approval_digest = ""
    package.status = "Kullanıcı incelemesinde"
    return proposal


def mark_explicit_approval(
    package: ChangePackage, actor: str, *, now: datetime | None = None
) -> ChangePackage:
    accepted = [item for item in package.proposals if item.decision == DECISION_ACCEPT]
    if not accepted:
        raise ChangePackageError(
            "Uygulama onayı için en az bir öneri 'Kabul et' olarak işaretlenmelidir."
        )
    actor = _clean(actor)
    if not actor:
        raise ChangePackageError("Onayı veren kişi/rol boş bırakılamaz.")
    for item in accepted:
        if item.target_kind in {"structured_record", "hardware_record", "new_test", "new_requirement"}:
            if item.target_kind != "new_test" and not item.requirement_id:
                raise ChangePackageError("Onaylanan belge önerisinde gerçek öğe kimliği eksik.")
            if item.target_kind != "structured_record" or (
                package.change_request.get("change_type")
                != simulation.CHANGE_REQUIREMENT_REMOVE
                or item.requirement_id
                != _clean((package.selected_item or {}).get("id"))
            ):
                if not _clean(item.proposed_text):
                    raise ChangePackageError(
                        f"{item.requirement_id}: Onaylanan öneri metni boş olamaz."
                    )
    package.approval_confirmed = True
    package.approval_actor = actor
    package.approval_at = _now(now).isoformat(timespec="seconds")
    package.approval_digest = package_decision_digest(package)
    package.status = "Uygulama onayı verildi"
    return package


def project_directory(
    package: ChangePackage,
    output_root: str | os.PathLike[str] | None = None,
) -> Path:
    root = Path(output_root) if output_root else DEFAULT_OUTPUT_ROOT
    return root / _safe(package.project_id)


def save_change_package(
    package: ChangePackage,
    output_root: str | os.PathLike[str] | None = None,
) -> Path:
    path = project_directory(package, output_root) / "drafts" / f"{package.change_id}.json"
    package.storage_path = str(path.resolve())
    _atomic_json(path, package.to_dict())
    return path


def load_change_package(path: str | os.PathLike[str]) -> ChangePackage:
    try:
        value = json.loads(Path(path).read_text(encoding="utf-8"))
    except (OSError, json.JSONDecodeError) as error:
        raise ChangePackageError(f"Değişiklik paketi okunamadı: {error}") from error
    if not isinstance(value, Mapping):
        raise ChangePackageError("Değişiklik paketi geçerli değil.")
    return ChangePackage.from_mapping(value)


def _current_state(project_dir: Path) -> dict[str, Any]:
    current = project_dir / "current.json"
    if not current.exists():
        return {"version": 1}
    try:
        value = json.loads(current.read_text(encoding="utf-8"))
        if not isinstance(value, Mapping):
            raise ValueError("güncel sürüm kaydı sözlük değil")
        result = dict(value)
        result["version"] = max(1, int(result.get("version", 1)))
        return result
    except Exception as error:
        raise ChangePackageError(f"Güncel sürüm bilgisi okunamadı: {error}") from error


def _validate_baseline(
    proposal: UpdateProposal,
    flat_data: Mapping[str, Mapping[str, Any]],
    hardware_data: Mapping[str, Mapping[str, Any]],
) -> None:
    if proposal.target_kind == "structured_record":
        record = flat_data.get(proposal.requirement_id)
        if not isinstance(record, Mapping):
            raise ChangePackageError(
                f"{proposal.requirement_id}: Güncellenecek yapılandırılmış kayıt bulunamadı."
            )
        if _clean(record.get("content")) != _clean(proposal.current_text):
            raise ChangePackageError(
                f"{proposal.requirement_id}: Belge içeriği öneri oluşturulduktan sonra değişmiş. "
                "İzlenebilirliği yeniden tarayıp yeni paket oluşturun."
            )
    elif proposal.target_kind == "hardware_record":
        record = hardware_data.get(proposal.requirement_id)
        if not isinstance(record, Mapping):
            raise ChangePackageError(
                f"{proposal.requirement_id}: Güncellenecek donanım kaydı bulunamadı."
            )
        current = record.get("description") or record.get("name")
        if _clean(current) != _clean(proposal.current_text):
            raise ChangePackageError(
                f"{proposal.requirement_id}: Donanım kaydı öneriden sonra değişmiş."
            )
    elif proposal.target_kind == "new_requirement":
        if proposal.requirement_id in flat_data:
            raise ChangePackageError(
                f"{proposal.requirement_id}: Yeni gereksinim kimliği mevcut kayıtlarda zaten kullanılıyor."
            )
        if proposal.document_type not in {"TID", "SGD", "STT"}:
            raise ChangePackageError(
                f"{proposal.requirement_id}: Yeni gereksinimin V-Model belge türü belirlenemedi."
            )


def _next_test_id(flat_data: Mapping[str, Any], document_type: str) -> str:
    prefix = {"KMTD": "AT", "SITET": "SITET", "AST": "SST"}[document_type]
    highest = 0
    pattern = re.compile(rf"^{re.escape(prefix)}-(\d+)$", re.IGNORECASE)
    for key, record in flat_data.items():
        identifier = _clean(record.get("ID") if isinstance(record, Mapping) else key)
        match = pattern.match(identifier)
        if match:
            highest = max(highest, int(match.group(1)))
    return f"{prefix}-{highest + 1:03d}"


def _new_test_type(package: ChangePackage) -> str:
    document_type = _clean((package.selected_item or {}).get("document_type"))
    return {"TID": "KMTD", "SGD": "SITET", "STT": "AST"}.get(
        document_type, "SITET"
    )


def _register_fonts() -> tuple[str, str]:
    try:
        from etki_analizi_raporlama import _register_pdf_fonts
        return _register_pdf_fonts()
    except Exception:
        return "Helvetica", "Helvetica-Bold"


def _write_version_pdf(
    path: Path, project_name: str, document_type: str,
    records: Sequence[Mapping[str, Any]], version: int,
    change_id: str, created_at: datetime,
) -> None:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import mm
    from reportlab.platypus import LongTable, Paragraph, SimpleDocTemplate, Spacer, TableStyle

    regular, bold = _register_fonts()
    title = DOCUMENT_TYPE_DEFINITIONS.get(document_type, {}).get(
        "document_title", document_type
    )
    title_style = ParagraphStyle(
        "VersionTitle", fontName=bold, fontSize=16, leading=20,
        textColor=colors.HexColor("#17365D"),
    )
    body = ParagraphStyle(
        "VersionBody", fontName=regular, fontSize=8, leading=10,
        textColor=colors.HexColor("#222222"),
    )
    head = ParagraphStyle(
        "VersionHead", fontName=bold, fontSize=8, leading=10,
        textColor=colors.white,
    )
    story = [
        Paragraph(escape(title), title_style),
        Paragraph(
            escape(
                f"Proje: {project_name} | Sürüm: v{version:04d} | "
                f"Değişiklik: {change_id} | Tarih: {created_at.strftime('%d.%m.%Y %H:%M')}"
            ),
            body,
        ),
        Spacer(1, 5 * mm),
    ]
    rows = [[
        Paragraph("Kimlik", head), Paragraph("Açıklama", head),
        Paragraph("Bağlı öğe", head), Paragraph("Değişiklik kaydı", head),
    ]]
    for record in records:
        rows.append([
            Paragraph(escape(_clean(record.get("ID"))), body),
            Paragraph(escape(_clean(record.get("content"))), body),
            Paragraph(escape(_clean(record.get("bound_to")) or "-"), body),
            Paragraph(
                escape(_clean(record.get("last_change_id")) or "Değişmedi"), body
            ),
        ])
    table = LongTable(
        rows, colWidths=[27 * mm, 105 * mm, 30 * mm, 30 * mm], repeatRows=1
    )
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#17365D")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [
            colors.white, colors.HexColor("#EEF1F4")
        ]),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8DEE5")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(table)
    document = SimpleDocTemplate(
        str(path), pagesize=A4, leftMargin=10 * mm, rightMargin=10 * mm,
        topMargin=13 * mm, bottomMargin=13 * mm,
        title=f"{title} v{version:04d}", author="Etki Analizi Uygulaması",
    )
    document.build(story)


def _write_version_docx(
    path: Path, project_name: str, document_type: str,
    records: Sequence[Mapping[str, Any]], version: int,
    change_id: str, created_at: datetime,
) -> None:
    from docx import Document
    from docx.shared import Pt

    title = DOCUMENT_TYPE_DEFINITIONS.get(document_type, {}).get(
        "document_title", document_type
    )
    document = Document()
    document.add_heading(title, level=0)
    document.add_paragraph(f"Proje: {project_name}")
    document.add_paragraph(
        f"Sürüm: v{version:04d} | Değişiklik isteği: {change_id} | "
        f"Tarih: {created_at.strftime('%d.%m.%Y %H:%M')}"
    )
    table = document.add_table(rows=1, cols=4)
    try:
        table.style = "Light Grid Accent 1"
    except Exception:
        table.style = "Table Grid"
    for cell, value in zip(
        table.rows[0].cells, ("Kimlik", "Açıklama", "Bağlı öğe", "Değişiklik kaydı")
    ):
        cell.text = value
    for record in records:
        cells = table.add_row().cells
        cells[0].text = _clean(record.get("ID"))
        cells[1].text = _clean(record.get("content"))
        cells[2].text = _clean(record.get("bound_to")) or "-"
        cells[3].text = _clean(record.get("last_change_id")) or "Değişmedi"
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Segoe UI"
            run.font.size = Pt(9)
    document.core_properties.title = f"{title} v{version:04d}"
    document.core_properties.subject = f"{change_id} değişiklik kaydı"
    document.save(path)


def _verify_version_documents(paths: Sequence[Path]) -> None:
    from docx import Document
    from pypdf import PdfReader

    for path in paths:
        if not path.exists() or path.stat().st_size == 0:
            raise ChangePackageError(f"Yeni sürüm belgesi oluşturulamadı: {path.name}")
        if path.suffix.lower() == ".pdf":
            reader = PdfReader(str(path))
            if not reader.pages:
                raise ChangePackageError(f"PDF belgesi sayfa içermiyor: {path.name}")
        elif path.suffix.lower() == ".docx":
            document = Document(str(path))
            if not document.paragraphs and not document.tables:
                raise ChangePackageError(f"Word belgesi içerik taşımıyor: {path.name}")


def compare_closure(
    package: ChangePackage,
    post_traceability: Mapping[str, Any],
    post_simulation: Mapping[str, Any] | None,
) -> dict[str, Any]:
    accepted = [
        item for item in package.proposals if item.decision == DECISION_ACCEPT
    ]
    resolved = [
        item.proposal_id for item in accepted
        if item.target_kind in {"structured_record", "hardware_record", "new_test", "new_requirement"}
    ]
    continuing = [
        item.proposal_id for item in package.proposals
        if item.decision in {DECISION_REJECT, DECISION_DEFER, DECISION_EDIT}
        and item.target_kind != "governance_action"
    ]
    before = package.baseline_traceability
    before_conflicts = {
        _digest(item) for item in before.get("conflicts", [])
    }
    after_conflicts = {
        _digest(item) for item in post_traceability.get("conflicts", [])
    }
    return {
        "status": "Kapanış kontrolü tamamlandı",
        "resolved_proposal_ids": resolved,
        "continuing_proposal_ids": continuing,
        "resolved_count": len(resolved),
        "continuing_count": len(continuing),
        "before": {
            "unlinked": list(before.get("unlinked_requirements", [])),
            "unverified": list(before.get("unverified_requirements", [])),
            "conflict_count": len(before.get("conflicts", [])),
        },
        "after": {
            "unlinked": list(post_traceability.get("unlinked_requirements", [])),
            "unverified": list(post_traceability.get("unverified_requirements", [])),
            "conflict_count": len(post_traceability.get("conflicts", [])),
        },
        "new_conflict_count": len(after_conflicts - before_conflicts),
        "closed_conflict_count": len(before_conflicts - after_conflicts),
        "rerun_status": _clean((post_simulation or {}).get("status")) or "Çalıştırılamadı",
        "rerun_impact_count": int(
            ((post_simulation or {}).get("summary") or {}).get("impact_count", 0) or 0
        ),
        "generated_at": _now().isoformat(timespec="seconds"),
    }


def apply_approved_changes(
    package: ChangePackage,
    flat_data: Mapping[str, Mapping[str, Any]],
    *,
    hardware_data: Mapping[str, Mapping[str, Any]] | None = None,
    source_paths: Sequence[str | os.PathLike[str]] | None = None,
    output_root: str | os.PathLike[str] | None = None,
    validator: Callable[
        [dict[str, dict[str, Any]], dict[str, dict[str, Any]], Path],
        Mapping[str, Any],
    ] | None = None,
    now: datetime | None = None,
) -> ApplicationResult:
    """Onaylı paketi geçici alanda üretip doğruladıktan sonra yeni sürüm olarak yayımlar."""
    with _LOCK:
        if not package.approval_confirmed:
            raise ChangePackageError(
                "Onaylanan değişiklikleri uygulamak için açık kullanıcı onayı gereklidir."
            )
        if package.approval_digest != package_decision_digest(package):
            raise ChangePackageError(
                "Paket onaydan sonra değişmiş. Kullanıcı onayını yeniden verin."
            )
        accepted = [
            item for item in package.proposals if item.decision == DECISION_ACCEPT
        ]
        if not accepted:
            raise ChangePackageError("Uygulanacak kabul edilmiş öneri bulunamadı.")
        for proposal in accepted:
            _validate_baseline(proposal, flat_data, hardware_data or {})

        project_dir = project_directory(package, output_root)
        project_dir.mkdir(parents=True, exist_ok=True)
        current_state = _current_state(project_dir)
        previous_version = int(current_state["version"])
        baseline_data_fingerprint = _digest({
            "flat_data": flat_data,
            "hardware_data": hardware_data or {},
        })
        expected_fingerprint = _clean(current_state.get("data_fingerprint"))
        if expected_fingerprint and expected_fingerprint != baseline_data_fingerprint:
            raise ChangePackageError(
                "Uygulamadaki belge verisi son yayımlanan sürümle uyuşmuyor. "
                "Önce güncel sürümü yükleyin veya izlenebilirliği yeniden tarayın; "
                "önceki değişikliklerin geri alınmasını önlemek için işlem durduruldu."
            )
        new_version = previous_version + 1
        final_dir = project_dir / "versions" / f"v{new_version:04d}"
        if final_dir.exists():
            raise ChangePackageError(
                f"Yeni sürüm klasörü zaten var: {final_dir}. İşlem güvenlik nedeniyle durduruldu."
            )
        created_at = _now(now)
        stage = Path(tempfile.mkdtemp(prefix=f".{package.change_id}.", dir=project_dir))
        try:
            backup_dir = stage / "baseline_before"
            documents_dir = stage / "documents"
            reports_dir = stage / "reports"
            backup_dir.mkdir(parents=True)
            documents_dir.mkdir(parents=True)
            reports_dir.mkdir(parents=True)
            _atomic_json(backup_dir / f"structured_data.v{previous_version:04d}.json", {
                "project_id": package.project_id,
                "version": previous_version,
                "captured_at": created_at,
                "flat_data": flat_data,
                "hardware_data": hardware_data or {},
            })
            _atomic_json(backup_dir / "traceability.before.json", package.baseline_traceability)
            copied_sources = []
            source_backup_dir = backup_dir / "source_files"
            for index, raw in enumerate(source_paths or (), start=1):
                source = Path(raw)
                if not source.is_file():
                    continue
                source_backup_dir.mkdir(parents=True, exist_ok=True)
                destination = source_backup_dir / f"{index:02d}-{source.name}"
                shutil.copy2(source, destination)
                copied_sources.append(str(destination.relative_to(stage)))

            new_flat = {
                str(key): deepcopy(dict(value))
                for key, value in flat_data.items() if isinstance(value, Mapping)
            }
            new_hardware = {
                str(key): deepcopy(dict(value))
                for key, value in (hardware_data or {}).items()
                if isinstance(value, Mapping)
            }
            modified: list[str] = []
            added: list[str] = []
            non_document_actions: list[str] = []
            affected_types: set[str] = set()
            removal_target = _clean((package.selected_item or {}).get("id"))
            for proposal in accepted:
                if proposal.target_kind == "structured_record":
                    record = new_flat[proposal.requirement_id]
                    if (
                        package.change_request.get("change_type")
                        == simulation.CHANGE_REQUIREMENT_REMOVE
                        and proposal.requirement_id == removal_target
                        and not _clean(proposal.proposed_text)
                    ):
                        del new_flat[proposal.requirement_id]
                    else:
                        if _clean(record.get("content")) == _clean(proposal.proposed_text):
                            raise ChangePackageError(
                                f"{proposal.requirement_id}: Onaylanan değişiklik mevcut içerikle aynı."
                            )
                        record["content"] = proposal.proposed_text.strip()
                        record["version"] = f"v{new_version:04d}"
                        record["last_change_id"] = package.change_id
                        record["updated_at"] = created_at.isoformat(timespec="seconds")
                    modified.append(proposal.requirement_id)
                    if proposal.document_type:
                        affected_types.add(proposal.document_type)
                elif proposal.target_kind == "hardware_record":
                    record = new_hardware[proposal.requirement_id]
                    record["description"] = proposal.proposed_text.strip()
                    record["version"] = f"v{new_version:04d}"
                    record["last_change_id"] = package.change_id
                    modified.append(proposal.requirement_id)
                elif proposal.target_kind == "new_test":
                    document_type = _new_test_type(package)
                    identifier = _next_test_id(new_flat, document_type)
                    new_flat[identifier] = {
                        "type": document_type,
                        "ID": identifier,
                        "content": proposal.proposed_text.strip(),
                        "bound_to": removal_target or "Yok",
                        "version": f"v{new_version:04d}",
                        "last_change_id": package.change_id,
                        "updated_at": created_at.isoformat(timespec="seconds"),
                    }
                    proposal.assigned_id = identifier
                    added.append(identifier)
                    affected_types.add(document_type)
                elif proposal.target_kind == "new_requirement":
                    identifier = proposal.requirement_id
                    if identifier in new_flat:
                        raise ChangePackageError(
                            f"{identifier}: Yeni gereksinim kimliği zaten kullanılıyor."
                        )
                    new_flat[identifier] = {
                        "type": proposal.document_type,
                        "ID": identifier,
                        "content": proposal.proposed_text.strip(),
                        "bound_to": "Yok",
                        "version": f"v{new_version:04d}",
                        "last_change_id": package.change_id,
                        "updated_at": created_at.isoformat(timespec="seconds"),
                    }
                    added.append(identifier)
                    affected_types.add(proposal.document_type)
                else:
                    non_document_actions.append(proposal.proposal_id)

            _atomic_json(stage / "structured_data.json", {
                "project_id": package.project_id,
                "project_name": package.project_name,
                "version": new_version,
                "change_id": package.change_id,
                "created_at": created_at,
                "flat_data": new_flat,
                "hardware_data": new_hardware,
            })
            _atomic_json(stage / "change_package.json", package.to_dict())
            created_documents: list[Path] = []
            for document_type in sorted(affected_types):
                if document_type not in RECORD_DOCUMENT_TYPES:
                    continue
                records = [
                    value for value in new_flat.values()
                    if _clean(value.get("type")) == document_type
                ]
                base_name = (
                    f"{document_type}_v{new_version:04d}_{_safe(package.change_id)}"
                )
                pdf_path = documents_dir / f"{base_name}.pdf"
                docx_path = documents_dir / f"{base_name}.docx"
                _write_version_pdf(
                    pdf_path, package.project_name, document_type, records,
                    new_version, package.change_id, created_at,
                )
                _write_version_docx(
                    docx_path, package.project_name, document_type, records,
                    new_version, package.change_id, created_at,
                )
                created_documents.extend((pdf_path, docx_path))
            if new_hardware != dict(hardware_data or {}):
                hardware_records = []
                for identifier, record in new_hardware.items():
                    linked = record.get("linked_requirements") or record.get("requirement_ids") or []
                    if isinstance(linked, str):
                        linked_text = linked
                    else:
                        linked_text = ", ".join(str(item) for item in linked)
                    hardware_records.append({
                        "ID": record.get("ID") or identifier,
                        "content": record.get("description") or record.get("name"),
                        "bound_to": linked_text or "Yok",
                        "last_change_id": record.get("last_change_id"),
                    })
                hardware_base = (
                    f"HARDWARE_v{new_version:04d}_{_safe(package.change_id)}"
                )
                hardware_pdf = documents_dir / f"{hardware_base}.pdf"
                hardware_docx = documents_dir / f"{hardware_base}.docx"
                _write_version_pdf(
                    hardware_pdf, package.project_name, "HARDWARE",
                    hardware_records, new_version, package.change_id, created_at,
                )
                _write_version_docx(
                    hardware_docx, package.project_name, "HARDWARE",
                    hardware_records, new_version, package.change_id, created_at,
                )
                hardware_path = documents_dir / (
                    f"HARDWARE_v{new_version:04d}_{_safe(package.change_id)}.json"
                )
                _atomic_json(hardware_path, {"hardware_data": new_hardware})
                created_documents.extend((hardware_pdf, hardware_docx, hardware_path))
            _verify_version_documents(created_documents)

            validation = dict(validator(new_flat, new_hardware, stage) if validator else {})
            post_traceability = dict(validation.get("post_traceability") or {})
            post_simulation = dict(validation.get("post_simulation") or {})
            closure = dict(validation.get("closure_summary") or {})
            report_paths = {
                str(key): str(value)
                for key, value in dict(validation.get("report_paths") or {}).items()
            }
            warnings = [_clean(item) for item in validation.get("warnings", []) if _clean(item)]
            if validator and not post_traceability:
                raise ChangePackageError(
                    "Son kontrol izlenebilirlik haritası üretmedi; yeni sürüm yayımlanmadı."
                )
            if post_traceability:
                _atomic_json(stage / "traceability.after.json", post_traceability)
            if post_simulation:
                _atomic_json(stage / "simulation.after.json", post_simulation)
            if closure:
                _atomic_json(stage / "closure_summary.json", closure)

            change_record = {
                "change_id": package.change_id,
                "project_id": package.project_id,
                "previous_version": previous_version,
                "new_version": new_version,
                "created_at": created_at,
                "approved_by": package.approval_actor,
                "approved_at": package.approval_at,
                "modified_item_ids": modified,
                "added_item_ids": added,
                "non_document_actions": non_document_actions,
                "copied_source_backups": copied_sources,
                "before_data_fingerprint": baseline_data_fingerprint,
                "after_data_fingerprint": _digest({
                    "flat_data": new_flat, "hardware_data": new_hardware,
                }),
                "decisions": [
                    {
                        "proposal_id": item.proposal_id,
                        "item_id": item.requirement_id,
                        "decision": item.decision,
                        "reason": item.rationale,
                        "before": item.current_text,
                        "after": item.proposed_text,
                    }
                    for item in package.proposals
                ],
                "closure": closure,
                "warnings": warnings,
            }
            _atomic_json(stage / "change_record.json", change_record)
            _atomic_json(stage / "COMMIT.json", {
                "status": "validated",
                "change_id": package.change_id,
                "version": new_version,
                "validated_at": _now().isoformat(timespec="seconds"),
            })
            final_dir.parent.mkdir(parents=True, exist_ok=True)
            os.replace(stage, final_dir)

            ledger_path = project_dir / "change_ledger.json"
            ledger = {"schema_version": "1.0", "project_id": package.project_id, "entries": []}
            if ledger_path.exists():
                try:
                    loaded = json.loads(ledger_path.read_text(encoding="utf-8"))
                    if isinstance(loaded, Mapping):
                        ledger.update(loaded)
                        ledger["entries"] = list(ledger.get("entries") or [])
                except Exception as error:
                    raise ChangePackageError(
                        f"Yeni sürüm hazırlandı ancak değişiklik defteri okunamadı: {error}"
                    ) from error
            ledger["entries"].append({
                "change_id": package.change_id,
                "version": new_version,
                "previous_version": previous_version,
                "approved_by": package.approval_actor,
                "approved_at": package.approval_at,
                "committed_at": _now().isoformat(timespec="seconds"),
                "path": str(final_dir.resolve()),
            })
            _atomic_json(ledger_path, ledger)
            _atomic_json(project_dir / "current.json", {
                "project_id": package.project_id,
                "version": new_version,
                "change_id": package.change_id,
                "version_directory": str(final_dir.resolve()),
                "data_fingerprint": _digest({
                    "flat_data": new_flat, "hardware_data": new_hardware,
                }),
                "updated_at": _now().isoformat(timespec="seconds"),
            })
            package.status = "Uygulandı ve doğrulandı"
            package.storage_path = str(
                (final_dir / "change_package.json").resolve()
            )
            final_report_paths = {
                key: str((final_dir / Path(value).relative_to(stage)).resolve())
                if Path(value).is_absolute() and str(value).startswith(str(stage))
                else str(value)
                for key, value in report_paths.items()
            }
            return ApplicationResult(
                change_id=package.change_id,
                project_id=package.project_id,
                previous_version=previous_version,
                new_version=new_version,
                status="Uygulandı ve doğrulandı",
                version_directory=str(final_dir.resolve()),
                backup_directory=str((final_dir / "baseline_before").resolve()),
                created_documents=[
                    str((final_dir / item.relative_to(stage)).resolve())
                    for item in created_documents
                ],
                modified_item_ids=modified,
                added_item_ids=added,
                non_document_actions=non_document_actions,
                new_flat_data=new_flat,
                new_hardware_data=new_hardware,
                post_traceability=post_traceability,
                post_simulation=post_simulation,
                closure_summary=closure,
                report_paths=final_report_paths,
                warnings=warnings,
            )
        except Exception:
            if stage.exists():
                shutil.rmtree(stage, ignore_errors=True)
            raise


__all__ = [
    "ApplicationResult", "CATEGORY_ACCEPTANCE", "CATEGORY_ASSUMPTION",
    "CATEGORY_CUSTOMER", "CATEGORY_DESIGN", "CATEGORY_MAIN", "CATEGORY_NEW_TEST",
    "CATEGORY_QUESTION", "CATEGORY_RISK", "CATEGORY_SUBSYSTEM", "CATEGORY_SYSTEM",
    "CATEGORY_TEST", "CATEGORY_VERIFICATION", "ChangePackage", "ChangePackageError",
    "DECISION_ACCEPT", "DECISION_DEFER", "DECISION_EDIT", "DECISION_REJECT",
    "DECISIONS", "UpdateProposal", "apply_approved_changes", "build_change_package",
    "compare_closure", "load_change_package", "mark_explicit_approval",
    "package_decision_digest", "project_directory", "save_change_package",
    "update_proposal",
]
