# -*- coding: utf-8 -*-
"""Faz 7 (mimari yeniden yapılandırma) — Arayüz.py'nin bölünmüş
parçalarından biri. Bkz. MIMARI_YENIDEN_YAPILANDIRMA_PLANI.md bölüm 3.
"""

import csv
import os
import sys

# Windows konsolu (cp1254) emoji/Unicode karakterleri basamadığı için
# çıktıyı UTF-8'e zorla; aksi halde print(...) ifadeleri çökertir.
for _stream in (sys.stdout, sys.stderr):
    try:
        _stream.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass

import time
import threading
import traceback
from datetime import datetime
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from PIL import Image, ImageTk
import ttkbootstrap as ttk
from ttkbootstrap.constants import *
from ttkbootstrap.style import Style
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase.ttfonts import TTFont 
from reportlab.pdfbase import pdfmetrics
from openpyxl import Workbook
import numpy as np
from langchain_huggingface import HuggingFaceEmbeddings
from app_identity import (
    APP_NAME, ICON_RELATIVE_PATH, apply_app_identity,
    prepare_process_identity, resource_path,
)

sys.path.append(os.path.dirname(os.path.abspath(__file__)))

try:
    import tid_generator_logic
    import sgd_generator_logic
    import stt_generator_logic
    import dgöygö_generator_logic
    import kmtd_generator_logic
    import sitet_generator_logic
    import alt_sistem_test_logic
    import dtet_ytet_generator_logic
    import hardware_list_logic
    import hardware_generator_logic
    import hardware_list_ui
    import donanim_kartlari_gorsel
    import donanim_kartlari_ui
    import donanim_kartlari_yonetim
    import etki_analizi_ui
    import etki_analizi_izlenebilirlik
    import etki_analizi_entegrasyon
    import etki_analizi_simulasyon
    import etki_analizi_degisim_paketi
    import etki_analizi_degisim_raporlama
    import donanim_kartlari_algilama
    import mimari_cerceve_ui
    import text_cleanup
    import html_generation
    import pdf_extraction
except ImportError as e:
    messagebox.showerror(
        "Modül Hatası",
        f"Gerekli bir modül yüklenemedi: {e}\nLütfen programı yeniden kurun veya bağımlılıkları kontrol edin."
    )
    sys.exit(1)

from .yardimcilar import pre_process_files, start1_time

class _DisaAktarimMixin:
    def _register_pdf_font(self):
        """Türkçe karakterleri destekleyen bir font ailesi kaydeder; (normal, kalın) döner."""
        candidates = [
            ("VModelFont", "VModelFont-Bold",
             r"C:\Windows\Fonts\arial.ttf", r"C:\Windows\Fonts\arialbd.ttf"),
        ]
        for reg_name, bold_name, reg_path, bold_path in candidates:
            try:
                if os.path.exists(reg_path):
                    pdfmetrics.registerFont(TTFont(reg_name, reg_path))
                    if os.path.exists(bold_path):
                        pdfmetrics.registerFont(TTFont(bold_name, bold_path))
                    else:
                        bold_name = reg_name
                    return reg_name, bold_name
            except Exception:
                continue
        # Yedek: reportlab ile gelen Bitstream Vera (Türkçe destekli)
        try:
            import reportlab
            fonts_dir = os.path.join(os.path.dirname(reportlab.__file__), "fonts")
            pdfmetrics.registerFont(TTFont("VModelFont", os.path.join(fonts_dir, "Vera.ttf")))
            pdfmetrics.registerFont(TTFont("VModelFont-Bold", os.path.join(fonts_dir, "VeraBd.ttf")))
            return "VModelFont", "VModelFont-Bold"
        except Exception:
            return "Helvetica", "Helvetica-Bold"

    def _generate_vmodel_pdf(self, path, proje_ismi):
        """flat_data'yı kullanarak V-Model yapısında biçimli bir PDF üretir."""
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib import colors
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
        )

        reg, bold = self._register_pdf_font()

        title_style = ParagraphStyle("vm_title", fontName=bold, fontSize=20,
                                     textColor=colors.HexColor("#1F3864"), leading=24, spaceAfter=4)
        sub_style = ParagraphStyle("vm_sub", fontName=reg, fontSize=11,
                                   textColor=colors.HexColor("#555555"), leading=15)
        legend_style = ParagraphStyle("vm_legend", fontName=reg, fontSize=9,
                                      textColor=colors.HexColor("#555555"), leading=13)
        sec_style = ParagraphStyle("vm_sec", fontName=bold, fontSize=13,
                                   textColor=colors.white, leading=16)
        cell_style = ParagraphStyle("vm_cell", fontName=reg, fontSize=9, leading=12)
        cellid_style = ParagraphStyle("vm_cellid", fontName=bold, fontSize=9, leading=12,
                                      textColor=colors.HexColor("#1F3864"))
        head_style = ParagraphStyle("vm_head", fontName=bold, fontSize=9, leading=12,
                                    textColor=colors.white)

        # Renkler: sol bacak (gereksinim) mavi, sağ bacak (test) yeşil tonları
        REQ_COLOR = colors.HexColor("#1F3864")
        REQ_LIGHT = colors.HexColor("#D6E0F0")
        TEST_COLOR = colors.HexColor("#375623")
        TEST_LIGHT = colors.HexColor("#E2EFDA")

        story = []
        story.append(Paragraph("V-MODEL TEKNİK DOKÜMAN PAKETİ", title_style))
        story.append(Paragraph(f"Proje: <b>{proje_ismi}</b>", sub_style))
        story.append(Paragraph(f"Oluşturulma: {datetime.now().strftime('%d.%m.%Y %H:%M')}", sub_style))
        story.append(Spacer(1, 4))
        story.append(HRFlowable(width="100%", thickness=1.5, color=colors.HexColor("#1F3864")))
        story.append(Spacer(1, 6))
        story.append(Paragraph(
            "<font color='#1F3864'>■</font> Sol Bacak: Gereksinim Dokümanları &nbsp;&nbsp;&nbsp; "
            "<font color='#375623'>■</font> Sağ Bacak: Test / Doğrulama Dokümanları", legend_style))
        story.append(Spacer(1, 12))

        # Sayfa genişliği ~ 170mm kullanılabilir
        col_widths = [26 * mm, 118 * mm, 26 * mm]
        any_section = False

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_section = True

            is_req = (leg == "req")
            head_color = REQ_COLOR if is_req else TEST_COLOR
            light_color = REQ_LIGHT if is_req else TEST_LIGHT
            bound_header = "Kaynak / Bağlı" if is_req else "Doğruladığı Madde"

            # Bölüm başlığı bandı
            sec_band = Table([[Paragraph(f"{title}  ({len(items)} madde)", sec_style)]],
                             colWidths=[sum(col_widths)])
            sec_band.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), head_color),
                ("LEFTPADDING", (0, 0), (-1, -1), 8),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
            ]))
            story.append(sec_band)

            # Tablo verisi
            data = [[Paragraph("ID", head_style),
                     Paragraph("Açıklama", head_style),
                     Paragraph(bound_header, head_style)]]
            for d in items:
                data.append([
                    Paragraph(str(d.get("ID", "")), cellid_style),
                    Paragraph(str(d.get("content", "")), cell_style),
                    Paragraph(str(d.get("bound_to", "-")), cell_style),
                ])

            tbl = Table(data, colWidths=col_widths, repeatRows=1)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), head_color),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, light_color]),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#BBBBBB")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)
            story.append(Spacer(1, 16))

        if not any_section:
            # flat_data boşsa ham metni sığdır
            story.append(Paragraph("Üretilen doküman verisi bulunamadı. Ham çıktı:", sub_style))
            story.append(Spacer(1, 6))
            for line in (self.last_generated_output or "").split("\n"):
                if line.strip():
                    story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;"), cell_style))

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=20 * mm, rightMargin=20 * mm,
            topMargin=18 * mm, bottomMargin=18 * mm,
            title=f"{proje_ismi} - V-Model Doküman Paketi"
        )
        doc.build(story)

    # ================================================================== #
    #  DOKÜMAN COPILOT — sağ panel, tek madde revizyonu (grounding'li)    #
    # ================================================================== #

    # --- Kıdemli Sistem Mühendisi system prompt (paydaş kuralları 1-4) ---
    COPILOT_SYSTEM_PROMPT = (
        "ROL: Sen kıdemli bir Sistem Mühendisisin (INCOSE / sistem mühendisliği "
        "standartlarına hâkim). Görevin, sana verilen TEK bir gereksinim veya test "
        "maddesini kullanıcının isteğine göre yeniden yazmaktır. Aşağıdaki kurallara "
        "İSTİSNASIZ uyarsın:\n"
        "1) SADE ve SAF: Yanıtın YALNIZCA revize edilmiş madde metnidir. 'Tabii ki', "
        "'İşte güncellenmiş hâli' gibi laf kalabalığı, giriş/kapanış cümlesi, açıklama, "
        "yorum, başlık, ID, numara, madde işareti, yıldız veya tırnak KESİNLİKLE YOK. "
        "Tek paragraf, en fazla 2-3 cümle.\n"
        "2) BİRİMLER: Tüm sayısal değerleri metrik/teknik birimlerle ver "
        "(ms, s, Hz, kHz, MHz, dB, dBm, m, km, %, °C). Birimsiz çıplak sayı bırakma.\n"
        "3) TEKNİK DİL: Günlük dil kullanma; sistem mühendisliği terminolojisi kullan. "
        "Cümleler net, ölçülebilir ve doğrulanabilir olmalı ('sistem ... -malıdır/-melidir', "
        "'doğrulanmalıdır', 'karşılanmalıdır').\n"
        "4) DSB KURALI (SAYI UYDURMAK KESİNLİKLE YASAK): Bir sayısal değeri YALNIZCA şu iki "
        "durumda yazabilirsin: (a) kullanıcı o sayıyı isteğinde AÇIKÇA vermişse, ya da "
        "(b) sayı zaten MEVCUT METİN'de varsa. Kullanıcı 'bir değer/süre/eşik/mesafe ekle' "
        "diyor ama sayının kendisini SÖYLEMİYORSA, o sayıyı TAHMİN ETME — makul/tipik bir "
        "değer bile olsa UYDURMA. Bunun yerine değerin geçtiği yere harfi harfine 'DSB' yaz "
        "(DSB = Daha Sonra Belirlenecek), birimi koru. "
        "Örnek: kullanıcı 'tespit süresi ekle' dedi ama süreyi vermediyse → "
        "'Sistem, hedefleri DSB ms içinde tespit etmelidir.'\n"
        "ÖRNEKLER (uy):\n"
        "- MEVCUT: 'Sistem, hedefleri tespit etmelidir.' | İSTEK: 'tespit süresi ekle' "
        "(kullanıcı sayı VERMEDİ) → DOĞRU: 'Sistem, hedefleri DSB ms içinde tespit etmelidir.' | "
        "YANLIŞ: 'Sistem, hedefleri 100 ms içinde tespit etmelidir.' (100 UYDURMA, yasak!).\n"
        "- MEVCUT: 'Sistem çalışmalıdır.' | İSTEK: 'çalışma sıcaklığı -20 ile +55 °C olsun' "
        "(değer VERİLDİ) → DOĞRU: 'Sistem, -20 ile +55 °C sıcaklık aralığında çalışmalıdır.'\n"
        "ÇIKTI: Yalnızca revize edilmiş madde metni. Başka HİÇBİR ŞEY yazma."
    )

    # DSB ripple'ının yayılacağı test tipi maddeler
    TEST_TYPES = ("KMTD", "SITET", "AST")
    # Revize edilince bağlı testleri yeniden üretilecek gereksinim tipleri
    REQ_TYPES = ("TID", "SGD", "STT")

    def _export_excel(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile çok sayfalı, biçimli Excel üretir."""
        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

        thin = Side(style="thin", color="CCCCCC")
        border = Border(left=thin, right=thin, top=thin, bottom=thin)

        def style_header(ws, ncols, color="1F3864"):
            fill = PatternFill("solid", fgColor=color)
            for c in range(1, ncols + 1):
                cell = ws.cell(row=1, column=c)
                cell.fill = fill
                cell.font = Font(color="FFFFFF", bold=True)
                cell.alignment = Alignment(horizontal="center", vertical="center")
                cell.border = border

        wb = Workbook()
        # --- Kapak ---
        info = wb.active
        info.title = "Proje Bilgisi"
        info.append(["V-MODEL İZLENEBİLİRLİK RAPORU"])
        info.append(["Proje", proje_ismi])
        info.append(["Tarih", datetime.now().strftime("%d.%m.%Y %H:%M")])
        info.append(["Toplam Madde", len(self.flat_data)])
        info["A1"].font = Font(bold=True, size=14, color="1F3864")
        info.column_dimensions["A"].width = 20
        info.column_dimensions["B"].width = 50

        leg_tr = {"req": "Gereksinim", "test": "Test"}

        # --- İzlenebilirlik Matrisi ---
        ws = wb.create_sheet("Izlenebilirlik Matrisi")
        ws.append(["ID", "Tür", "Bacak", "Açıklama", "Bağlı Olduğu"])
        for type_key, title, leg in self.VMODEL_SECTIONS:
            for d in self.flat_data.values():
                if d.get("type") == type_key:
                    ws.append([d.get("ID", ""), type_key, leg_tr.get(leg, leg),
                               d.get("content", ""), d.get("bound_to", "Yok")])
        style_header(ws, 5)
        for col, w in zip("ABCDE", [16, 12, 12, 80, 22]):
            ws.column_dimensions[col].width = w
        for r in range(2, ws.max_row + 1):
            ws.cell(row=r, column=4).alignment = Alignment(wrap_text=True, vertical="top")

        # --- Her tür için ayrı sayfa ---
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            safe = type_key.replace("/", "-")[:31]
            sh = wb.create_sheet(safe)
            sh.append([f"{type_key} ID", "Açıklama", "Bağlı Olduğu"])
            for d in items:
                sh.append([d.get("ID", ""), d.get("content", ""), d.get("bound_to", "Yok")])
            style_header(sh, 3, "1F3864" if leg == "req" else "375623")
            for col, w in zip("ABC", [16, 90, 22]):
                sh.column_dimensions[col].width = w
            for r in range(2, sh.max_row + 1):
                sh.cell(row=r, column=2).alignment = Alignment(wrap_text=True, vertical="top")

        wb.save(path)

    def _export_html(self, path, proje_ismi):
        """flat_data + VMODEL_SECTIONS ile bağımsız, biçimli HTML raporu üretir."""
        import html as _html
        esc = lambda s: _html.escape(str(s))
        leg_color = {"req": "#1F3864", "test": "#375623"}
        leg_light = {"req": "#D6E0F0", "test": "#E2EFDA"}

        p = []
        p.append("<!DOCTYPE html><html lang='tr'><head><meta charset='utf-8'>")
        p.append(f"<title>{esc(proje_ismi)} - V-Model Doküman Paketi</title>")
        p.append("<style>"
                 "body{font-family:'Segoe UI',Arial,sans-serif;margin:24px;color:#222;background:#f7f9fc}"
                 "h1{color:#1F3864;margin-bottom:0}.sub{color:#666;margin:2px 0 14px}"
                 "table{border-collapse:collapse;width:100%;margin-bottom:22px;background:#fff;"
                 "box-shadow:0 1px 3px rgba(0,0,0,.08)}"
                 "th,td{border:1px solid #ccc;padding:7px 9px;text-align:left;vertical-align:top;font-size:14px}"
                 ".sec{color:#fff;padding:8px 12px;font-weight:bold;font-size:15px;margin-top:8px;border-radius:3px 3px 0 0}"
                 "td.id{font-weight:bold;white-space:nowrap}"
                 ".legend span{display:inline-block;margin-right:18px}"
                 ".chip{display:inline-block;width:12px;height:12px;border-radius:2px;margin-right:5px;vertical-align:middle}"
                 "</style></head><body>")
        p.append("<h1>V-Model Teknik Doküman Paketi</h1>")
        p.append(f"<div class='sub'>Proje: <b>{esc(proje_ismi)}</b> &nbsp;|&nbsp; "
                 f"{datetime.now().strftime('%d.%m.%Y %H:%M')} &nbsp;|&nbsp; Toplam {len(self.flat_data)} madde</div>")
        p.append("<div class='legend'>"
                 "<span><span class='chip' style='background:#1F3864'></span>Sol Bacak: Gereksinim</span>"
                 "<span><span class='chip' style='background:#375623'></span>Sağ Bacak: Test / Doğrulama</span></div><br>")

        any_sec = False
        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            any_sec = True
            col = leg_color.get(leg, "#333"); light = leg_light.get(leg, "#eee")
            bound_h = "Kaynak / Bağlı" if leg == "req" else "Doğruladığı Madde"
            p.append(f"<div class='sec' style='background:{col}'>{esc(title)} ({len(items)} madde)</div>")
            p.append(f"<table><tr style='background:{col};color:#fff'>"
                     f"<th>ID</th><th>Açıklama</th><th>{bound_h}</th></tr>")
            for i, d in enumerate(items):
                bg = "#ffffff" if i % 2 == 0 else light
                p.append(f"<tr style='background:{bg}'><td class='id'>{esc(d.get('ID',''))}</td>"
                         f"<td>{esc(d.get('content',''))}</td><td>{esc(d.get('bound_to','-'))}</td></tr>")
            p.append("</table>")
        if not any_sec:
            p.append("<p><i>Üretilen doküman verisi bulunamadı.</i></p>")
        p.append("</body></html>")

        with open(path, "w", encoding="utf-8") as f:
            f.write("".join(p))

    def _export_docx(self, path, proje_ismi, template_path=None):
        """Kurumsal .docx şablonunu doldurur (veya sıfırdan Word üretir). Başarı=True."""
        try:
            from docx import Document
        except ImportError:
            messagebox.showerror(
                "Eksik Kütüphane",
                "Word (.docx) çıktısı için 'python-docx' gerekli.\n\n"
                "Terminalde şunu çalıştırın:\n    pip install python-docx")
            return False

        doc = Document(template_path) if template_path else Document()

        if template_path:
            mapping = {"{{PROJE_ADI}}": proje_ismi, "{{TARIH}}": datetime.now().strftime("%d.%m.%Y")}
            for para in doc.paragraphs:
                for k, v in mapping.items():
                    if k in para.text:
                        para.text = para.text.replace(k, v)
            doc.add_page_break()
        else:
            doc.add_heading("V-MODEL TEKNİK DOKÜMAN PAKETİ", level=0)
            doc.add_paragraph(f"Proje: {proje_ismi}")
            doc.add_paragraph(f"Tarih: {datetime.now().strftime('%d.%m.%Y')}")

        for type_key, title, leg in self.VMODEL_SECTIONS:
            items = [d for d in self.flat_data.values() if d.get("type") == type_key]
            if not items:
                continue
            doc.add_heading(title, level=2)
            t = doc.add_table(rows=1, cols=3)
            try:
                t.style = "Light Grid Accent 1"
            except Exception:
                t.style = "Table Grid"
            h = t.rows[0].cells
            h[0].text, h[1].text, h[2].text = "ID", "Açıklama", "Bağlı Olduğu"
            for d in items:
                c = t.add_row().cells
                c[0].text = str(d.get("ID", ""))
                c[1].text = str(d.get("content", ""))
                c[2].text = str(d.get("bound_to", "Yok"))
        doc.save(path)
        return True

    def _export_doors_csv(self, path, proje_ismi):
        """V-Model maddelerini IBM DOORS'a aktarılabilir UTF-8 CSV olarak üretir."""
        leg_names = {"req": "Requirement", "test": "Test"}

        with open(path, "w", encoding="utf-8-sig", newline="") as csv_file:
            writer = csv.writer(
                csv_file,
                dialect="excel",
                quoting=csv.QUOTE_ALL,
                lineterminator="\r\n",
            )
            writer.writerow([
                "Object Identifier",
                "Object Text",
                "Requirement Type",
                "V-Model Leg",
                "Linked Object",
                "Project",
            ])

            for type_key, _title, leg in self.VMODEL_SECTIONS:
                for item in self.flat_data.values():
                    if item.get("type") != type_key:
                        continue
                    writer.writerow([
                        item.get("ID", ""),
                        item.get("content", ""),
                        type_key,
                        leg_names.get(leg, leg),
                        item.get("bound_to", "Yok"),
                        proje_ismi,
                    ])

    def download_docs(self):
        fmt = self.format_combo.get().lower()
        if fmt != "şablon" and not self.last_generated_output and not self.flat_data:
            messagebox.showwarning("Uyarı", "İndirilecek içerik yok. Önce doküman üretin.")
            return

        proje_ismi = self.entry_widgets["proje_ismi"].get().strip().replace(" ", "_") or "yapay_zeka_dokumanlar"

        if fmt == "excel":
            ext, ftypes = ".xlsx", [("Excel Workbook", "*.xlsx")]
        elif fmt == "doors":
            ext, ftypes = ".csv", [("IBM DOORS CSV", "*.csv")]
        elif fmt in ("şablon", "docx"):
            ext, ftypes = ".docx", [("Word Document", "*.docx")]
        else:
            ext, ftypes = f".{fmt}", [(f"{fmt.upper()} files", f"*.{fmt}")]

        init_name = f"{proje_ismi}_Dokumanlar{ext}"
        path = filedialog.asksaveasfilename(
            defaultextension=ext,
            filetypes=ftypes + [("All files", "*.*")],
            initialfile=init_name
        )
        if not path:
            return

        try:
            if fmt == "txt":
                with open(path, "w", encoding="utf-8") as f:
                    f.write(self.last_generated_output or "")

            elif fmt == "pdf":
                self._generate_vmodel_pdf(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "excel":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Excel'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_excel(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "html":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "HTML'e aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_html(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje")

            elif fmt == "doors":
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "DOORS'a aktarılacak veri yok. Önce doküman üretin.")
                    return
                self._export_doors_csv(
                    path,
                    self.entry_widgets["proje_ismi"].get().strip() or "Proje"
                )

            elif fmt in ("şablon", "docx"):
                if not self.flat_data:
                    messagebox.showwarning("Uyarı", "Word'e yerleştirilecek veri yok. Önce doküman üretin.")
                    return
                tmpl = self.template_file_path if fmt == "şablon" else None
                if fmt == "şablon" and not tmpl:
                    messagebox.showerror("Hata", "Şablon dosyası seçilmedi (.docx).")
                    return
                ok = self._export_docx(path, self.entry_widgets["proje_ismi"].get().strip() or "Proje", tmpl)
                if not ok:
                    return

            resolved_path = os.path.abspath(path)
            if resolved_path not in self.generated_document_paths:
                self.generated_document_paths.append(resolved_path)
            messagebox.showinfo("Başarılı", f"Dosya kaydedildi:\n{path}")
        except Exception as e:
            messagebox.showerror("Hata", f"Kaydetme hatası: {e}")
            self.update_status_text(f"Kaydetme hatası: {e}", is_error=True)

